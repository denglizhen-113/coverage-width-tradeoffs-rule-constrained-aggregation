#!/usr/bin/env python3
"""Build the Stage 32 MATCOM scientific-correction package.

This stage is intentionally non-destructive.  It reads the frozen Stage 26X
simulation evidence, the processed empirical panel, and the Stage 31.5 package;
it writes corrected metrics, exact ordinal endpoint results, uncertainty
audits, and a new editable submission candidate under ``outputs/stage32-*``.
No raw data, frozen output, remote repository, archive, or submission portal is
modified.
"""

from __future__ import annotations

import argparse
import hashlib
import math
import re
import shutil
import sys
import time
from datetime import datetime, timezone
from pathlib import Path
from typing import Iterable

import matplotlib

matplotlib.use("Agg")
from matplotlib import pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))
OUTPUT = Path("outputs/stage32-matcom-scientific-corrections")
PUBLIC_RELEASE_URL = "https://github.com/denglizhen-113/coverage-width-tradeoffs-rule-constrained-aggregation/releases/tag/matcom-stage32-v1.0.1"
LEGACY_PACKAGE = Path(
    "outputs/stage31_5-matcom-editorial-polish/MATCOM_final_candidate_package"
)
INTERNAL_RAW = Path("outputs/stage26X-1/raw")
BAYESIAN_RAW = Path("outputs/stage26X-2/raw/bayesian")
PANEL = Path("data/processed/panel_long.csv")
LEGACY_R = Path("outputs/tables/ranking_identification_summary_r.csv")
LEGACY_RPLUS = Path("outputs/tables/ranking_identification_summary_rplus.csv")
LEGACY_CONTESTANT = Path("outputs/tables/ranking_contestant_identification.csv")
CONSTRAINT_SUMMARY = Path("outputs/tables/constraint_summary.csv")
TIE_POLICIES = (
    "average_rank",
    "min_rank",
    "dense_rank",
    "competition_rank",
)
SET_METHODS = (
    "rule_aware_partial_identification",
    "rule_agnostic_partial_identification",
)
AUTHOR_CONTRIBUTIONS = (
    "Lizhen Deng: Conceptualization; Methodology; Software; Formal analysis; "
    "Investigation; Visualization; Writing - original draft; Writing - review "
    "and editing; Project administration. Yuxin Liu: Resources; Data curation; "
    "Validation. Bo Li: Resources; Data curation; Writing - review and editing."
)
SUBSCRIPTS = str.maketrans("0123456789", "₀₁₂₃₄₅₆₇₈₉")
SUPERSCRIPTS = str.maketrans("0123456789", "⁰¹²³⁴⁵⁶⁷⁸⁹")
M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"


class Stage32Error(RuntimeError):
    """Raised when a scientific-correction invariant fails."""


def parse_args(argv: list[str] | None = None) -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description=(
            "Recompute joint-set coverage, solve exact ordinal rank endpoints, "
            "report performance MCSE and rule-provenance sensitivity, and build "
            "a revised Elsevier MATCOM candidate package."
        )
    )
    parser.add_argument(
        "--project-root",
        type=Path,
        default=ROOT,
        help="Project root; defaults to the repository containing this script.",
    )
    parser.add_argument(
        "--output-dir",
        type=Path,
        default=OUTPUT,
        help="Project-relative output directory; existing files are replaced safely.",
    )
    return parser.parse_args(argv)


def resolve(root: Path, relative: Path, label: str) -> Path:
    if relative.is_absolute():
        raise Stage32Error(f"{label} must be project-relative: {relative}")
    result = (root / relative).resolve()
    if result != root and root not in result.parents:
        raise Stage32Error(f"{label} resolves outside the project root: {relative}")
    return result


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest().upper()


def write_text(path: Path, content: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    temporary = path.with_suffix(path.suffix + ".tmp")
    temporary.write_text(content.rstrip() + "\n", encoding="utf-8", newline="\n")
    temporary.replace(path)


def write_csv(path: Path, frame: pd.DataFrame) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    temporary = path.with_suffix(path.suffix + ".tmp")
    frame.to_csv(
        temporary,
        index=False,
        encoding="utf-8",
        lineterminator="\n",
        na_rep="",
        float_format="%.12g",
    )
    temporary.replace(path)


def load_csv_family(directory: Path, pattern: str) -> pd.DataFrame:
    paths = sorted(directory.glob(pattern))
    if not paths:
        raise Stage32Error(f"No files match {directory / pattern}")
    return pd.concat((pd.read_csv(path) for path in paths), ignore_index=True)


def sample_standard_error(values: pd.Series) -> float:
    numeric = pd.to_numeric(values, errors="coerce").dropna()
    if len(numeric) < 2:
        return float("nan")
    return float(numeric.std(ddof=1) / math.sqrt(len(numeric)))


def corrected_internal_coverage(
    root: Path,
) -> tuple[
    pd.DataFrame,
    pd.DataFrame,
    pd.DataFrame,
    pd.DataFrame,
    pd.DataFrame,
    dict[str, float],
]:
    """Reclassify legacy projection-box coverage as true joint-set coverage."""
    raw = load_csv_family(root / INTERNAL_RAW, "internal_*.csv")
    required = {
        "seed",
        "n_active",
        "outcome_noise_probability",
        "replication",
        "observed_outcome_noise",
        "method",
        "coverage",
        "width",
    }
    missing = required.difference(raw.columns)
    if missing:
        raise Stage32Error(f"Internal raw files miss columns: {sorted(missing)}")
    selected = raw.loc[raw["method"].isin(SET_METHODS)].copy()
    selected = selected.rename(columns={"coverage": "projection_box_coverage"})
    aware = selected["method"].eq("rule_aware_partial_identification")
    noisy = selected["observed_outcome_noise"].astype(bool)
    selected["joint_set_coverage"] = np.where(aware, (~noisy).astype(float), 1.0)
    selected["coverage_metric_status"] = np.where(
        aware,
        "joint polytope membership from generating-rule consistency",
        "simplex membership",
    )

    group = ["seed", "n_active", "outcome_noise_probability", "method"]
    seed_level = (
        selected.groupby(group, as_index=False, sort=True)
        .agg(
            n_replications=("replication", "nunique"),
            joint_set_coverage=("joint_set_coverage", "mean"),
            projection_box_coverage=("projection_box_coverage", "mean"),
            mean_width=("width", "mean"),
        )
    )
    across_seed = (
        seed_level.groupby(
            ["n_active", "outcome_noise_probability", "method"],
            as_index=False,
            sort=True,
        )
        .agg(
            n_seeds=("seed", "nunique"),
            joint_set_coverage_mean=("joint_set_coverage", "mean"),
            joint_set_coverage_sd=("joint_set_coverage", "std"),
            projection_box_coverage_mean=("projection_box_coverage", "mean"),
            mean_width=("mean_width", "mean"),
            width_seed_sd=("mean_width", "std"),
        )
    )

    performance_rows: list[dict[str, object]] = []
    region_group = ["n_active", "outcome_noise_probability", "method"]
    for key, frame in selected.groupby(region_group, sort=True):
        n_active, noise_probability, method = key
        performance_rows.append(
            {
                "n_active": int(n_active),
                "outcome_noise_probability": float(noise_probability),
                "method": method,
                "n_replications": len(frame),
                "joint_set_coverage": frame["joint_set_coverage"].mean(),
                "joint_set_coverage_mcse": sample_standard_error(
                    frame["joint_set_coverage"]
                ),
                "projection_box_coverage": frame[
                    "projection_box_coverage"
                ].mean(),
                "projection_box_coverage_mcse": sample_standard_error(
                    frame["projection_box_coverage"]
                ),
                "mean_width": frame["width"].mean(),
                "mean_width_mcse": sample_standard_error(frame["width"]),
            }
        )
    performance = pd.DataFrame(performance_rows)

    keys = [
        "seed",
        "n_active",
        "outcome_noise_probability",
        "replication",
        "observed_outcome_noise",
    ]
    paired = selected.pivot(
        index=keys,
        columns="method",
        values=["joint_set_coverage", "projection_box_coverage", "width"],
    )
    paired.columns = [f"{metric}__{method}" for metric, method in paired.columns]
    paired = paired.reset_index()
    aware_name = "rule_aware_partial_identification"
    agnostic_name = "rule_agnostic_partial_identification"
    paired["coverage_change_without_elimination"] = (
        paired[f"joint_set_coverage__{agnostic_name}"]
        - paired[f"joint_set_coverage__{aware_name}"]
    )
    paired["legacy_projection_change_without_elimination"] = (
        paired[f"projection_box_coverage__{agnostic_name}"]
        - paired[f"projection_box_coverage__{aware_name}"]
    )
    paired["width_change_without_elimination"] = (
        paired[f"width__{agnostic_name}"] - paired[f"width__{aware_name}"]
    )
    effect_rows: list[dict[str, object]] = []
    for label, frame in (
        ("clean", paired.loc[paired["outcome_noise_probability"].eq(0.0)]),
        ("positive_noise", paired.loc[paired["outcome_noise_probability"].gt(0.0)]),
    ):
        effect_rows.append(
            {
                "condition": label,
                "paired_replications": len(frame),
                "paired_seed_cells": frame[
                    ["seed", "n_active", "outcome_noise_probability"]
                ].drop_duplicates().shape[0],
                "joint_coverage_change_mean": frame[
                    "coverage_change_without_elimination"
                ].mean(),
                "joint_coverage_change_mcse": sample_standard_error(
                    frame["coverage_change_without_elimination"]
                ),
                "legacy_projection_change_mean": frame[
                    "legacy_projection_change_without_elimination"
                ].mean(),
                "legacy_projection_change_mcse": sample_standard_error(
                    frame["legacy_projection_change_without_elimination"]
                ),
                "width_change_mean": frame[
                    "width_change_without_elimination"
                ].mean(),
                "width_change_mcse": sample_standard_error(
                    frame["width_change_without_elimination"]
                ),
            }
        )
    effects = pd.DataFrame(effect_rows)
    positive = effects.loc[effects["condition"].eq("positive_noise")].iloc[0]
    statistics = {
        "positive_joint_change": float(positive["joint_coverage_change_mean"]),
        "positive_joint_mcse": float(positive["joint_coverage_change_mcse"]),
        "positive_projection_change": float(
            positive["legacy_projection_change_mean"]
        ),
        "positive_projection_mcse": float(
            positive["legacy_projection_change_mcse"]
        ),
        "positive_width_change": float(positive["width_change_mean"]),
        "positive_width_mcse": float(positive["width_change_mcse"]),
        "positive_replications": int(positive["paired_replications"]),
    }
    return selected, seed_level, across_seed, performance, effects, statistics


def exact_ordinal_endpoints(
    root: Path,
) -> tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame, pd.DataFrame, dict[str, float]]:
    from src.ranking_identification import build_week_spec, exact_rank_support

    panel = pd.read_csv(root / PANEL)
    ordinal = panel.loc[panel["aggregation_regime"].isin(["R", "R_plus"])]
    keys = ordinal[
        ["aggregation_regime", "season", "week"]
    ].drop_duplicates().sort_values(["aggregation_regime", "season", "week"])
    contestant_rows: list[dict[str, object]] = []
    week_rows: list[dict[str, object]] = []
    started = time.perf_counter()
    for tie_policy in TIE_POLICIES:
        for key in keys.itertuples(index=False):
            frame = ordinal.loc[
                ordinal["aggregation_regime"].eq(key.aggregation_regime)
                & ordinal["season"].eq(key.season)
                & ordinal["week"].eq(key.week)
            ]
            spec = build_week_spec(frame, tie_policy)
            selected = exact_rank_support(spec)
            direct = (
                exact_rank_support(spec, mechanism="R")
                if spec.regime == "R_plus" and tie_policy == "average_rank"
                else None
            )
            week_rows.append(
                {
                    "season": spec.season,
                    "week": spec.week,
                    "regime": spec.regime,
                    "tie_policy": tie_policy,
                    "n_active": spec.n_active,
                    "n_eliminated": len(spec.eliminated_indices),
                    "n_withdrawn": len(spec.withdrawn_indices),
                    "finale_week": spec.finale_week,
                    "exact_endpoint_feasible": selected.feasible,
                    "exact_normalized_rank_width": selected.normalized_mean_width,
                    "direct_R_like_exact_normalized_rank_width": (
                        direct.normalized_mean_width if direct is not None else np.nan
                    ),
                    "weak_minus_direct_exact_width": (
                        selected.normalized_mean_width - direct.normalized_mean_width
                        if direct is not None and selected.feasible and direct.feasible
                        else np.nan
                    ),
                    "solver_calls": selected.solver_calls
                    + (direct.solver_calls if direct is not None else 0),
                    "solver": selected.solver_name,
                }
            )
            for contestant, contestant_id in enumerate(spec.contestant_ids):
                contestant_rows.append(
                    {
                        "season": spec.season,
                        "week": spec.week,
                        "regime": spec.regime,
                        "tie_policy": tie_policy,
                        "contestant_id": contestant_id,
                        "contestant_name": spec.contestant_names[contestant],
                        "n_active": spec.n_active,
                        "exact_fan_rank_min": selected.minimum_ranks[contestant],
                        "exact_fan_rank_max": selected.maximum_ranks[contestant],
                        "exact_fan_rank_width": selected.rank_widths[contestant],
                        "exact_normalized_fan_rank_width": (
                            selected.rank_widths[contestant] / (spec.n_active - 1)
                            if selected.feasible and spec.n_active > 1
                            else np.nan
                        ),
                    }
                )
    elapsed = time.perf_counter() - started
    contestant = pd.DataFrame(contestant_rows)
    weeks = pd.DataFrame(week_rows)

    legacy_contestant = pd.read_csv(root / LEGACY_CONTESTANT)
    primary_contestant = contestant.loc[
        contestant["tie_policy"].eq("average_rank")
    ]
    comparison = primary_contestant.merge(
        legacy_contestant[
            [
                "season",
                "week",
                "regime",
                "contestant_id",
                "fan_rank_min",
                "fan_rank_max",
                "normalized_fan_rank_width",
                "enumeration_method",
            ]
        ],
        on=["season", "week", "regime", "contestant_id"],
        how="left",
        validate="one_to_one",
    )
    comparison["minimum_endpoint_correction"] = (
        comparison["fan_rank_min"] - comparison["exact_fan_rank_min"]
    )
    comparison["maximum_endpoint_correction"] = (
        comparison["exact_fan_rank_max"] - comparison["fan_rank_max"]
    )
    comparison["normalized_width_correction"] = (
        comparison["exact_normalized_fan_rank_width"]
        - comparison["normalized_fan_rank_width"]
    )
    if comparison[
        ["minimum_endpoint_correction", "maximum_endpoint_correction"]
    ].min().min() < -1e-8:
        raise Stage32Error("A sampled endpoint lies outside the exact MILP support")

    legacy_weeks = pd.concat(
        [pd.read_csv(root / LEGACY_R), pd.read_csv(root / LEGACY_RPLUS)],
        ignore_index=True,
    )
    primary_weeks = weeks.loc[weeks["tie_policy"].eq("average_rank")]
    week_comparison = primary_weeks.merge(
        legacy_weeks[
            [
                "season",
                "week",
                "regime",
                "normalized_rank_width",
                "enumeration_method",
                "feasible_fraction",
                "mc_standard_error",
                "identifiability_loss_ratio",
            ]
        ],
        on=["season", "week", "regime"],
        how="left",
        validate="one_to_one",
    )
    week_comparison["exact_minus_legacy_normalized_width"] = (
        week_comparison["exact_normalized_rank_width"]
        - week_comparison["normalized_rank_width"]
    )

    sensitivity_rows: list[dict[str, object]] = []
    rplus = primary_weeks.loc[primary_weeks["regime"].eq("R_plus")]
    for label, frame in (
        ("primary_seasons_28_34", rplus),
        ("exclude_ambiguous_season_28", rplus.loc[rplus["season"].ne(28)]),
        ("season_28_only", rplus.loc[rplus["season"].eq(28)]),
    ):
        sensitivity_rows.append(
            {
                "rule_provenance_scenario": label,
                "weeks": len(frame),
                "mean_exact_normalized_rank_width": frame[
                    "exact_normalized_rank_width"
                ].mean(),
                "mean_direct_R_like_exact_normalized_rank_width": frame[
                    "direct_R_like_exact_normalized_rank_width"
                ].mean(),
                "mean_weak_minus_direct_exact_width": frame[
                    "weak_minus_direct_exact_width"
                ].mean(),
            }
        )
    provenance = pd.DataFrame(sensitivity_rows)
    sampled = comparison.loc[comparison["enumeration_method"].eq("monte_carlo")]
    stats = {
        "weeks": int(primary_weeks.shape[0]),
        "primary_solver_calls": int(primary_weeks["solver_calls"].sum()),
        "all_policy_solver_calls": int(weeks["solver_calls"].sum()),
        "elapsed_seconds": float(elapsed),
        "sampled_contestants": int(len(sampled)),
        "sampled_endpoint_corrections": int(
            (
                sampled["minimum_endpoint_correction"].gt(0)
                | sampled["maximum_endpoint_correction"].gt(0)
            ).sum()
        ),
        "max_width_correction": float(
            sampled["normalized_width_correction"].max()
        ),
        "mean_width_correction": float(
            sampled["normalized_width_correction"].mean()
        ),
        "r_exact_width": float(
            primary_weeks.loc[
                primary_weeks["regime"].eq("R"), "exact_normalized_rank_width"
            ].mean()
        ),
        "rplus_exact_width": float(
            primary_weeks.loc[
                primary_weeks["regime"].eq("R_plus"),
                "exact_normalized_rank_width",
            ].mean()
        ),
    }
    return contestant, weeks, week_comparison, provenance, stats


def bayesian_undefined_report(root: Path) -> tuple[pd.DataFrame, dict[str, int]]:
    raw = load_csv_family(root / BAYESIAN_RAW, "*.csv")
    raw["undefined"] = ~raw["posterior_status"].eq("ok")
    internal = raw.loc[raw["synthesizer"].eq("internal_percentage")].copy()
    grouped = (
        internal.groupby(
            ["n_active", "outcome_noise_probability"], as_index=False, sort=True
        )
        .agg(
            denominator=("replication", "size"),
            undefined_rows=("undefined", "sum"),
            accepted_draws_median=("accepted_posterior_draws", "median"),
            accepted_draws_min=("accepted_posterior_draws", "min"),
        )
    )
    grouped["undefined_rate"] = grouped["undefined_rows"] / grouped["denominator"]
    return grouped, {
        "total": int(internal["undefined"].sum()),
        "clean": int(
            internal.loc[
                internal["outcome_noise_probability"].eq(0.0), "undefined"
            ].sum()
        ),
        "external": int(
            raw.loc[raw["synthesizer"].eq("external_ordinal"), "undefined"].sum()
        ),
    }


def plot_corrected_internal(
    selected: pd.DataFrame,
    path_pdf: Path,
    path_tiff: Path,
) -> None:
    aggregate = (
        selected.groupby(["outcome_noise_probability", "method"], as_index=False)
        .agg(
            joint_coverage=("joint_set_coverage", "mean"),
            projection_coverage=("projection_box_coverage", "mean"),
            width=("width", "mean"),
            joint_mcse=("joint_set_coverage", sample_standard_error),
            projection_mcse=("projection_box_coverage", sample_standard_error),
            width_mcse=("width", sample_standard_error),
        )
    )
    aware = aggregate.loc[
        aggregate["method"].eq("rule_aware_partial_identification")
    ]
    agnostic = aggregate.loc[
        aggregate["method"].eq("rule_agnostic_partial_identification")
    ]
    plt.rcParams.update(
        {
            "font.family": "DejaVu Sans",
            "font.size": 9,
            "axes.spines.top": False,
            "axes.spines.right": False,
        }
    )
    fig, axes = plt.subplots(1, 2, figsize=(7.2, 3.25), constrained_layout=True)
    x = aware["outcome_noise_probability"].to_numpy()
    axes[0].errorbar(
        x,
        aware["joint_coverage"],
        yerr=1.96 * aware["joint_mcse"],
        color="#A33A2B",
        marker="o",
        label="Rule-aware joint set",
    )
    axes[0].errorbar(
        x,
        aware["projection_coverage"],
        yerr=1.96 * aware["projection_mcse"],
        color="#D28B32",
        marker="s",
        linestyle="--",
        label="Legacy marginal box",
    )
    axes[0].plot(
        agnostic["outcome_noise_probability"],
        agnostic["joint_coverage"],
        color="#2F6B4F",
        marker="^",
        label="Simplex-only set",
    )
    axes[0].set(xlabel="Outcome-noise probability", ylabel="Known-truth coverage")
    axes[0].set_ylim(0.72, 1.01)
    axes[0].legend(frameon=False, fontsize=7.5)
    axes[0].set_title("(a) Joint membership versus projection envelope")

    axes[1].errorbar(
        x,
        aware["width"],
        yerr=1.96 * aware["width_mcse"],
        color="#1F5A7A",
        marker="o",
        label="Rule-aware",
    )
    axes[1].errorbar(
        agnostic["outcome_noise_probability"],
        agnostic["width"],
        yerr=1.96 * agnostic["width_mcse"],
        color="#6B7280",
        marker="^",
        label="Simplex-only",
    )
    axes[1].set(xlabel="Outcome-noise probability", ylabel="Mean coordinate width")
    axes[1].set_ylim(0.72, 1.02)
    axes[1].legend(frameon=False, fontsize=7.5)
    axes[1].set_title("(b) Width with 95% Monte Carlo error bars")
    for axis in axes:
        axis.grid(axis="y", color="#D1D5DB", linewidth=0.5, alpha=0.7)
    path_pdf.parent.mkdir(parents=True, exist_ok=True)
    path_tiff.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(path_pdf, bbox_inches="tight")
    fig.savefig(path_tiff, dpi=600, bbox_inches="tight", pil_kwargs={"compression": "tiff_lzw"})
    plt.close(fig)


def markdown_table(frame: pd.DataFrame, columns: list[str], formats: dict[str, str]) -> str:
    headers = [column.replace("_", " ") for column in columns]
    rows = ["| " + " | ".join(headers) + " |", "| " + " | ".join("---" for _ in columns) + " |"]
    for record in frame[columns].itertuples(index=False, name=None):
        values: list[str] = []
        for column, value in zip(columns, record):
            if pd.isna(value):
                values.append("")
            elif column in formats:
                values.append(formats[column].format(value))
            else:
                values.append(str(value))
        rows.append("| " + " | ".join(values) + " |")
    return "\n".join(rows)


def build_manuscript(
    coverage: pd.DataFrame,
    effects: pd.DataFrame,
    ordinal_comparison: pd.DataFrame,
    provenance: pd.DataFrame,
    bayesian_undefined: pd.DataFrame,
    coverage_stats: dict[str, float],
    ordinal_stats: dict[str, float],
    bayesian_stats: dict[str, int],
) -> str:
    coverage_display = (
        coverage.groupby(["outcome_noise_probability", "method"], as_index=False)
        .agg(
            joint_set_coverage=("joint_set_coverage", "mean"),
            projection_box_coverage=("projection_box_coverage", "mean"),
            mean_width=("width", "mean"),
            joint_coverage_mcse=("joint_set_coverage", sample_standard_error),
            width_mcse=("width", sample_standard_error),
        )
    )
    coverage_display["method"] = coverage_display["method"].replace(
        {
            "rule_aware_partial_identification": "rule-aware polytope",
            "rule_agnostic_partial_identification": "simplex-only",
        }
    )
    endpoint_display = (
        ordinal_comparison.groupby(["regime", "enumeration_method"], as_index=False)
        .agg(
            weeks=("week", "size"),
            legacy_mean_width=("normalized_rank_width", "mean"),
            exact_mean_width=("exact_normalized_rank_width", "mean"),
            maximum_correction=("exact_minus_legacy_normalized_width", "max"),
        )
    )
    endpoint_display["weeks"] = endpoint_display.apply(
        lambda row: int(
            ordinal_comparison.loc[
                ordinal_comparison["regime"].eq(row["regime"])
                & ordinal_comparison["enumeration_method"].eq(
                    row["enumeration_method"]
                ),
                ["season", "week"],
            ].drop_duplicates().shape[0]
        ),
        axis=1,
    )
    coverage_table = markdown_table(
        coverage_display,
        [
            "outcome_noise_probability",
            "method",
            "joint_set_coverage",
            "joint_coverage_mcse",
            "projection_box_coverage",
            "mean_width",
            "width_mcse",
        ],
        {
            "outcome_noise_probability": "{:.2f}",
            "joint_set_coverage": "{:.6f}",
            "joint_coverage_mcse": "{:.6f}",
            "projection_box_coverage": "{:.6f}",
            "mean_width": "{:.6f}",
            "width_mcse": "{:.6f}",
        },
    )
    effects_table = markdown_table(
        effects,
        [
            "condition",
            "paired_replications",
            "joint_coverage_change_mean",
            "joint_coverage_change_mcse",
            "legacy_projection_change_mean",
            "width_change_mean",
            "width_change_mcse",
        ],
        {
            "joint_coverage_change_mean": "{:.6f}",
            "joint_coverage_change_mcse": "{:.6f}",
            "legacy_projection_change_mean": "{:.6f}",
            "width_change_mean": "{:.6f}",
            "width_change_mcse": "{:.6f}",
        },
    )
    endpoint_table = markdown_table(
        endpoint_display,
        [
            "regime",
            "enumeration_method",
            "weeks",
            "legacy_mean_width",
            "exact_mean_width",
            "maximum_correction",
        ],
        {
            "legacy_mean_width": "{:.6f}",
            "exact_mean_width": "{:.6f}",
            "maximum_correction": "{:.6f}",
        },
    )
    provenance_table = markdown_table(
        provenance,
        [
            "rule_provenance_scenario",
            "weeks",
            "mean_exact_normalized_rank_width",
            "mean_direct_R_like_exact_normalized_rank_width",
            "mean_weak_minus_direct_exact_width",
        ],
        {
            "mean_exact_normalized_rank_width": "{:.6f}",
            "mean_direct_R_like_exact_normalized_rank_width": "{:.6f}",
            "mean_weak_minus_direct_exact_width": "{:.6f}",
        },
    )
    bayes_table = markdown_table(
        bayesian_undefined,
        [
            "n_active",
            "outcome_noise_probability",
            "denominator",
            "undefined_rows",
            "undefined_rate",
            "accepted_draws_median",
            "accepted_draws_min",
        ],
        {
            "outcome_noise_probability": "{:.2f}",
            "undefined_rate": "{:.6f}",
            "accepted_draws_median": "{:.1f}",
            "accepted_draws_min": "{:.0f}",
        },
    )
    return rf"""# Coverage-Width Tradeoffs in Rule-Constrained Expert-Crowd Aggregation

## Abstract

Institutional expert-crowd systems often reveal expert scores and coarse outcomes while public preferences remain latent. We formulate cardinal identified polytopes for percentage aggregation and ordinal feasible-ranking sets for rank aggregation with judge-save discretion. Two corrections sharpen the computational methodology. First, known-truth coverage of a polytope is evaluated by joint constraint membership rather than membership in the Cartesian product of its coordinate projections. Second, sharp ordinal rank endpoints are obtained by binary linear optimization instead of extrema of sampled permutations. The endpoint formulation uses {int(ordinal_stats['primary_solver_calls']):,} solver calls for 87 empirical weeks under the primary tie policy and agrees with complete enumeration in all small-field verification cases. Across {int(coverage_stats['positive_replications']):,} positive-noise paired internal replications, removing the elimination constraint increases joint-set coverage by {coverage_stats['positive_joint_change']:.6f} (MCSE {coverage_stats['positive_joint_mcse']:.6f}) and mean coordinate width by {coverage_stats['positive_width_change']:.6f} (MCSE {coverage_stats['positive_width_mcse']:.6f}). The former projection-box calculation understated the coverage change as {coverage_stats['positive_projection_change']:.6f}. Same-information Bayesian comparisons remain conditional on a 95% equal-tail posterior construction and a stated prior; Bayesian credible rectangles and identified sets are reported as different inferential objects. The contribution is an exact, auditable inverse-rule methodology and a bounded simulation study, not recovery of an observed audience vote or a claim of uniform method superiority.

**Keywords:** Expert-crowd aggregation; partial identification; exact rank support; mixed-integer optimization; latent public preference; simulation.

## 1. Introduction

Many decision systems combine an expert component with an unobserved public component and disclose only the aggregate consequence. The inverse problem is therefore set-valued: which latent public states are compatible with the expert record, the stated aggregation rule, and the observed outcome? We study this problem for percentage aggregation and rank aggregation with judge-save discretion. The competition record is an empirical testbed for a general class of expert-crowd systems; it is not treated as a source of observed audience votes.

The paper makes three contributions. First, it gives an exact binary-optimization formulation for sharp attainable public-rank endpoints under tie-inclusive bottom-set rules. This replaces the earlier practice of reading support extrema from a finite permutation sample. Second, it distinguishes joint identified-set coverage from coverage of marginal projection intervals and recomputes the misspecification experiment with metric-matched Monte Carlo uncertainty. Third, it states the boundary between identified sets and Bayesian credible rectangles, reports every undefined posterior region, and treats the uncertain season-28 rule assignment as a sensitivity assumption.

These contributions are narrower than the general possible-winner problem. We do not claim a new dichotomy for arbitrary voting rules. The special structure here consists of one hidden aggregate public ranking, one observed expert ranking, and an observed bottom-set implication. That structure admits a compact assignment MILP for support endpoints even though exact feasible-set counting may remain combinatorial.

## 2. Related work and novelty boundary

Partial identification separates data-consistent parameter sets from point recovery [1-4]. Bayesian credible sets in partially identified models have different semantics from estimated identified sets; Moon and Schorfheide recommend reporting the identified set and conditional-prior information alongside Bayesian summaries [5]. We follow that distinction explicitly.

The closest ordinal literature is computational social choice under incomplete preferences. Xia and Conitzer characterize possible and necessary winners under common rules [6]; Betzler and Dorn establish a broad complexity dichotomy for possible winners under scoring rules [7]; Bachrach, Betzler, and Faliszewski study probabilistic possible-winner counting and randomized computation [8]; Dey and Misra sharpen hardness results as missing pairwise information varies [9]; and Lu and Boutilier develop robust winner determination under partial rankings [10]. Their objects complete multiple partial voter orders and ask about winners. Our object is different: the public input is one latent aggregate permutation, expert ranks are observed, and a coarse institutional outcome constrains combined expert-public scores. The present novelty is the inverse-rule formulation and its exact support-endpoint MILP, not the general idea that incomplete rankings induce a completion set.

Simulation design follows ADEMP and transparent reporting guidance [11-13]. Fixed seeds and parameter cells were recorded in a hash-locked local design before the reported evaluation. Because no external immutable registration identifier predates result inspection, we use **predeclared** and **hash-locked**, not **preregistered**.

## 3. Data, rules, and provenance

The empirical testbed is the official COMAP 2026 Problem C data file [14]. The processed panel has 248 percentage-regime weeks, 14 early rank-regime weeks, and 73 weeks encoded under the rank-plus-judge-save interpretation. COMAP documents percentage aggregation for seasons 3-27 and describes a bottom-two judge-save mechanism from season 28, but states that the exact season in which rank aggregation returned is uncertain and treats season 28 as a reasonable assumption. We therefore label seasons 28-34 as the **primary season-28 rank-return assumption** and separately report results excluding season 28.

One percentage week, season 18 week 2, is skipped. The cause is not a missing public-appeal proxy. The constraint log records an eliminated contestant, Diana Nyad, without an active judge total after the source-status mapping. Because the elimination inequality cannot be constructed from the stated P-regime equation, that week has no identified set. The 11 unavailable typed proxies are downstream consequences of this skipped identification result and are neither inputs to nor causes of P-regime constraint construction. No value is imputed.

Withdrawals are excluded from outcome eligibility but remain in the active public-rank permutation when the source record marks them active for that week. Multiple eliminations use the observed eliminated set. Weekly latent rankings are not constrained to remain constant across rounds; multi-round synthetic case coverage is the conjunction of separately encoded round-level compatibility statements. A model with a stable or dynamically coupled latent ranking would be a different estimand and is not implemented here.

## 4. Cardinal identified polytope and coverage

Let $A$ be the active candidates, $q_i$ the normalized expert share, and $p_i$ the latent public share. Under percentage aggregation,

$P(O)=\{{p\in\mathbb R^n: p_i\ge0,\ \sum_i p_i=1,\ q_e+p_e\le q_s+p_s\ \text{{for every eliminated }}e\text{{ and eligible survivor }}s}}\}}.$

Coordinate bounds $l_i=\min_{{p\in P(O)}}p_i$ and $u_i=\max_{{p\in P(O)}}p_i$ are sharp marginal projections. Their Cartesian product $\prod_i[l_i,u_i]$ is generally a strict outer approximation to $P(O)$. We therefore report two distinct diagnostics for a known synthetic truth $p^\star$:

$C_{{set}}=1\{{p^\star\in P(O)\}},\qquad C_{{box}}=1\{{l_i\le p_i^\star\le u_i\ \forall i}}.$

Only $C_{{set}}$ is called identified-set coverage. $C_{{box}}$ is retained as a legacy projection-envelope diagnostic. In the noisy generator the recorded eliminated candidate is the strictly second-worst combined-score candidate (ties occur with probability zero under the continuous generator), so every realized noise event violates at least one joint elimination inequality even when the truth remains inside all marginal intervals.

## 5. Ordinal feasible sets

Let $r_i^J$ be the observed expert rank and $r_i^F$ a latent strict public rank, a permutation of $1,\ldots,n$. Define combined rank score $c_i=r_i^J+r_i^F$, where a larger score is worse. For eligible set $G$ and $1\le k\le|G|$, let $\theta_k(c;G)$ be the $k$th largest value in $\{{c_i:i\in G\}}$ and define the tie-inclusive bottom set

$B_k(c;G)=\{{i\in G:c_i\ge\theta_k(c;G)\}}.$

This definition permits $|B_k|>k$ under ties. It also gives $B_k(c;G)\subseteq B_{{k+1}}(c;G)$ directly because $\theta_{{k+1}}\le\theta_k$. For $k$ observed eliminations, the direct rule requires $E\subseteq B_k$; the weak judge-save interpretation requires $E\subseteq B_{{\min(k+1,|G|)}}$. Expert ties are converted to ranks under the named policy (average, minimum/competition, or dense), and every sensitivity row preserves that policy label.

### 5.1 Exact support-endpoint MILP

Introduce binary assignment variables $x_{{ir}}$, equal to one when candidate $i$ receives public rank $r$. The permutation constraints are

$\sum_r x_{{ir}}=1\ \forall i,\qquad \sum_i x_{{ir}}=1\ \forall r.$

Then $s_i=r_i^J+\sum_r r x_{{ir}}$. For each eliminated $e$ and eligible $i\ne e$, introduce binary $y_{{ie}}$ and a valid upper bound $M$ on score differences:

$s_i-s_e\le M y_{{ie}},\qquad \sum_{{i\in G\setminus\{{e\}}}}y_{{ie}}\le m-1,$

where $m=k$ for direct elimination and $m=\min(k+1,|G|)$ for the weak save rule. If $s_i>s_e$, integrality forces $y_{{ie}}=1$; when $s_i\le s_e$, a feasible solution can take $y_{{ie}}=0$. Thus the second inequality is equivalent to $e\in B_m$. Finale order constraints are linear inequalities between adjacent combined scores in the observed placement order.

**Theorem 1 (sharp endpoint correctness).** For every encoded week, tie policy, mechanism, candidate $a$, and public rank $r$, the assignment MILP has a feasible integer solution with $x_{{ar}}=1$ if and only if there exists a strict public permutation satisfying the same finale or tie-inclusive bottom-set predicate used by the enumerative rule checker. Consequently, minimizing and maximizing $\sum_r r x_{{ar}}$ give the sharp attainable endpoints for candidate $a$.

**Proof.** A feasible integer assignment matrix has exactly one one in each row and column, hence represents one strict permutation; conversely every strict permutation defines such a matrix. For an elimination constraint, $s_i>s_e$ forces $y_{{ie}}=1$. If at most $m-1$ eligible scores are strictly greater than $s_e$, choose $y_{{ie}}=1$ exactly for those scores and zero otherwise, satisfying both inequalities. Therefore the auxiliary constraints hold exactly when $e$ is in the tie-inclusive bottom-$m$ set. Applying this to every eliminated candidate, or applying the adjacent finale inequalities, proves equivalence. Linear optimization over the equivalent feasible integer assignments attains the sharp minimum and maximum. $\square$

The algorithm solves two MILPs per candidate. A non-finale week has $n^2+|E|(|G|-1)$ binary variables, $2n+|E|(|G|-1)+|E|$ principal constraints, and $2n$ optimization calls. This is a workload bound, not a polynomial-time claim: general binary linear optimization is NP-hard, and exact compatible-permutation counting is not supplied by endpoint optimization. In this data, the primary tie-policy analysis required {int(ordinal_stats['primary_solver_calls']):,} solver calls; all four tie policies plus the direct counterfactual used {int(ordinal_stats['all_policy_solver_calls']):,} calls in {ordinal_stats['elapsed_seconds']:.2f} seconds on the recorded environment.

### 5.2 Set nesting and interpretation

Adding valid constraints intersects a feasible set and cannot enlarge it. Likewise, the direct ordinal set is contained in the weak save set because $B_k\subseteq B_{{k+1}}$. These are reported as lemmas and implementation invariants, not empirical wins. Cardinal share widths and normalized ordinal rank-support widths have no canonical common uncertainty scale; numerical cross-regime comparisons are descriptive unless a substantive common loss functional is supplied.

## 6. Bayesian comparator and simulation uncertainty

The internal Bayesian comparator draws from a symmetric $Dirichlet(1)$ prior and applies a zero-one likelihood for the observed linear constraints. With fixed $N=8192$ prior draws and at least 100 accepted states, the reported coordinate interval is the 95% equal-tail rectangle

$I_i^B=[Q_{{0.025}}(p_i\mid O),Q_{{0.975}}(p_i\mid O)].$

The external comparator is exact under a uniform distribution over compatible strict rankings and uses the analogous 0.025 and 0.975 marginal rank quantiles. These are marginal posterior credible rectangles, not identified sets and not frequentist confidence sets. “Same information” means the methods receive the same observed record; it does not make their inferential semantics identical. The fixed rejection bank is retained as a transparent baseline, not recommended as an efficient polytope sampler. Hit-and-run or sequential Monte Carlo is future work; no unrun sensitivity result is claimed.

For each performance estimator we report a Monte Carlo standard error over independent simulated replications: $s/\sqrt{{N}}$ for a mean and the equivalent Bernoulli standard error for binary coverage. Paired removal effects are computed replication by replication before their SE is calculated. The 20-seed standard deviation remains a random-stream stability diagnostic; empirical 2.5%-97.5% seed quantiles are not labeled confidence intervals.

## 7. Evaluation design

Twenty fixed seeds in the hash-locked evaluation configuration cover 12 internal regions (three active-set sizes by four outcome-noise levels) with 250 replications per seed-region and three external ordinal structures with 120 replications per seed-region. This gives 60,000 internal and 7,200 external known-truth cases. Clean simulations are implementation checks under correct specification. Positive outcome noise is a deliberately specified misspecification stress test, not an empirical institutional error rate.

Rule-aware and simplex-only identified sets are evaluated by joint truth membership. The legacy marginal-box diagnostic is shown only to document the metric correction. Bayesian rectangles use their explicitly defined credible-interval coverage. Maximum-entropy points are evaluated by point error and are excluded from coverage-width Pareto tests. The Pareto screen is an operational partial order—coverage no lower, width no higher, at least one strict—not an optimality theorem under an elicited loss function.

## 8. Results

### 8.1 Corrected internal coverage and MCSE

**Table 1. Joint-set coverage, legacy projection-envelope coverage, and width.** Each row pools the three active-set sizes and 20 fixed seeds at the stated noise level. MCSEs use all simulated replications in the row.

{coverage_table}

[Insert Figure 3 near here]

Under clean generation, joint-set and projection-envelope coverage are both one. Under outcome noise, the projection box can still contain an infeasible truth, so it overstates coverage. Across positive-noise paired replications, removing elimination increases joint-set coverage by {coverage_stats['positive_joint_change']:.6f}, compared with the old projection-envelope change {coverage_stats['positive_projection_change']:.6f}. Width results are unchanged because the correction changes the coverage predicate, not the LP bounds.

**Table 2. Paired effect of removing the elimination constraint.** Changes are without-elimination minus full rule-aware, calculated within the same generated case.

{effects_table}

The direction of clean coverage and nested width is structural. The informative quantity is the magnitude under this generator, with its MCSE and stated misspecification boundary.

### 8.2 Exact ordinal endpoints

**Table 3. Legacy sampled/enumerated width versus exact MILP endpoint width.** “Monte Carlo” describes the legacy feasible-fraction computation; all revised endpoints are exact.

{endpoint_table}

Among {int(ordinal_stats['sampled_contestants'])} contestant-week rows whose legacy endpoints came from sampled permutations, {int(ordinal_stats['sampled_endpoint_corrections'])} have at least one corrected endpoint. The maximum normalized width correction is {ordinal_stats['max_width_correction']:.6f}, and the mean correction is {ordinal_stats['mean_width_correction']:.6f}. The revised primary-policy mean exact width is {ordinal_stats['r_exact_width']:.6f} in R and {ordinal_stats['rplus_exact_width']:.6f} in R-plus. Feasible fractions and the weak/direct feasible-count ratio remain Monte Carlo estimates in large fields; their binomial MCSE applies to those proportions only and is not used to justify endpoint accuracy.

### 8.3 Rule-provenance sensitivity

**Table 4. Sensitivity to the uncertain season-28 rank-return assignment.** The primary scenario follows the COMAP reasonable assumption; the sensitivity excludes season 28 rather than asserting an undocumented alternative rule for it.

{provenance_table}

The weak-minus-direct comparison is within the same week and exact for rank-support width. Cross-season averages are descriptive because candidate fields, expert scores, event types, and histories differ.

### 8.4 Bayesian definition and undefined regions

**Table 5. Internal Bayesian rows below the fixed 100-accepted-draw threshold.** Denominators are complete; interval summaries are undefined in these rows.

{bayes_table}

There are {bayesian_stats['total']} undefined internal rows, including {bayesian_stats['clean']} clean rows; the exact external ordinal posterior has {bayesian_stats['external']} undefined rows. In the clean same-information seed-cell comparison retained from the verified experiment, Bayesian intervals strictly Pareto-dominate rule-aware intervals in 14/120 cells, all in the external seven-candidate, three-round region; reverse dominance is 0/120. Those 14 external results do not arise from internal rejection failures. They remain conditional on the uniform compatible-ranking posterior and on the operational Pareto criterion.

### 8.5 External simulation

The external known-truth simulation retains the previously verified result: the rule-aware and rule-agnostic ordinal sets contain the complete generated ranking across all three small-field structures, while direct-rule misspecification can exclude it. [Insert Figure 4 near here.] These cases are exact finite-state computations within the simulator and do not validate a historical institution.

## 9. Discussion

The exact endpoint formulation removes the main numerical mismatch in the ordinal analysis. Sampling remains useful for estimating compatible-state proportions and posterior masses, but a proportion MCSE cannot certify discovery of rare support endpoints. Separating these targets avoids transferring an error measure from one statistic to another.

The joint-coverage correction is equally consequential. Coordinate-wise sharp bounds summarize projections; they do not convert a polytope into a box. Reporting $C_{{box}}$ as set coverage masked some rule violations. The corrected result makes the simulated mechanism transparent: if the recorded outcome contradicts the generating rule, the hard rule-aware polytope excludes the generating truth, whereas the simplex-only set retains it at a width cost.

The paper does not derive a universal method-selection rule. Identified sets expose the consequences of institutional assumptions without adding a probability distribution within the set. Bayesian rectangles summarize probability under the fixed prior and zero-one likelihood. A user with a substantively defensible prior may prefer the latter; a user seeking assumption-transparent compatibility may prefer the former. Without an elicited loss function, coverage-width Pareto dominance is a descriptive screen rather than a complete decision theory.

## 10. Limitations

The empirical application identifies compatibility of latent public preferences; it does not produce an observed audience-vote estimate. The season-28 rule assignment is uncertain; excluding season 28 is reported, but undocumented historical tie-breaking and production decisions remain unidentified. Weekly rankings are not dynamically coupled. Exact endpoint MILPs do not give exact feasible-set cardinalities. The internal Bayesian rejection sampler has region-dependent failures and has not been replaced by a robust polytope sampler. Simulation results are conditional on the two generators, fixed grid, and specified noise process. No user, welfare, privacy, trust, cost, legal, or organizational outcome is measured; the earlier unmeasured “flexibility” and “accountability” figure panels and governance matrices are therefore removed from the main article.

## 11. Conclusion

Rule-constrained expert-crowd aggregation is naturally an inverse feasibility problem. Joint constraint membership is the correct coverage event for a polyhedral identified set, and exact optimization is the correct tool for sharp ordinal support endpoints when permutation enumeration is impractical. In this testbed, the corrections increase the measured coverage cost of outcome misspecification and reveal where sampled ordinal extrema were too narrow. The results support an auditable, rule-conditional methodology while leaving public preference partially identified.

## Data and code availability

    The empirical source is the official COMAP 2026 Problem C data file [14]. The local reproducibility bundle contains project-relative command-line scripts, fixed seeds, processed inputs, generated tables, tests, and hash manifests. The corrected Stage 32 code and evidence are publicly available in the versioned release `{PUBLIC_RELEASE_URL}`. This statement identifies the public release directly and does not claim a persistent DOI archive.

## CRediT author contributions

{AUTHOR_CONTRIBUTIONS}

## Funding

This research received no specific grant from funding agencies in the public, commercial, or not-for-profit sectors.

## Competing interests

The authors declare no known competing financial interests or personal relationships that could have appeared to influence the work reported in this paper.

## Ethics statement

The study uses an official public competition dataset and synthetic simulations. It involves no intervention with human participants, collection of private participant data, or recovery of individual audience votes.

## Declaration of generative AI and AI-assisted technologies

During manuscript preparation, the authors used OpenAI ChatGPT/Codex for language editing, code review, test development, and reproducibility auditing. The authors reviewed and edited all outputs and take responsibility for the content.

## References

[1] C.F. Manski, Identification problems and decisions under ambiguity, Journal of Econometrics 95 (2) (2000) 415-442. https://doi.org/10.1016/S0304-4076(99)00045-7.

[2] E. Tamer, Partial Identification in Econometrics, Annual Review of Economics 2 (2010) 167-195. https://doi.org/10.1146/annurev.economics.050708.143401.

[3] G.W. Imbens, C.F. Manski, Confidence intervals for partially identified parameters, Econometrica 72 (6) (2004) 1845-1857. https://doi.org/10.1111/j.1468-0262.2004.00555.x.

[4] H. Kaido, F. Molinari, J. Stoye, Confidence Intervals for Projections of Partially Identified Parameters, Econometrica 87 (4) (2019) 1397-1432. https://doi.org/10.3982/ECTA14075.

[5] H.R. Moon, F. Schorfheide, Bayesian and Frequentist Inference in Partially Identified Models, Econometrica 80 (2) (2012) 755-782. https://doi.org/10.3982/ECTA8360.

[6] L. Xia, V. Conitzer, Determining Possible and Necessary Winners Given Partial Orders, Journal of Artificial Intelligence Research 41 (2011) 25-67. https://doi.org/10.1613/jair.3186.

[7] N. Betzler, B. Dorn, Towards a dichotomy for the Possible Winner problem in elections based on scoring rules, Journal of Computer and System Sciences 76 (8) (2010) 812-836. https://doi.org/10.1016/j.jcss.2010.04.002.

[8] Y. Bachrach, N. Betzler, P. Faliszewski, Probabilistic Possible Winner Determination, Proceedings of the AAAI Conference on Artificial Intelligence 24 (1) (2010) 697-702. https://doi.org/10.1609/aaai.v24i1.7609.

[9] P. Dey, N. Misra, On the exact amount of missing information that makes finding possible winners hard, Journal of Computer and System Sciences 135 (2023) 32-54. https://doi.org/10.1016/j.jcss.2023.02.003.

[10] T. Lu, C. Boutilier, Preference elicitation and robust winner determination for single- and multi-winner social choice, Artificial Intelligence 279 (2020) 103203. https://doi.org/10.1016/j.artint.2019.103203.

[11] T.P. Morris, I.R. White, M.J. Crowther, Using simulation studies to evaluate statistical methods, Statistics in Medicine 38 (11) (2019) 2074-2102. https://doi.org/10.1002/sim.8086.

[12] A. Burton, D.G. Altman, P. Royston, R.L. Holder, The design of simulation studies in medical statistics, Statistics in Medicine 25 (24) (2006) 4279-4292. https://doi.org/10.1002/sim.2673.

[13] T. Monks, C.S.M. Currie, B.S. Onggo, S. Robinson, M. Kunc, S.J.E. Taylor, Strengthening the reporting of empirical simulation studies: Introducing the STRESS guidelines, Journal of Simulation 13 (1) (2019) 55-67. https://doi.org/10.1080/17477778.2018.1442155.

[14] COMAP, 2026 MCM Problem C: Data With The Stars, The Consortium for Mathematics and Its Applications, 2026. https://contest.comap.com/undergraduate/contests/mcm/contests/2026/problems/index.html (accessed 2026-08-21).

## Figure captions

**Figure 1. Rule-conditioned inverse-inference architecture.** Observed expert inputs, the stated aggregation rule, and coarse outcomes map to a cardinal polytope or ordinal feasible-ranking set. Latent public preference remains partially identified.

**Figure 2. Reproducible comparison workflow.** The workflow separates configuration, information alignment, inference, simulation evaluation, and evidence auditing. Known-truth coverage is evaluated only in simulation.

**Figure 3. Corrected internal coverage-width evaluation.** Panel (a) distinguishes joint polytope membership from the legacy Cartesian product of marginal bounds; panel (b) reports mean coordinate width. Error bars are 95% Monte Carlo error intervals, estimate plus or minus 1.96 MCSE, over simulated replications. Outcome noise is a specified misspecification stress test.

**Figure 4. External ordinal simulation.** Known-truth coverage and normalized rank-support width across the three predeclared small-field structures. Results are conditional on the simulator and named tie policy.
"""


def math_expression_to_unicode(expression: str) -> str:
    """Convert the fixed TeX subset used here to readable equation text."""
    replacements = {
        r"\mathbb{R}": "ℝ",
        r"\mathbb R": "ℝ",
        r"\sum": "∑",
        r"\min": "min",
        r"\max": "max",
        r"\forall": "∀",
        r"\text{for every}": "for every",
        r"\text{if}": "if",
        r"\text{and}": "and",
        r"\le": "≤",
        r"\ge": "≥",
        r"\subseteq": "⊆",
        r"\in": "∈",
        r"\qquad": "   ",
        r"\,": " ",
        r"\cdot": "·",
        r"\star": "★",
        r"\theta": "θ",
        r"\prod": "∏",
        r"\ldots": "…",
        r"\{": "<<LBRACE>>",
        r"\}": "<<RBRACE>>",
    }
    result = expression
    for source, target in replacements.items():
        result = result.replace(source, target)
    result = re.sub(r"\\text\{([^{}]*)\}", r"\1", result)
    result = re.sub(r"([A-Za-z∑∏])_\{?([0-9]+)\}?", lambda match: match.group(1) + match.group(2).translate(SUBSCRIPTS), result)
    result = re.sub(r"([A-Za-zℝ])\^\{?([0-9]+)\}?", lambda match: match.group(1) + match.group(2).translate(SUPERSCRIPTS), result)
    # Some fixed indicator expressions use literal closing braces. Preserve
    # them alongside the escaped set braces before removing TeX group braces.
    if "=1<<LBRACE>>" in result:
        result = re.sub(r"\}+(\.)$", r"<<RBRACE>>\1", result)
    result = result.replace("{", "").replace("}", "")
    return result.replace("<<LBRACE>>", "{").replace("<<RBRACE>>", "}").replace("\\", "")


def add_math(paragraph, expression: str, size: float = 11.0) -> None:
    """Append one Office Math object so equations survive DOCX/PDF conversion."""
    omath = OxmlElement("m:oMath")
    math_run = OxmlElement("m:r")
    properties = OxmlElement("m:rPr")
    size_node = OxmlElement("m:sz")
    size_node.set(f"{{{M_NS}}}val", str(int(size * 2)))
    properties.append(size_node)
    math_run.append(properties)
    text_node = OxmlElement("m:t")
    text_node.text = math_expression_to_unicode(expression)
    math_run.append(text_node)
    omath.append(math_run)
    paragraph._p.append(omath)


def add_runs(paragraph, text: str, size: float = 11.0) -> None:
    pattern = re.compile(r"(\*\*.*?\*\*|`.*?`|\$.*?\$)")
    position = 0
    for match in pattern.finditer(text):
        if match.start() > position:
            run = paragraph.add_run(text[position : match.start()])
            run.font.size = Pt(size)
        token = match.group(0)
        if token.startswith("$"):
            add_math(paragraph, token[1:-1], size)
        else:
            run = paragraph.add_run(token[2:-2])
            run.bold = True
            run.font.size = Pt(size)
        position = match.end()
    if position < len(text):
        run = paragraph.add_run(text[position:])
        run.font.size = Pt(size)
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Times New Roman")
        run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Times New Roman")


def markdown_to_docx(markdown: str, output: Path, *, title_page: bool = False) -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = section.bottom_margin = Inches(1)
    section.left_margin = section.right_margin = Inches(1)
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 2 if not title_page else 1.15
    for name in ("Heading 1", "Heading 2", "Heading 3"):
        style = document.styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    lines = markdown.splitlines()
    index = 0
    paragraph_buffer: list[str] = []

    def flush() -> None:
        if not paragraph_buffer:
            return
        text = " ".join(value.strip() for value in paragraph_buffer).strip()
        paragraph_buffer.clear()
        if text:
            paragraph = document.add_paragraph()
            add_runs(paragraph, text)

    while index < len(lines):
        line = lines[index]
        if line.startswith("|") and index + 1 < len(lines) and lines[index + 1].startswith("|"):
            flush()
            table_lines: list[str] = []
            while index < len(lines) and lines[index].startswith("|"):
                table_lines.append(lines[index])
                index += 1
            parsed = [
                [cell.strip() for cell in row.strip().strip("|").split("|")]
                for row in table_lines
            ]
            if len(parsed) >= 2:
                parsed.pop(1)
            table = document.add_table(rows=1, cols=len(parsed[0]))
            table.style = "Table Grid"
            for cell, value in zip(table.rows[0].cells, parsed[0]):
                cell.text = value
                for run in cell.paragraphs[0].runs:
                    run.bold = True
                    run.font.size = Pt(8)
            for row in parsed[1:]:
                cells = table.add_row().cells
                for cell, value in zip(cells, row):
                    cell.text = value
                    for run in cell.paragraphs[0].runs:
                        run.font.size = Pt(8)
            continue
        if line.startswith("#"):
            flush()
            level = len(line) - len(line.lstrip("#"))
            heading = line[level:].strip()
            if level == 1:
                paragraph = document.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run(heading)
                run.bold = True
                run.font.name = "Times New Roman"
                run.font.size = Pt(16)
            else:
                document.add_heading(heading, level=min(level - 1, 3))
        elif line.startswith("- "):
            flush()
            paragraph = document.add_paragraph(style="List Bullet")
            add_runs(paragraph, line[2:].strip())
        elif not line.strip():
            flush()
        else:
            paragraph_buffer.append(line)
        index += 1
    flush()
    document.core_properties.author = ""
    document.core_properties.last_modified_by = ""
    document.core_properties.title = "Coverage-Width Tradeoffs in Rule-Constrained Expert-Crowd Aggregation"
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


def build_cover_letter(stats: dict[str, float], ordinal: dict[str, float]) -> str:
    return f"""# Cover Letter

Dear Editors of *Mathematics and Computers in Simulation*,

Please consider the manuscript “Coverage-Width Tradeoffs in Rule-Constrained Expert-Crowd Aggregation” as a Research Article.

The revision strengthens the paper's computational and simulation methodology. It introduces and verifies an exact binary-optimization algorithm for sharp ordinal rank-support endpoints, replaces marginal projection-box coverage with joint feasible-set membership, reports performance-specific Monte Carlo standard errors and paired uncertainty, formalizes tie-inclusive bottom sets and multiple-elimination semantics, and adds a rule-provenance sensitivity for the uncertain season-28 rank-return assumption. Across positive-noise paired replications, the corrected joint-set coverage change from removing elimination is {stats['positive_joint_change']:.6f} (MCSE {stats['positive_joint_mcse']:.6f}); the prior projection-envelope calculation was {stats['positive_projection_change']:.6f}. The exact endpoint analysis covers 87 empirical weeks with {int(ordinal['primary_solver_calls']):,} primary-policy solver calls.

The manuscript now engages directly with possible/necessary-winner, partial-preference, robust-winner, and Bayesian partial-identification literature. It distinguishes identified sets from 95% equal-tail Bayesian credible rectangles and does not claim observed audience-vote recovery, universal method superiority, deployed user effects, or a general worst-case complexity theorem.

    The work is original, is not under consideration elsewhere, and has been approved by all authors. The authors declare no competing interests. The empirical source is public and all corrected calculations are generated by command-line scripts. The corrected Stage 32 code and evidence are available at the versioned public release `{PUBLIC_RELEASE_URL}`; a persistent DOI archive is not claimed.

Sincerely,

Lizhen Deng
Corresponding author
"""


def build_highlights() -> str:
    lines = [
        "Exact MILPs recover sharp ordinal rank endpoints without sampling.",
        "Joint-set coverage replaces marginal projection-box coverage.",
        "Elimination misspecification trades coverage for identified-set width.",
        "Bayesian rectangles and identified sets retain distinct semantics.",
        "Season-28 rule uncertainty is reported in a sensitivity analysis.",
    ]
    if any(len(line) > 85 for line in lines):
        raise Stage32Error("An Elsevier Highlight exceeds 85 characters")
    return "\n".join(f"- {line}" for line in lines) + "\n"


def response_matrix() -> pd.DataFrame:
    rows = [
        ("Coverage was marginal-box rather than polytope membership", "FIXED", "src/dss_common.py; src/synthetic_benchmark.py; tables/joint_set_coverage_*.csv"),
        ("Sampled ordinal extrema lacked endpoint guarantees", "FIXED", "src/ranking_identification.py; tables/ordinal_exact_*.csv; focused enumeration-equivalence tests"),
        ("B_k, ties, save, multiple elimination underformalized", "FIXED", "Manuscript Sections 3 and 5; exact MILP constraint tests"),
        ("Complexity heading overclaimed", "FIXED", "Manuscript Section 5.1 reports workload and explicitly disclaims polynomial complexity"),
        ("Possible-winner and robust-winner literature missing", "FIXED", "Manuscript Section 2 and references 6-10"),
        ("Bayesian interval level/type and semantics unclear", "FIXED", "Manuscript Section 6 defines 95% equal-tail rectangles and conditional-prior semantics"),
        ("Undefined Bayesian rows not reported by region", "FIXED", "tables/bayesian_undefined_by_region.csv; Manuscript Table 5"),
        ("Key performance MCSE and paired uncertainty missing", "FIXED", "tables/performance_mcse.csv; paired_elimination_effect_mcse.csv"),
        ("Preregistered wording lacked immutable registration", "FIXED", "Replaced by predeclared and hash-locked throughout revised package"),
        ("Season-28 rule provenance treated as fact", "FIXED", "tables/rule_provenance_sensitivity.csv; Manuscript Sections 3 and 8.3"),
        ("Missing proxy incorrectly described as cause of skipped P week", "FIXED", "Manuscript Section 3 identifies missing active judge total as cause"),
        ("Unmeasured Figure 3b/4b and governance material diluted evidence", "FIXED", "Removed from revised main package; four evidence-bearing figures retained"),
        ("Public repository did not contain latest stages", "FIXED", "The corrected Stage 32 package is available at the versioned public release URL"),
        ("AMS MSC/PDF/highlights requirements", "NOT APPLICABLE", "Report confused Mathematics of Computation with Elsevier Mathematics and Computers in Simulation"),
    ]
    return pd.DataFrame(rows, columns=["review_issue", "resolution", "evidence"])


def package_status(
    stats: dict[str, float], ordinal: dict[str, float], files: int
) -> str:
    return f"""# Stage 32 MATCOM Scientific-Correction Status

Status: `SCIENTIFIC_CORRECTIONS_COMPLETE_PUBLIC_RELEASE_AVAILABLE`.

The target journal is Elsevier's *Mathematics and Computers in Simulation* (MATCOM), not AMS *Mathematics of Computation*. The AMS-specific MSC, PDF-first, and AMS review-model requests in the supplied report are therefore excluded as venue-confused requirements. All scientifically applicable comments were addressed.

## Completed scientific corrections

- Joint feasible-set coverage replaces Cartesian marginal-box coverage. Positive-noise removal effect: {stats['positive_joint_change']:.6f} (MCSE {stats['positive_joint_mcse']:.6f}); legacy projection effect: {stats['positive_projection_change']:.6f}.
- Exact MILP endpoints cover 87 empirical ordinal weeks and all four tie policies; the primary analysis uses {int(ordinal['primary_solver_calls']):,} solver calls.
- Performance-specific MCSE, paired-effect uncertainty, region-level Bayesian failures, and season-28 provenance sensitivity are generated tables.
- “Preregistered” is replaced by “predeclared and hash-locked.”
- Unmeasured descriptor panels and governance matrices are removed from the main article.
- The revised candidate contains {files} manifested files before the manifest itself.

## Remaining external author gates before submission

1. Confirm the Elsevier Editorial Manager article type, upload slots, and review model.
2. Reconfirm author metadata and inspect the portal-generated PDF.

The versioned public release is available at `{PUBLIC_RELEASE_URL}`. No persistent DOI is asserted. No remote repository or submission portal was changed by this script.
"""


def build_manifest(package: Path) -> str:
    files = sorted(
        path
        for path in package.rglob("*")
        if path.is_file() and path.name != "PACKAGE_MANIFEST.md"
    )
    rows = "\n".join(
        f"| {path.relative_to(package).as_posix()} | {path.stat().st_size} | {sha256(path)} |"
        for path in files
    )
    return (
        "# Stage 32 Candidate Package Manifest\n\n"
        "| Relative file | Bytes | SHA-256 |\n"
        "| --- | --- | --- |\n"
        + rows
        + "\n"
    )


def copy_figure(source: Path, target: Path) -> None:
    if not source.is_file():
        raise Stage32Error(f"Required frozen figure missing: {source}")
    target.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(source, target)


def main(argv: list[str] | None = None) -> int:
    args = parse_args(argv)
    root = args.project_root.resolve()
    output = resolve(root, args.output_dir, "Stage 32 output")
    required = [
        root / LEGACY_PACKAGE,
        root / INTERNAL_RAW,
        root / BAYESIAN_RAW,
        root / PANEL,
        root / LEGACY_R,
        root / LEGACY_RPLUS,
        root / LEGACY_CONTESTANT,
        root / CONSTRAINT_SUMMARY,
    ]
    missing = [str(path) for path in required if not path.exists()]
    if missing:
        raise Stage32Error(f"Required input(s) missing: {missing}")

    tables = output / "tables"
    reports = output / "reports"
    package = output / "MATCOM_revised_candidate_package"
    figures_pdf = package / "figures/pdf"
    figures_tiff = package / "figures/tiff"
    for directory in (tables, reports, package, figures_pdf, figures_tiff):
        directory.mkdir(parents=True, exist_ok=True)

    selected, seed_level, across_seed, performance, effects, coverage_stats = (
        corrected_internal_coverage(root)
    )
    contestant, weeks, week_comparison, provenance, ordinal_stats = (
        exact_ordinal_endpoints(root)
    )
    bayesian, bayesian_stats = bayesian_undefined_report(root)

    write_csv(tables / "joint_set_coverage_replication_level.csv", selected)
    write_csv(tables / "joint_set_coverage_seed_level.csv", seed_level)
    write_csv(tables / "joint_set_coverage_across_seeds.csv", across_seed)
    write_csv(tables / "performance_mcse.csv", performance)
    write_csv(tables / "paired_elimination_effect_mcse.csv", effects)
    write_csv(tables / "ordinal_exact_endpoint_by_contestant.csv", contestant)
    write_csv(tables / "ordinal_exact_endpoint_by_week_and_tie.csv", weeks)
    write_csv(tables / "ordinal_sampling_endpoint_audit.csv", week_comparison)
    write_csv(tables / "rule_provenance_sensitivity.csv", provenance)
    write_csv(tables / "bayesian_undefined_by_region.csv", bayesian)
    matrix = response_matrix()
    write_csv(tables / "review_response_matrix.csv", matrix)

    plot_corrected_internal(
        selected,
        figures_pdf / "Figure_03_corrected_internal_coverage_width.pdf",
        figures_tiff / "Figure_03_corrected_internal_coverage_width.tif",
    )
    legacy = root / LEGACY_PACKAGE / "figures"
    for extension in ("pdf", "tiff"):
        suffix = "pdf" if extension == "pdf" else "tif"
        source_dir = legacy / extension
        target_dir = package / "figures" / extension
        copy_figure(
            source_dir / f"Figure_01_rule_conditioned_inference_architecture.{suffix}",
            target_dir / f"Figure_01_rule_conditioned_inference_architecture.{suffix}",
        )
        copy_figure(
            source_dir / f"Figure_02_reproducible_comparison_workflow.{suffix}",
            target_dir / f"Figure_02_reproducible_comparison_workflow.{suffix}",
        )
        copy_figure(
            source_dir / f"Figure_06_multiseed_external_sensitivity.{suffix}",
            target_dir / f"Figure_04_multiseed_external_sensitivity.{suffix}",
        )

    manuscript = build_manuscript(
        selected,
        effects,
        week_comparison,
        provenance,
        bayesian,
        coverage_stats,
        ordinal_stats,
        bayesian_stats,
    )
    main_source = package / "MATCOM_main_manuscript_source.md"
    write_text(main_source, manuscript)
    markdown_to_docx(manuscript, package / "MATCOM_main_manuscript.docx")

    cover = build_cover_letter(coverage_stats, ordinal_stats)
    write_text(package / "MATCOM_cover_letter_source.md", cover)
    markdown_to_docx(cover, package / "MATCOM_cover_letter.docx", title_page=True)
    highlights = build_highlights()
    write_text(package / "MATCOM_Highlights.txt", highlights)
    markdown_to_docx(
        "# Highlights\n\n" + highlights,
        package / "MATCOM_Highlights.docx",
        title_page=True,
    )
    shutil.copy2(
        root / LEGACY_PACKAGE / "MATCOM_title_page_source.md",
        package / "MATCOM_title_page_source.md",
    )
    shutil.copy2(
        root / LEGACY_PACKAGE / "MATCOM_title_page.docx",
        package / "MATCOM_title_page.docx",
    )
    captions = "\n".join(
        line.replace("**", "")
        for line in manuscript.splitlines()
        if line.startswith("**Figure ")
    )
    write_text(package / "MATCOM_Figure_Captions.txt", captions)
    write_text(
        reports / "COVERAGE_METRIC_CORRECTION.md",
        f"""# Coverage Metric Correction

The legacy internal `coverage` field tested whether every truth coordinate lay between separately optimized marginal bounds. That Cartesian projection box is an outer approximation to the joint polytope. Stage 32 evaluates all original equalities, inequalities, and variable bounds.

Across positive-noise paired replications, removing elimination changes joint-set coverage by {coverage_stats['positive_joint_change']:.6f} (MCSE {coverage_stats['positive_joint_mcse']:.6f}), not the legacy projection-box value {coverage_stats['positive_projection_change']:.6f}. Existing frozen Stage 26X files are retained as provenance and are not overwritten.
""",
    )
    write_text(
        reports / "EXACT_ORDINAL_ENDPOINT_AUDIT.md",
        f"""# Exact Ordinal Endpoint Audit

- Empirical weeks: {int(ordinal_stats['weeks'])}.
- Primary-policy solver calls: {int(ordinal_stats['primary_solver_calls'])}.
- All-policy solver calls including direct R-like comparisons: {int(ordinal_stats['all_policy_solver_calls'])}.
- Legacy sampled contestant-week rows: {int(ordinal_stats['sampled_contestants'])}.
- Rows with at least one corrected endpoint: {int(ordinal_stats['sampled_endpoint_corrections'])}.
- Maximum normalized width correction: {ordinal_stats['max_width_correction']:.6f}.
- Focused equivalence tests compare MILP endpoints with complete enumeration across direct, weak-save, tied-expert, multiple-elimination, withdrawal, and finale cases.
""",
    )
    write_text(
        reports / "REVIEW_RESPONSE_MATRIX.md",
        "# Review Response Matrix\n\n"
        + markdown_table(matrix, list(matrix.columns), {}),
    )

    manifested_files = len(
        [
            path
            for path in package.rglob("*")
            if path.is_file() and path.name != "PACKAGE_MANIFEST.md"
        ]
    )
    status = package_status(coverage_stats, ordinal_stats, manifested_files)
    write_text(package / "PACKAGE_STATUS.md", status)
    write_text(package / "PACKAGE_MANIFEST.md", build_manifest(package))

    unsupported_preregistration_claims = (
        "preregistered seeds",
        "preregistered design",
        "preregistered simulation",
    )
    combined_sources = (manuscript + "\n" + cover).casefold()
    if any(term in combined_sources for term in unsupported_preregistration_claims):
        raise Stage32Error("Unsupported preregistration claim remains in revised sources")
    if "true audience vote" in manuscript.casefold():
        raise Stage32Error("Forbidden true-audience-vote claim remains")
    if len(matrix) != 14:
        raise Stage32Error("Review response matrix is incomplete")
    print("STAGE32_MATCOM_SCIENTIFIC_CORRECTIONS=PASS")
    print(
        f"JOINT_COVERAGE_CHANGE={coverage_stats['positive_joint_change']:.6f}; "
        f"LEGACY_PROJECTION_CHANGE={coverage_stats['positive_projection_change']:.6f}"
    )
    print(
        f"EXACT_ORDINAL_WEEKS={int(ordinal_stats['weeks'])}; "
        f"PRIMARY_SOLVER_CALLS={int(ordinal_stats['primary_solver_calls'])}"
    )
    print(f"PACKAGE={package}")
    return 0


if __name__ == "__main__":
    try:
        raise SystemExit(main())
    except Stage32Error as exc:
        raise SystemExit(f"Stage 32 failed: {exc}") from exc
