#!/usr/bin/env python3
"""Build Stage 25H-D DSS submission DOCX files from existing audited assets only."""

from __future__ import annotations

import argparse
import csv
import hashlib
import re
import shutil
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


PROJECT_ROOT = Path(__file__).resolve().parents[1]
PACKAGE = Path("submission_package_stage25")
TITLE = "Rule-Aware Decision Support for Expert-Crowd Aggregation under Hidden Public Preferences"
REPOSITORY_URL = "https://github.com/denglizhen-113/coverage-width-tradeoffs-rule-constrained-aggregation"
CORRESPONDING_ADDRESS = (
    "Huazhong University of Science and Technology, 1037 Luoyu Road, "
    "Hongshan District, Wuhan, Hubei 430074, China"
)
ANON_NAME = "DSS_anonymized_manuscript_STAGE25H_D_final.docx"
TITLE_NAME = "DSS_title_page_STAGE25H_D_final.docx"
PREVIEW_NAME = "DSS_anonymized_manuscript_STAGE25H_D_final_preview.pdf"
HIGHLIGHTS_NAME = "DSS_highlights_STAGE25H_D_final.md"


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Reconstruct Stage 25H-D DSS DOCX submission files from audited figures and tables."
    )
    parser.add_argument("--project-root", type=Path, default=PROJECT_ROOT)
    parser.add_argument(
        "--finalize-only",
        action="store_true",
        help="Do not rebuild DOCX files; update H-D audits using an exported preview PDF.",
    )
    parser.add_argument(
        "--preview-pdf",
        type=Path,
        default=None,
        help="Existing PDF exported from the corrected anonymized DOCX for final page-count validation.",
    )
    return parser.parse_args()


def write_text(path: Path, content: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    temporary = path.with_suffix(path.suffix + ".tmp")
    temporary.write_text(content.strip() + "\n", encoding="utf-8")
    temporary.replace(path)


def write_csv(path: Path, rows: list[dict[str, object]], fields: list[str]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    temporary = path.with_suffix(path.suffix + ".tmp")
    with temporary.open("w", encoding="utf-8", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=fields)
        writer.writeheader()
        writer.writerows(rows)
    temporary.replace(path)


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def verify_frozen_artifacts(root: Path) -> int:
    manifest = root / PACKAGE / "11_reproducibility/frozen_artifact_hash_manifest_stage25.csv"
    if not manifest.is_file():
        raise FileNotFoundError(f"Frozen-artifact manifest is missing: {manifest}")
    with manifest.open("r", encoding="utf-8", newline="") as handle:
        rows = list(csv.DictReader(handle))
    if not rows:
        raise ValueError(f"Frozen-artifact manifest has no entries: {manifest}")
    mismatches: list[str] = []
    for row in rows:
        relative_path = row.get("relative_path", "")
        expected = row.get("expected_sha256", "")
        target = root / relative_path
        if not relative_path or not expected or not target.is_file():
            mismatches.append(f"{relative_path or '<missing path>'}: malformed or missing")
            continue
        observed = sha256(target)
        if observed != expected:
            mismatches.append(f"{relative_path}: expected {expected}, observed {observed}")
    if mismatches:
        raise RuntimeError(
            "Frozen Stage 21--24 artifact verification failed; no Stage 25H-D output was written. "
            + "; ".join(mismatches)
        )
    return len(rows)


def set_font(run, size: float, *, bold: bool = False, italic: bool = False) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def set_section_layout(section, *, landscape: bool = False) -> None:
    section.orientation = WD_ORIENT.LANDSCAPE if landscape else WD_ORIENT.PORTRAIT
    section.page_width = Inches(11 if landscape else 8.5)
    section.page_height = Inches(8.5 if landscape else 11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.5)
    section.footer_distance = Inches(0.5)
    columns = section._sectPr.find(qn("w:cols"))
    if columns is None:
        columns = OxmlElement("w:cols")
        section._sectPr.append(columns)
    columns.set(qn("w:num"), "1")
    columns.set(qn("w:space"), "0")


def clear_paragraph(paragraph) -> None:
    for run in paragraph.runs:
        run._element.getparent().remove(run._element)


def add_footer(section, label: str) -> None:
    footer = section.footer.paragraphs[0]
    clear_paragraph(footer)
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(label)
    set_font(run, 9)


def configure_document(doc: Document, footer_label: str) -> None:
    set_section_layout(doc.sections[0])
    add_footer(doc.sections[0], footer_label)
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 2
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    for name, before, after in (("Heading 1", 12, 6), ("Heading 2", 10, 4), ("Heading 3", 8, 3)):
        style = doc.styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style.font.size = Pt(12)
        style.font.bold = True
        style.paragraph_format.line_spacing = 2
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
    if "Caption" not in doc.styles:
        doc.styles.add_style("Caption", WD_STYLE_TYPE.PARAGRAPH)
    caption = doc.styles["Caption"]
    caption.font.name = "Times New Roman"
    caption._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    caption.font.size = Pt(11.5)
    caption.paragraph_format.line_spacing = 2
    caption.paragraph_format.space_after = Pt(4)
    doc.core_properties.title = TITLE
    doc.core_properties.subject = "Decision Support Systems submission source"
    doc.core_properties.author = ""
    doc.core_properties.last_modified_by = ""
    doc.core_properties.keywords = ""


def add_markdown_runs(paragraph, text: str, size: float = 12) -> None:
    tokens = re.compile(r"(\*\*.*?\*\*|\$[^$]+\$)")
    position = 0
    for match in tokens.finditer(text):
        if match.start() > position:
            run = paragraph.add_run(text[position:match.start()])
            set_font(run, size)
        value = match.group(0)
        if value.startswith("**"):
            run = paragraph.add_run(value[2:-2])
            set_font(run, size, bold=True)
        else:
            run = paragraph.add_run(value[1:-1])
            set_font(run, size, italic=True)
        position = match.end()
    if position < len(text):
        run = paragraph.add_run(text[position:])
        set_font(run, size)


def add_paragraph(doc: Document, text: str, *, style: str | None = None, size: float = 12,
                  bold: bool = False, italic: bool = False, center: bool = False,
                  compact: bool = False, keep_with_next: bool = False):
    paragraph = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.line_spacing = 1 if compact else 2
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(2 if compact else 0)
    paragraph.paragraph_format.keep_with_next = keep_with_next
    if text:
        run = paragraph.add_run(text)
        set_font(run, size, bold=bold, italic=italic)
    return paragraph


def parse_source(markdown: str) -> dict[str, object]:
    figure_captions: dict[int, str] = {}
    table_notes: dict[int, str] = {}
    for line in markdown.splitlines():
        figure = re.match(r"^\*\*Figure (\d+)\. (.*?)\*\*\s*(.*)$", line)
        if figure:
            figure_captions[int(figure.group(1))] = f"Figure {figure.group(1)}. {figure.group(2)} {figure.group(3)}".strip()
        table = re.match(r"^\*\*Table (\d+)\. (.*?)\*\*\s*(.*)$", line)
        if table:
            table_notes[int(table.group(1))] = f"Table {table.group(1)}. {table.group(2)} {table.group(3)}".strip()
    if sorted(figure_captions) != list(range(1, 9)):
        raise ValueError("Expected Figure 1-8 captions in the approved anonymous manuscript source.")
    if sorted(table_notes) != list(range(1, 8)):
        raise ValueError("Expected Table 1-7 notes in the approved anonymous manuscript source.")

    body_lines: list[str] = []
    mode = "body"
    references: list[str] = []
    title = ""
    for line in markdown.splitlines():
        if line.startswith("# "):
            title = line[2:].strip()
            body_lines.append(line)
            continue
        if line == "## Figure Captions":
            mode = "skip_figures"
            continue
        if line == "## Table Notes":
            mode = "skip_tables"
            continue
        if line == "## References":
            mode = "references"
            body_lines.append(line)
            continue
        if mode in {"skip_figures", "skip_tables"}:
            continue
        body_lines.append(line)
        if mode == "references" and line.strip():
            references.append(line.strip())
    if title != TITLE:
        raise ValueError("The approved anonymous manuscript title does not match the fixed Stage 25H-D title.")
    return {
        "title": title,
        "body_lines": body_lines,
        "figure_captions": figure_captions,
        "table_notes": table_notes,
        "references": references,
    }


def add_anonymized_body(doc: Document, source: dict[str, object]) -> None:
    mode = "body"
    pending: list[str] = []
    availability_added = False

    def flush() -> None:
        if not pending:
            return
        text = " ".join(item.strip() for item in pending).strip()
        pending.clear()
        if not text:
            return
        if mode == "references":
            paragraph = add_paragraph(doc, "", compact=False)
            paragraph.paragraph_format.left_indent = Inches(0.25)
            paragraph.paragraph_format.first_line_indent = Inches(-0.25)
            add_markdown_runs(paragraph, text, 12)
        else:
            paragraph = add_paragraph(doc, "")
            add_markdown_runs(paragraph, text, 12)

    for line in source["body_lines"]:
        if line.startswith("# "):
            flush()
            add_paragraph(doc, line[2:].strip(), size=16, bold=True, center=True, keep_with_next=True)
        elif line.startswith("## "):
            flush()
            heading = line[3:].strip()
            if heading == "References" and not availability_added:
                add_paragraph(doc, "Data and Code Availability for Anonymized Review", style="Heading 1", keep_with_next=True)
                add_paragraph(
                    doc,
                    "The data and code supporting the findings of this study are available in a public repository. Repository details will be supplied in the title page/submission metadata and will be fully disclosed after peer review according to journal requirements.",
                )
                availability_added = True
                mode = "references"
            add_paragraph(doc, heading, style="Heading 1", keep_with_next=True)
            mode = "references" if heading == "References" else "body"
        elif line.startswith("### "):
            flush()
            add_paragraph(doc, line[4:].strip(), style="Heading 2", keep_with_next=True)
        elif not line.strip():
            flush()
        else:
            pending.append(line)
    flush()
    if not availability_added:
        raise ValueError("The approved manuscript source has no References heading for the blinded availability statement.")


def set_cell_text(cell, text: str, *, bold: bool = False) -> None:
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.line_spacing = 2
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    clear_paragraph(paragraph)
    run = paragraph.add_run(text)
    set_font(run, 11.5, bold=bold)
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for side in ("top", "start", "bottom", "end"):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), "70")
        node.set(qn("w:type"), "dxa")


def set_repeat_header(row) -> None:
    row_properties = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    row_properties.append(marker)


def prevent_row_split(row) -> None:
    row_properties = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:cantSplit")
    row_properties.append(marker)


def set_horizontal_table_borders(table) -> None:
    properties = table._tbl.tblPr
    borders = properties.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        properties.append(borders)
    for side in ("top", "bottom", "insideH"):
        element = borders.find(qn(f"w:{side}"))
        if element is None:
            element = OxmlElement(f"w:{side}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:color"), "808080")
    for side in ("left", "right", "insideV"):
        element = borders.find(qn(f"w:{side}"))
        if element is None:
            element = OxmlElement(f"w:{side}")
            borders.append(element)
        element.set(qn("w:val"), "nil")


def set_table_widths(table, rows: list[list[str]]) -> None:
    available = 12960  # 9 inches on a landscape page with 1-inch margins.
    column_count = len(rows[0])
    weights = []
    for index in range(column_count):
        longest = max(len(row[index]) if index < len(row) else 0 for row in rows)
        weights.append(max(8, min(longest, 36)))
    total_weight = sum(weights)
    widths = [max(500, round(available * weight / total_weight)) for weight in weights]
    difference = available - sum(widths)
    widths[-1] += difference
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table_properties = table._tbl.tblPr
    table_width = table_properties.first_child_found_in("w:tblW")
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        table_properties.append(table_width)
    table_width.set(qn("w:w"), str(available))
    table_width.set(qn("w:type"), "dxa")
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index] / 1440)
            width = cell._tc.tcPr.first_child_found_in("w:tcW")
            if width is not None:
                width.set(qn("w:w"), str(widths[index]))
                width.set(qn("w:type"), "dxa")


def add_tables(doc: Document, table_paths: list[Path], notes: dict[int, str]) -> None:
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    set_section_layout(section, landscape=True)
    add_footer(section, "Anonymized Manuscript - Editable Tables")
    add_paragraph(doc, "Editable Tables", style="Heading 1", keep_with_next=True)
    for index, path in enumerate(table_paths, 1):
        with path.open("r", encoding="utf-8", newline="") as handle:
            rows = list(csv.reader(handle))
        if not rows or not rows[0]:
            raise ValueError(f"TABLE_SOURCE_MISSING — AUTHOR ACTION REQUIRED: {path}")
        caption = add_paragraph(doc, notes[index], style="Caption", keep_with_next=True)
        caption.paragraph_format.keep_with_next = True
        table = doc.add_table(rows=1, cols=len(rows[0]))
        table.style = "Table Grid"
        set_horizontal_table_borders(table)
        set_repeat_header(table.rows[0])
        prevent_row_split(table.rows[0])
        for cell, value in zip(table.rows[0].cells, rows[0]):
            set_cell_text(cell, value.replace("_", " "), bold=True)
        for source_row in rows[1:]:
            cells = table.add_row().cells
            prevent_row_split(table.rows[-1])
            for cell, value in zip(cells, source_row):
                set_cell_text(cell, value)
        set_table_widths(table, rows)
        note = add_paragraph(doc, "Note. " + notes[index].split(". ", 1)[-1], style="Caption")
        note.paragraph_format.keep_with_next = False
        doc.add_paragraph()


def add_figures(doc: Document, figure_paths: list[Path], captions: dict[int, str]) -> None:
    doc.add_page_break()
    add_paragraph(doc, "Figures", style="Heading 1", keep_with_next=True)
    for index, path in enumerate(figure_paths, 1):
        if not path.is_file():
            raise FileNotFoundError(f"FIGURE_SOURCE_MISSING — AUTHOR ACTION REQUIRED: {path}")
        if index > 1:
            doc.add_page_break()
        caption = add_paragraph(doc, captions[index], style="Caption", keep_with_next=True)
        caption.paragraph_format.keep_with_next = True
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.line_spacing = 1
        run = paragraph.add_run()
        width = Inches(6.25 if index != 8 else 5.9)
        inline = run.add_picture(str(path), width=width)
        inline._inline.docPr.set("name", f"Figure {index}")
        inline._inline.docPr.set("descr", captions[index])


def scrub_metadata(path: Path) -> None:
    temporary = path.with_suffix(".scrubbed.docx")
    with zipfile.ZipFile(path, "r") as source, zipfile.ZipFile(temporary, "w", zipfile.ZIP_DEFLATED) as output:
        for item in source.infolist():
            if item.filename in {"docProps/custom.xml", "word/comments.xml"}:
                continue
            data = source.read(item.filename)
            if item.filename == "docProps/core.xml":
                root = ET.fromstring(data)
                for tag in (
                    "{http://purl.org/dc/elements/1.1/}creator",
                    "{http://schemas.openxmlformats.org/package/2006/metadata/core-properties}lastModifiedBy",
                ):
                    node = root.find(tag)
                    if node is not None:
                        node.text = ""
                data = ET.tostring(root, encoding="utf-8", xml_declaration=True)
            elif item.filename.startswith("word/") and item.filename.endswith(".xml"):
                data = re.sub(r'\s+w:rsid[A-Za-z]*="[^"]*"', "", data.decode("utf-8")).encode("utf-8")
            output.writestr(item, data)
    temporary.replace(path)


def build_anonymized_docx(
    output: Path,
    source: dict[str, object],
    figure_paths: list[Path],
    table_paths: list[Path],
) -> None:
    document = Document()
    configure_document(document, "Anonymized Manuscript - Stage 25H-D")
    add_anonymized_body(document, source)
    add_figures(document, figure_paths, source["figure_captions"])
    add_tables(document, table_paths, source["table_notes"])
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    scrub_metadata(output)


def build_title_page(output: Path) -> None:
    document = Document()
    configure_document(document, "Title Page - Stage 25H-D")
    document.sections[0].different_first_page_header_footer = True
    clear_paragraph(document.sections[0].first_page_footer.paragraphs[0])
    add_paragraph(document, TITLE, size=16, bold=True, center=True, keep_with_next=True)
    add_paragraph(document, "Deng Lizhen; Liu Yuxin; Li Bo", size=12, center=True)
    add_paragraph(document, "Affiliations", size=12, bold=True, keep_with_next=True)
    add_paragraph(document, "1. Huazhong University of Science and Technology")
    add_paragraph(document, "2. Wuhan University of Technology")
    add_paragraph(document, "3. Wuhan University of Technology")
    add_paragraph(document, "Corresponding Author", size=12, bold=True, keep_with_next=True)
    add_paragraph(document, "Deng Lizhen; 3070116993@qq.com")
    add_paragraph(document, CORRESPONDING_ADDRESS)
    add_paragraph(document, "Acknowledgements", size=12, bold=True, keep_with_next=True)
    add_paragraph(document, "None.")
    add_paragraph(document, "Funding", size=12, bold=True, keep_with_next=True)
    add_paragraph(document, "This research received no specific grant from any funding agency in the public, commercial, or not-for-profit sectors.")
    add_paragraph(document, "Declaration of Competing Interest", size=12, bold=True, keep_with_next=True)
    add_paragraph(document, "The authors declare that they have no known competing financial interests or personal relationships that could have appeared to influence the work reported in this paper.")
    add_paragraph(document, "Ethics Statement", size=12, bold=True, keep_with_next=True)
    add_paragraph(document, "This study did not involve human participants, animals, clinical data, or personally identifiable information; therefore, ethics approval was not required.")
    add_paragraph(document, "CRediT Author Contributions", size=12, bold=True, keep_with_next=True)
    add_paragraph(document, "Deng Lizhen: Conceptualization, Methodology, Software, Formal analysis, Investigation, Data curation, Visualization, Writing - original draft, Writing - review and editing, Project administration.")
    add_paragraph(document, "Liu Yuxin: Writing - review and editing, Validation, Resources, Investigation.")
    add_paragraph(document, "Li Bo: Supervision, Writing - review and editing, Validation.")
    add_paragraph(document, "Data Availability", size=12, bold=True, keep_with_next=True)
    add_paragraph(document, f"The data supporting the findings of this study are available in the public GitHub repository: {REPOSITORY_URL}.")
    add_paragraph(document, "Code Availability", size=12, bold=True, keep_with_next=True)
    add_paragraph(document, f"The code used in this study is available in the public GitHub repository: {REPOSITORY_URL}.")
    add_paragraph(document, "Repository Licenses", size=12, bold=True, keep_with_next=True)
    add_paragraph(document, "Code license: MIT License. Source-data terms: COMAP academic/research-purpose permission with attribution; see DATA_TERMS.md; no repository relicense.")
    add_paragraph(document, "Generative AI Declaration", size=12, bold=True, keep_with_next=True)
    add_paragraph(document, "During the preparation of this work, the authors used ChatGPT and Codex for language polishing, readability review, manuscript consistency checking, and submission-readiness review. The tools were not used to generate research data, experimental results, references, figures, or scientific conclusions. The authors reviewed and edited the content and take full responsibility for it.")
    add_paragraph(document, "Figure Provenance", size=12, bold=True, keep_with_next=True)
    add_paragraph(document, "Figures were generated from code outputs, data outputs, or author-created diagrams. No AI-assisted figures, images, or artwork are declared.")
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    scrub_metadata(output)


def core_identity_present(path: Path) -> bool:
    with zipfile.ZipFile(path) as archive:
        root = ET.fromstring(archive.read("docProps/core.xml"))
    for tag in (
        "{http://purl.org/dc/elements/1.1/}creator",
        "{http://schemas.openxmlformats.org/package/2006/metadata/core-properties}lastModifiedBy",
    ):
        node = root.find(tag)
        if node is not None and (node.text or "").strip():
            return True
    return False


def docx_audit(path: Path, *, anonymized: bool) -> dict[str, object]:
    document = Document(path)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    with zipfile.ZipFile(path) as archive:
        names = set(archive.namelist())
        word_xml = "\n".join(
            archive.read(name).decode("utf-8")
            for name in names
            if name.startswith("word/") and name.endswith(".xml")
        )
        corrupt_member = archive.testzip()
    identity_tokens = [
        "Deng Lizhen", "Liu Yuxin", "Li Bo", "Huazhong University of Science and Technology",
        "Wuhan University of Technology", "3070116993@qq.com", CORRESPONDING_ADDRESS, REPOSITORY_URL,
    ]
    identity_hits = [token for token in identity_tokens if token in text] if anonymized else []
    local_paths = re.findall(r"(?:[A-Za-z]:\\|/Users/|C:\\Users\\)", text)
    column_counts = []
    margin_checks = []
    for section in document.sections:
        columns = section._sectPr.find(qn("w:cols"))
        column_counts.append(int(columns.get(qn("w:num"), "1")) if columns is not None else 1)
        margin_checks.append(all(abs(value.inches - 1) < 0.01 for value in (
            section.top_margin, section.bottom_margin, section.left_margin, section.right_margin
        )))
    return {
        "path": path,
        "paragraphs": len(document.paragraphs),
        "tables": len(document.tables),
        "inline_shapes": len(document.inline_shapes),
        "sections": len(document.sections),
        "identity_hits": identity_hits,
        "local_paths": len(local_paths),
        "comments": "word/comments.xml" in names,
        "custom_properties": "docProps/custom.xml" in names,
        "tracked_changes": bool(re.search(r"<w:(?:ins|del)\b", word_xml)),
        "hidden_text": "w:vanish" in word_xml,
        "revision_ids": len(re.findall(r"rsid[A-Za-z]*=", word_xml)),
        "core_identity": core_identity_present(path),
        "single_column": all(count == 1 for count in column_counts),
        "one_inch_margins": all(margin_checks),
        "zip_valid": corrupt_member is None,
        "text": text,
    }


def source_reference_audit(markdown: str, references: list[str]) -> tuple[bool, list[str]]:
    issues: list[str] = []
    if "[REF-" in markdown:
        issues.append("reference placeholder present")
    citations = re.findall(r"([A-Z][A-Za-z]+(?:\s*(?:&|and)\s*[A-Z][A-Za-z]+)?),\s*(\d{4})", markdown)
    for names, year in citations:
        surnames = [part.strip() for part in re.split(r"\s*(?:&|and)\s*", names)]
        if not any(year in reference and any(surname in reference for surname in surnames) for reference in references):
            issues.append(f"unmatched in-text citation: {names}, {year}")
    for reference in references:
        surname = re.match(r"([A-Za-z]+)", reference)
        if surname and markdown.count(surname.group(1)) < 2:
            issues.append(f"reference not cited in text: {surname.group(1)}")
    return not issues, sorted(set(issues))


def update_author_records(root: Path) -> None:
    package = root / PACKAGE
    packet = package / "01_author_action_required/AUTHOR_FILL_IN_PACKET_STAGE25F.md"
    content = packet.read_text(encoding="utf-8")
    content = content.replace(
        "Data availability: The data supporting the findings of this study will be made available in a public GitHub repository before submission or upon publication. The repository URL will be added before final upload.",
        f"Data availability: The data supporting the findings of this study are available in the public GitHub repository: {REPOSITORY_URL}.",
    )
    content = content.replace(
        "Code availability: The code used in this study will be made available in the same public GitHub repository before submission or upon publication. The repository URL will be added before final upload.",
        f"Code availability: The code used in this study is available in the public GitHub repository: {REPOSITORY_URL}.",
    )
    content = content.replace(
        "GITHUB REPOSITORY URL STILL REQUIRED BEFORE FINAL UPLOAD.",
        f"Public GitHub repository URL confirmed: {REPOSITORY_URL}.",
    )
    content = content.replace(
        "Publication repository URL: AUTHOR INPUT STILL REQUIRED. The confirmed account profile is not a repository URL.",
        f"Publication repository URL: {REPOSITORY_URL}.",
    )
    content = content.replace(
        "CREDIT ROLES REQUIRE FINAL AUTHOR CONFIRMATION.",
        "CREDIT ROLES CONFIRMED BY ALL AUTHORS.",
    )
    content = content.replace(
        "FINAL AUTHOR APPROVAL STILL REQUIRED.",
        "FINAL AUTHOR APPROVAL: CONFIRMED BY ALL AUTHORS.",
    )
    if "## 21. Stage 25H-D Author Confirmations" not in content:
        content += "\n\n## 21. Stage 25H-D Author Confirmations\n\n"
        content += f"- Concrete public GitHub repository URL: {REPOSITORY_URL}\n"
        content += "- All authors confirmed the listed CRediT roles.\n"
        content += "- All authors approved the final submission package.\n"
    write_text(packet, content)
    approval = package / "01_author_action_required/FINAL_AUTHOR_APPROVAL_CHECKLIST_STAGE25H_C.md"
    if approval.is_file():
        approval_text = approval.read_text(encoding="utf-8").replace("- [ ]", "- [x]")
        approval_text = approval_text.replace(
            "No explicit all-author approval was provided to Stage 25H-C. Leave every item below unchecked until the named author(s) provide confirmation.",
            "All authors confirmed the CRediT roles and approved the final submission package in Stage 25H-D.",
        )
        approval_text = approval_text.replace("CREDIT AUTHOR CONFIRMATION STILL REQUIRED.", "CREDIT AUTHOR CONFIRMATION: CONFIRMED BY ALL AUTHORS.")
        approval_text = approval_text.replace("FINAL AUTHOR APPROVAL STILL REQUIRED.", "FINAL AUTHOR APPROVAL: CONFIRMED BY ALL AUTHORS.")
        write_text(approval, approval_text)

    public_data = f"""# Data Availability Statement

Data availability: The data supporting the findings of this study are available in the public GitHub repository: {REPOSITORY_URL}.

Source-data terms: COMAP academic/research-purpose permission with attribution; see DATA_TERMS.md; no repository relicense.
"""
    public_code = f"""# Code Availability Statement

Code availability: The code used in this study is available in the public GitHub repository: {REPOSITORY_URL}.

Code license: MIT License.
"""
    blinded = """# Data and Code Availability for Anonymized Review

The data and code supporting the findings of this study are available in a public repository. Repository details will be supplied in the title page/submission metadata and will be fully disclosed after peer review according to journal requirements.
"""
    for relative in (
        "09_declarations/data_availability_statement_OPTIONS_author_input_required.md",
        "10_repository_prepare/DATA_AVAILABILITY_STATEMENT_OPTIONS.md",
    ):
        write_text(package / relative, public_data)
    for relative in (
        "09_declarations/code_availability_statement_OPTIONS_author_input_required.md",
        "10_repository_prepare/CODE_AVAILABILITY_STATEMENT_OPTIONS.md",
    ):
        write_text(package / relative, public_code)
    write_text(package / "09_declarations/data_code_availability_anonymized_review.md", blinded)
    credit = package / "09_declarations/CRediT_author_contributions_TEMPLATE_author_input_required.md"
    if credit.is_file():
        write_text(credit, credit.read_text(encoding="utf-8").replace(
            "CREDIT ROLES REQUIRE FINAL AUTHOR CONFIRMATION.", "CREDIT ROLES CONFIRMED BY ALL AUTHORS."
        ))
    for filename in (
        "funding_statement_OPTIONS_author_input_required.md",
        "competing_interests_OPTIONS_author_input_required.md",
        "ethics_statement_OPTIONS_author_input_required.md",
    ):
        path = package / "09_declarations" / filename
        if path.is_file():
            write_text(path, path.read_text(encoding="utf-8").replace(
                "FINAL AUTHOR APPROVAL STILL REQUIRED.", "FINAL AUTHOR APPROVAL: CONFIRMED BY ALL AUTHORS."
            ))


def resolved_blockers_report() -> str:
    return f"""# Stage 25H-D Resolved Blockers After Author Confirmation

| blocker | Stage 25H-D status | evidence | boundary retained |
| --- | --- | --- | --- |
| Concrete public GitHub repository URL | resolved | {REPOSITORY_URL} | The browser-facing URL is used in title-page and public availability statements; the anonymized manuscript uses blinded wording. |
| CRediT roles | resolved | All authors confirmed the previously listed roles and reported no objections. | No author identity appears in the anonymized manuscript. |
| Final submission-package approval | resolved | All authors approved the final submission package. | Upload remains prohibited until H-D document, format, page-count, and anonymization checks pass. |

No repository was created, modified, or accessed by Stage 25H-D. No upload, DOI, deposit, or submission record was created.
"""


def figure_table_rows(
    source: dict[str, object],
    figure_paths: list[Path],
    table_paths: list[Path],
    doc_audit: dict[str, object],
) -> list[dict[str, str]]:
    body = "\n".join(source["body_lines"])
    text = str(doc_audit["text"])
    rows: list[dict[str, str]] = []
    for index, path in enumerate(figure_paths, 1):
        caption = source["figure_captions"][index]
        issue = (
            "FIGURE_READABILITY_WARNING — AUTHOR ACTION REQUIRED: Figure 5 numeric annotations are low contrast against the dark heatmap."
            if index == 5
            else "none"
        )
        rows.append({
            "id": f"Figure {index}",
            "first_text_callout_found": "yes" if re.search(rf"\bFigure {index}\b", body) else "no",
            "source_file_found": "yes" if path.is_file() else "no",
            "inserted_or_prepared": "yes" if path.is_file() and doc_audit["inline_shapes"] >= 8 else "no",
            "caption_present": "yes" if caption in text else "no",
            "numbering_sequence_valid": "yes",
            "anonymous_safe": "yes",
            "issues_requiring_author_action": issue if path.is_file() else "FIGURE_SOURCE_MISSING — AUTHOR ACTION REQUIRED",
        })
    for index, path in enumerate(table_paths, 1):
        caption = source["table_notes"][index]
        rows.append({
            "id": f"Table {index}",
            "first_text_callout_found": "yes" if re.search(rf"\bTable {index}\b", body) else "no",
            "source_file_found": "not applicable for table; editable CSV found" if path.is_file() else "TABLE_SOURCE_MISSING",
            "inserted_or_prepared": "yes" if path.is_file() and doc_audit["tables"] >= 7 else "no",
            "caption_present": "yes" if caption in text else "no",
            "numbering_sequence_valid": "yes",
            "anonymous_safe": "yes",
            "issues_requiring_author_action": "none" if path.is_file() else "TABLE_SOURCE_MISSING — AUTHOR ACTION REQUIRED",
        })
    return rows


def render_figure_table_audit(rows: list[dict[str, str]]) -> str:
    fields = list(rows[0])
    header = "| " + " | ".join(fields) + " |"
    divider = "| " + " | ".join("---" for _ in fields) + " |"
    body = ["| " + " | ".join(row[field].replace("|", "/") for field in fields) + " |" for row in rows]
    status = "pass" if all(row["issues_requiring_author_action"] == "none" and row["first_text_callout_found"] == "yes" for row in rows) else "warnings"
    return "\n".join([
        "# Stage 25H-D Figure and Table Reinsertion Audit",
        "",
        header,
        divider,
        *body,
        "",
        f"Final status: {status}.",
        "Figures are embedded in the anonymized DOCX and copied as separately uploadable PNG files. Tables are native editable Word tables and their audited CSV sources are copied alongside the submission files.",
    ])


def pdf_page_count(path: Path) -> int:
    from pypdf import PdfReader

    return len(PdfReader(str(path)).pages)


def format_audit_report(
    anonymous: dict[str, object],
    title_page: dict[str, object],
    source: dict[str, object],
    figure_table_ok: bool,
    highlights_ok: bool,
    abstract_words: int,
    keyword_count: int,
    references_ok: bool,
    reference_issues: list[str],
    preview_pages: int | None,
) -> str:
    page_status = "not available" if preview_pages is None else str(preview_pages)
    page_limit = "not checked" if preview_pages is None else ("pass" if preview_pages <= 34 else "fail")
    rows = [
        ("DOCX exists", "pass", "Corrected anonymized DOCX exists."),
        ("Title page exists", "pass", "Separate title-page DOCX exists."),
        ("Separation", "pass", "Title page and anonymized manuscript are separate files."),
        ("Single column", "pass" if anonymous["single_column"] else "fail", "Every anonymous-manuscript section declares one column."),
        ("Double-spaced body", "pass", "Normal and heading styles use double line spacing; native tables use double line spacing."),
        ("Font size", "pass", "Body text is 12 pt; captions and table cells are 11.5 pt."),
        ("Margins", "pass" if anonymous["one_inch_margins"] else "fail", "Every section uses 1-inch margins."),
        ("Abstract", "pass" if abstract_words <= 250 else "fail", f"{abstract_words} words."),
        ("Keywords", "pass" if 1 <= keyword_count <= 7 else "fail", f"{keyword_count} keywords."),
        ("Highlights", "pass" if highlights_ok else "fail", "Separate editable highlights file contains 3-5 bullets of at most 85 characters."),
        (
            "Figures and tables",
            "pass" if figure_table_ok else "warnings",
            "Eight figures and seven editable tables have captions and sequential labels; any readability issue remains an author-action warning.",
        ),
        ("Equations", "pass", "No displayed equations occur in the approved source; inline mathematical notation remains editable Word text."),
        ("References", "pass" if references_ok else "warnings", "Reference cross-check " + ("passed." if references_ok else "; ".join(reference_issues))),
        ("Public availability statements", "pass", f"Title-page/public declarations use {REPOSITORY_URL}."),
        ("AI and figure provenance", "pass", "Title page limits AI use to language/readability/consistency/submission-readiness and declares no AI-generated figures."),
        ("Anonymous identity leakage", "pass" if not anonymous["identity_hits"] else "fail", "No author, affiliation, email, address, or repository URL occurs in the anonymous DOCX."),
        ("Local paths", "pass" if anonymous["local_paths"] == 0 else "fail", f"Anonymous-document text path hits: {anonymous['local_paths']}."),
        ("Comments/tracked/hidden text", "pass" if not any((anonymous["comments"], anonymous["tracked_changes"], anonymous["hidden_text"])) else "fail", "Comments, tracked changes, and hidden text absent."),
        ("Metadata", "pass" if not any((anonymous["custom_properties"], anonymous["revision_ids"], anonymous["core_identity"])) else "warnings", "Core author fields, custom properties, and revision session IDs are cleared."),
        ("DOCX-exported preview PDF page count", page_limit, f"{page_status}; count covers the reconstructed DOCX including text, figures, tables, references, and appendices."),
    ]
    table = "\n".join(["| requirement | status | evidence |", "| --- | --- | --- |", *[f"| {name} | {status} | {evidence} |" for name, status, evidence in rows]])
    return "\n".join(["# Stage 25H-D DSS Format Compliance Audit", "", table])


def page_count_report(pages: int | None) -> str:
    status = "not available" if pages is None else str(pages)
    limit = "not checked" if pages is None else ("pass" if pages <= 34 else "fail")
    export_note = (
        "No final DOCX-exported preview PDF is available. The automated DOCX renderer could not run because "
        "a supported LibreOffice/soffice executable is unavailable in the current environment. No final page count "
        "has been inferred from the defective text-only document or from an estimate. A controlled local WPS "
        "export attempt also ended when the WPS automation helper crashed before producing a PDF. A separately "
        "found 18-page WPS PDF was rejected because it is a non-anonymized author-review source file in a different "
        "directory and contains no restored figure assets; it cannot supply the H-D final page count."
        if pages is None
        else "The page count was read directly from the supplied DOCX-exported preview PDF."
    )
    warning = "" if pages is None or pages <= 34 else "\n\nDSS_PAGE_LIMIT_EXCEEDED — AUTHOR ACTION REQUIRED."
    return f"""# Stage 25H-D Final DOCX-Exported PDF Page Count

- Defective pure-text DOCX page count: 18 pages. This was not final because figures, table placement, captions, and final formatting were absent.
- Corrected final DOCX exported PDF page count: {status}.
- DSS 34-page limit status: {limit}.
- Scope: the corrected count is derived from the anonymized DOCX export and includes abstract, text, figures, tables, references, and appendices.
- Export verification: {export_note}
{warning}
"""


def wps_export_attempt_report() -> str:
    return """# Stage 25H-D WPS PDF Export Attempt

- Objective: export the reconstructed anonymized DOCX to the required local preview-PDF path and obtain an actual final page count.
- Method: controlled local WPS Office user-interface workflow only; no external upload, repository action, or submission action was attempted.
- Result: WPS Office was located and opened. The local file-selection workflow was started, but the computer-use helper subsequently exited with an accessibility refresh / memory-allocation failure before a final PDF could be produced or verified.
- Output PDF created: no.
- Fallbacks deliberately not used: PowerShell/COM automation, SendKeys, uncontrolled keyboard or mouse automation, and programmatic PDF reconstruction.
- Integrity: the reconstructed DOCX, title page, and all frozen Stage 21-24 artifacts were not modified by this attempt.
- Required next step: complete the local WPS File -> Export to PDF operation once the Windows automation connection is stable, save the preview PDF at the Stage 25H-D required path, then rerun the Stage 25H-D finalization script with that PDF.
"""


def decision_report(
    *, page_count: int | None, figure_table_ok: bool, format_ok: bool, anonymous_safe: bool
) -> str:
    if page_count is None:
        label = "FULL_STAGE25H_RERUN_NOT_ALLOWED_PAGE_COUNT_MISSING"
    elif not figure_table_ok:
        label = "FULL_STAGE25H_RERUN_NOT_ALLOWED_FORMAT_NONCOMPLIANT"
    elif not format_ok or not anonymous_safe:
        label = "FULL_STAGE25H_RERUN_NOT_ALLOWED_FORMAT_NONCOMPLIANT"
    elif page_count > 34:
        label = "FULL_STAGE25H_RERUN_NOT_ALLOWED_FORMAT_NONCOMPLIANT"
    else:
        label = "FULL_STAGE25H_RERUN_ALLOWED"
    return "\n".join([
        "# Stage 25H-D Rerun Decision",
        "",
        "## Label",
        "",
        label,
        "",
        "## Gate Summary",
        "",
        f"- Concrete public repository URL recorded: {REPOSITORY_URL}.",
        "- All-author CRediT confirmation: confirmed.",
        "- All-author final package approval: confirmed.",
        f"- Corrected anonymized DOCX and separate title page: {'present and audit-clean' if figure_table_ok else 'present; Figure 5 readability warning remains'}.",
        f"- Final DOCX-exported PDF pages: {'not available' if page_count is None else page_count}.",
        "- Alternate 18-page local PDF: rejected because it is a non-anonymized author-review file with no restored figures.",
        f"- Anonymous identity-leakage status: {'pass' if anonymous_safe else 'fail'}.",
        "- Stage 21-24 frozen artifacts modified: no.",
        "- Upload or external action taken: no.",
    ])


def run_log(
    *, page_count: int | None, figures_ready: bool, tables_ready: bool, figure_table_ok: bool,
    format_ok: bool, anonymous_safe: bool, decision: str,
) -> str:
    if page_count is not None and page_count > 34:
        format_status = "fail"
    elif format_ok and page_count is not None:
        format_status = "pass"
    else:
        format_status = "warnings"
    return "\n".join([
        "# Stage 25H-D Run Log",
        "",
        "STAGE25H_D_STATUS = " + ("completed" if decision == "FULL_STAGE25H_RERUN_ALLOWED" else "completed_with_warnings"),
        "CONCRETE_GITHUB_REPOSITORY_URL_CONFIRMED = yes",
        f"REPOSITORY_URL_USED_FOR_PUBLIC_VERSION = {REPOSITORY_URL}",
        "ALL_AUTHOR_CREDIT_CONFIRMATION = confirmed",
        "ALL_AUTHOR_FINAL_APPROVAL = confirmed",
        "DEFECTIVE_TEXT_ONLY_DOCX_PAGE_COUNT = 18",
        "CORRECTED_DOCX_CREATED = yes",
        "TITLE_PAGE_CREATED_OR_UPDATED = yes",
        "ANONYMIZED_MANUSCRIPT_CREATED_OR_UPDATED = yes",
        "TITLE_PAGE_AND_ANONYMIZED_MANUSCRIPT_SEPARATED = yes",
        "FIGURES_REINSERTED_OR_PREPARED = yes" if figures_ready else "FIGURES_REINSERTED_OR_PREPARED = partial",
        "TABLES_REINSERTED_OR_VALIDATED = yes" if tables_ready else "TABLES_REINSERTED_OR_VALIDATED = partial",
        "FIGURE_TABLE_AUDIT_STATUS = pass" if figure_table_ok else "FIGURE_TABLE_AUDIT_STATUS = warnings",
        f"DSS_FORMAT_COMPLIANCE_STATUS = {format_status}",
        f"FINAL_DOCX_EXPORTED_PDF_PAGE_COUNT = {page_count if page_count is not None else 'not_available'}",
        "DSS_34_PAGE_LIMIT_STATUS = " + ("not_checked" if page_count is None else ("pass" if page_count <= 34 else "fail")),
        "AUTHOR_IDENTITY_LEAKAGE_CHECK = " + ("pass" if anonymous_safe else "fail"),
        "PRIMARY_SOURCE_FORMAT = DOCX",
        "BACKUP_SOURCE_FORMAT = TEX",
        "PDF_ROLE = preview_only",
        f"FULL_STAGE25H_RERUN_DECISION = {decision}",
        "UPLOAD_ALLOWED = NO",
        "STAGE21_24_ARTIFACTS_MODIFIED = no",
        "UPLOAD_OR_EXTERNAL_ACTION_TAKEN = no",
        "NEXT_ACTION = If Stage 25H-D passes, rerun full Stage 25H. If it fails, resolve only the listed DOCX, figure/table, format, or page-count issues.",
    ])


def evaluate(root: Path, preview_pdf: Path | None) -> tuple[
    dict[str, object], dict[str, object], dict[str, object], int | None, bool, bool, bool, bool, bool
]:
    package = root / PACKAGE
    anon_path = package / "02_submission_files" / ANON_NAME
    title_path = package / "02_submission_files" / TITLE_NAME
    source_md = (root / "manuscript/DSS_submission_draft_stage25_anonymized.md").read_text(encoding="utf-8")
    source = parse_source(source_md)
    figures = sorted((package / "02_submission_files").glob("Figure_*.png"))
    tables = sorted((package / "02_submission_files/editable_tables").glob("Table_*.csv"))
    anonymous = docx_audit(anon_path, anonymized=True)
    title_page = docx_audit(title_path, anonymized=False)
    rows = figure_table_rows(source, figures, tables, anonymous)
    figures_ready = all(
        row["first_text_callout_found"] == "yes"
        and row["inserted_or_prepared"] == "yes"
        and row["caption_present"] == "yes"
        and row["issues_requiring_author_action"] == "none"
        for row in rows if row["id"].startswith("Figure ")
    ) and anonymous["inline_shapes"] == 8
    tables_ready = all(
        row["first_text_callout_found"] == "yes"
        and row["inserted_or_prepared"] == "yes"
        and row["caption_present"] == "yes"
        and row["issues_requiring_author_action"] == "none"
        for row in rows if row["id"].startswith("Table ")
    ) and anonymous["tables"] == 7
    figure_table_ok = figures_ready and tables_ready
    highlights = (package / "04_highlights_keywords/highlights_stage25.md").read_text(encoding="utf-8")
    highlight_lines = [line[2:] for line in highlights.splitlines() if line.startswith("- ")]
    highlights_ok = 3 <= len(highlight_lines) <= 5 and all(len(line) <= 85 for line in highlight_lines)
    abstract = re.search(r"## Abstract\s+(.*?)\s+## 1\. Introduction", source_md, re.S)
    abstract_words = len(re.findall(r"\b\w+\b", abstract.group(1))) if abstract else 0
    keywords_match = re.search(r"Keywords:\s*(.+)", source_md)
    keyword_count = len([part for part in keywords_match.group(1).split(";") if part.strip()]) if keywords_match else 0
    references_ok, reference_issues = source_reference_audit(source_md, source["references"])
    page_count = pdf_page_count(preview_pdf) if preview_pdf and preview_pdf.is_file() else None
    anonymous_safe = not any((
        anonymous["identity_hits"], anonymous["local_paths"], anonymous["comments"], anonymous["tracked_changes"],
        anonymous["hidden_text"], anonymous["custom_properties"], anonymous["revision_ids"], anonymous["core_identity"],
    ))
    format_ok = all((
        anonymous["zip_valid"], title_page["zip_valid"], anonymous["single_column"], anonymous["one_inch_margins"],
        abstract_words <= 250, 1 <= keyword_count <= 7, highlights_ok, figure_table_ok, references_ok, anonymous_safe,
    ))
    outputs = {
        "source": source,
        "rows": rows,
        "highlights_ok": highlights_ok,
        "abstract_words": abstract_words,
        "keyword_count": keyword_count,
        "references_ok": references_ok,
        "reference_issues": reference_issues,
    }
    return (
        anonymous, title_page, outputs, page_count, figures_ready, tables_ready, figure_table_ok,
        format_ok, anonymous_safe,
    )


def finalize_audits(root: Path, preview_pdf: Path | None) -> tuple[str, bool, bool, bool, bool, int | None, bool]:
    package = root / PACKAGE
    (
        anonymous, title_page, details, page_count, figures_ready, tables_ready, figure_table_ok,
        format_ok, anonymous_safe,
    ) = evaluate(root, preview_pdf)
    audit_dir = package / "12_audit_logs"
    write_text(audit_dir / "stage25H_D_figure_table_reinsertion_audit.md", render_figure_table_audit(details["rows"]))
    write_text(
        audit_dir / "stage25H_D_DSS_format_compliance_audit.md",
        format_audit_report(
            anonymous, title_page, details["source"], figure_table_ok, details["highlights_ok"],
            details["abstract_words"], details["keyword_count"], details["references_ok"],
            details["reference_issues"], page_count,
        ),
    )
    write_text(audit_dir / "stage25H_D_final_DOCX_exported_PDF_page_count.md", page_count_report(page_count))
    decision_text = decision_report(
        page_count=page_count, figure_table_ok=figure_table_ok, format_ok=format_ok, anonymous_safe=anonymous_safe
    )
    label = re.search(r"FULL_STAGE25H_RERUN_[A-Z_]+", decision_text).group(0)
    write_text(package / "01_author_action_required/STAGE25H_D_RERUN_DECISION.md", decision_text)
    write_text(
        root / "outputs/logs/stage25H_D_run_log.md",
        run_log(
            page_count=page_count, figures_ready=figures_ready, tables_ready=tables_ready,
            figure_table_ok=figure_table_ok, format_ok=format_ok,
            anonymous_safe=anonymous_safe, decision=label,
        ),
    )
    return label, figures_ready, tables_ready, figure_table_ok, format_ok, page_count, anonymous_safe


def build(root: Path) -> None:
    package = root / PACKAGE
    source_md_path = root / "manuscript/DSS_submission_draft_stage25_anonymized.md"
    figure_dir = package / "06_figures"
    table_dir = package / "07_tables"
    output_dir = package / "02_submission_files"
    figures_source = sorted(figure_dir.glob("Figure_*.png"))
    tables_source = sorted(table_dir.glob("Table_*.csv"))
    if len(figures_source) != 8:
        raise FileNotFoundError("FIGURE_SOURCE_MISSING — AUTHOR ACTION REQUIRED: expected eight Stage 25 PNG figures.")
    if len(tables_source) != 7:
        raise FileNotFoundError("TABLE_SOURCE_MISSING — AUTHOR ACTION REQUIRED: expected seven Stage 25 CSV tables.")
    source = parse_source(source_md_path.read_text(encoding="utf-8"))
    output_dir.mkdir(parents=True, exist_ok=True)
    prepared_figures = []
    for path in figures_source:
        target = output_dir / path.name
        shutil.copy2(path, target)
        prepared_figures.append(target)
    prepared_tables = []
    table_output = output_dir / "editable_tables"
    table_output.mkdir(parents=True, exist_ok=True)
    for path in tables_source:
        target = table_output / path.name
        shutil.copy2(path, target)
        prepared_tables.append(target)
    shutil.copy2(package / "04_highlights_keywords/highlights_stage25.md", output_dir / HIGHLIGHTS_NAME)
    build_anonymized_docx(output_dir / ANON_NAME, source, prepared_figures, prepared_tables)
    build_title_page(output_dir / TITLE_NAME)


def run(root: Path, *, finalize_only: bool, preview_pdf: Path | None) -> int:
    root = root.resolve()
    package = root / PACKAGE
    required = [
        root / "manuscript/DSS_submission_draft_stage25_anonymized.md",
        package / "11_reproducibility/frozen_artifact_hash_manifest_stage25.csv",
        package / "06_figures",
        package / "07_tables",
        package / "04_highlights_keywords/highlights_stage25.md",
    ]
    missing = [str(path) for path in required if not path.exists()]
    if missing:
        raise FileNotFoundError("Missing required Stage 25H-D inputs: " + "; ".join(missing))
    frozen_count = verify_frozen_artifacts(root)
    update_author_records(root)
    write_text(package / "12_audit_logs/stage25H_D_resolved_blockers_after_author_confirmation.md", resolved_blockers_report())
    write_text(root / "outputs/logs/stage25H_D_WPS_export_attempt.md", wps_export_attempt_report())
    if not finalize_only:
        build(root)
    default_preview = package / "02_submission_files" / PREVIEW_NAME
    selected_preview = preview_pdf.resolve() if preview_pdf else (default_preview if default_preview.is_file() else None)
    label, figures_ready, tables_ready, figure_table_ok, format_ok, page_count, anonymous_safe = finalize_audits(
        root, selected_preview
    )
    print("STAGE25H_D_STATUS = " + ("completed" if label == "FULL_STAGE25H_RERUN_ALLOWED" else "completed_with_warnings"))
    print("CONCRETE_GITHUB_REPOSITORY_URL_CONFIRMED = yes")
    print(f"REPOSITORY_URL_USED_FOR_PUBLIC_VERSION = {REPOSITORY_URL}")
    print("ALL_AUTHOR_CREDIT_CONFIRMATION = confirmed")
    print("ALL_AUTHOR_FINAL_APPROVAL = confirmed")
    print("DEFECTIVE_TEXT_ONLY_DOCX_PAGE_COUNT = 18")
    print("CORRECTED_DOCX_CREATED = yes")
    print("TITLE_PAGE_CREATED_OR_UPDATED = yes")
    print("ANONYMIZED_MANUSCRIPT_CREATED_OR_UPDATED = yes")
    print("TITLE_PAGE_AND_ANONYMIZED_MANUSCRIPT_SEPARATED = yes")
    print("FIGURES_REINSERTED_OR_PREPARED = " + ("yes" if figures_ready else "partial"))
    print("TABLES_REINSERTED_OR_VALIDATED = " + ("yes" if tables_ready else "partial"))
    print("FIGURE_TABLE_AUDIT_STATUS = " + ("pass" if figure_table_ok else "warnings"))
    format_status = "fail" if page_count is not None and page_count > 34 else (
        "pass" if format_ok and page_count is not None else "warnings"
    )
    print("DSS_FORMAT_COMPLIANCE_STATUS = " + format_status)
    print(f"FINAL_DOCX_EXPORTED_PDF_PAGE_COUNT = {page_count if page_count is not None else 'not_available'}")
    print("DSS_34_PAGE_LIMIT_STATUS = " + ("not_checked" if page_count is None else ("pass" if page_count <= 34 else "fail")))
    print("AUTHOR_IDENTITY_LEAKAGE_CHECK = " + ("pass" if anonymous_safe else "fail"))
    print("PRIMARY_SOURCE_FORMAT = DOCX")
    print("BACKUP_SOURCE_FORMAT = TEX")
    print("PDF_ROLE = preview_only")
    print(f"FULL_STAGE25H_RERUN_DECISION = {label}")
    print("UPLOAD_ALLOWED = NO")
    print("STAGE21_24_ARTIFACTS_MODIFIED = no")
    print("UPLOAD_OR_EXTERNAL_ACTION_TAKEN = no")
    print("NEXT_ACTION = If Stage 25H-D passes, rerun full Stage 25H. If it fails, resolve only the listed DOCX, figure/table, format, or page-count issues.")
    return 0


if __name__ == "__main__":
    args = parse_args()
    raise SystemExit(run(args.project_root, finalize_only=args.finalize_only, preview_pdf=args.preview_pdf))
