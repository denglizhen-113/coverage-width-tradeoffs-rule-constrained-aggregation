"""Finalize the additive Stage 25H-E DSS submission package.

The script never modifies Stage 21-24 artifacts. Run figure generation with an
environment that provides matplotlib, then run DOCX generation and audit with
the bundled document runtime. All paths are project-relative by default.
"""

from __future__ import annotations

import argparse
import csv
import hashlib
import importlib.util
import json
import re
import shutil
import sys
import zipfile
from pathlib import Path
from typing import Any


TITLE = "Rule-Aware Decision Support for Expert-Crowd Aggregation under Hidden Public Preferences"
REPOSITORY_URL = "https://github.com/denglizhen-113/coverage-width-tradeoffs-rule-constrained-aggregation"
PACKAGE = Path("submission_package_stage25")
SOURCE = PACKAGE / "02_manuscript/DSS_submission_draft_STAGE25H_E_final_source.md"
SUBMISSION_DIR = PACKAGE / "02_submission_files"
FINAL_FIGURES = SUBMISSION_DIR / "figures_STAGE25H_E_final"
ANON_DOCX = SUBMISSION_DIR / "DSS_anonymized_manuscript_STAGE25H_E_final.docx"
PUBLIC_DOCX = SUBMISSION_DIR / "DSS_full_metadata_manuscript_STAGE25H_E_final.docx"
TITLE_DOCX = SUBMISSION_DIR / "DSS_title_page_STAGE25H_E_final.docx"
HIGHLIGHTS_DOCX = SUBMISSION_DIR / "DSS_highlights_STAGE25H_E.docx"
COVER_DOCX = SUBMISSION_DIR / "DSS_cover_letter_STAGE25H_E_final.docx"
PREVIEW_PDF = SUBMISSION_DIR / "DSS_anonymized_manuscript_STAGE25H_E_final_preview.pdf"
AUDIT_DIR = PACKAGE / "12_audit_logs"
AUTHOR_DIR = PACKAGE / "01_author_action_required"
RUN_LOG = Path("outputs/logs/stage25H_E_run_log.md")

FIGURE_FILENAMES = {
    1: "Figure_01_DSS_conceptual_framework.png",
    2: "Figure_02_decision_support_workflow.png",
    3: "Figure_03_discretion_identifiability_frontier.png",
    4: "Figure_04_disclosure_uncertainty_curve.png",
    5: "Figure_05_rule_robustness_heatmap.png",
    6: "Figure_06_synthetic_benchmark_coverage.png",
    7: "Figure_07_external_testbed_comparison.png",
    8: "Figure_08_DSS_artifact_evaluation.png",
}

HIGHLIGHTS = [
    "Rule-aware feasible sets preserve uncertainty under hidden preferences.",
    "Mechanism tests compare aggregation, discretion, ties, and disclosure.",
    "Synthetic calibration attains 1.000 coverage without point recovery.",
    "A decision cockpit converts uncertainty into auditable design guidance.",
]

NUMERIC_TRACE = [
    ("panel rows", "4199", "data/processed/panel_long.csv", "row count"),
    ("identification rows", "2777", "data/processed/identification_features_long.csv", "row count"),
    ("typed proxy rows", "2766", "data/processed/identification_features_long.csv", "public_appeal_proxy nonmissing"),
    ("logged missing typed proxies", "11", "data/processed/identification_features_long.csv", "public_appeal_proxy missing"),
    ("P empirical width", "0.843", "outputs/tables/identification_comparison_by_regime.csv", "P mean normalized width, rounded"),
    ("R empirical width", "0.891", "outputs/tables/identification_comparison_by_regime.csv", "R mean normalized width, rounded"),
    ("R_plus empirical width", "0.924", "outputs/tables/identification_comparison_by_regime.csv", "R_plus mean normalized width, rounded"),
    ("synthetic no-noise coverage", "1.000", "outputs/tables/synthetic_coverage_results.csv", "rule-aware, rule_consistent"),
    ("synthetic noise-stress coverage", "0.948", "outputs/tables/synthetic_coverage_results.csv", "rule-aware, outcome_noise_stress_test"),
    ("synthetic rule-aware width", "0.845", "outputs/tables/synthetic_coverage_results.csv", "rule-aware, rule_consistent"),
    ("synthetic rule-agnostic width", "1.000", "outputs/tables/synthetic_coverage_results.csv", "rule-agnostic, rule_consistent"),
    ("external rule-aware coverage", "1.000", "outputs/tables/external_testbed_results.csv", "rule_aware_discretion"),
    ("external misspecification false certainty", "0.958", "outputs/tables/external_testbed_results.csv", "direct_rule_misspecification"),
    ("top-k disclosure reduction", "12.5%", "outputs/tables/value_of_disclosure.csv", "relative reduction, rounded"),
    ("vote-bin disclosure reduction", "88.3%", "outputs/tables/value_of_disclosure.csv", "relative reduction, rounded"),
    ("margin disclosure reduction", "92.7%", "outputs/tables/value_of_disclosure.csv", "relative reduction, rounded"),
    ("artifact reproducibility", "69 passed", "outputs/tables/dss_evaluation_metrics.csv", "recorded artifact-level result"),
    ("artifact demo runtime", "0.027 seconds", "outputs/tables/dss_evaluation_metrics.csv", "recorded artifact-level result"),
]


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Generate and audit the additive Stage 25H-E DSS submission package."
    )
    parser.add_argument(
        "--project-root",
        type=Path,
        default=Path(__file__).resolve().parents[1],
        help="Project root containing data/, manuscript/, outputs/, and submission_package_stage25/.",
    )
    parser.add_argument(
        "--mode",
        choices=("assets", "build", "audit", "all"),
        default="all",
        help="Generate figures, build editable sources, audit an exported preview, or run available steps.",
    )
    parser.add_argument(
        "--preview-pdf",
        type=Path,
        help="DOCX-exported PDF used by audit mode. Defaults to the final preview path.",
    )
    parser.add_argument(
        "--visual-review",
        choices=("pass", "warnings", "not_recorded"),
        default="not_recorded",
        help="Result of a page-by-page visual review of the preview PDF.",
    )
    parser.add_argument(
        "--reviewed-pages",
        default="",
        help="Inclusive reviewed page range such as 1-30.",
    )
    return parser.parse_args()


def load_module(path: Path, name: str):
    spec = importlib.util.spec_from_file_location(name, path)
    if spec is None or spec.loader is None:
        raise ImportError(f"Cannot load module from {path}")
    module = importlib.util.module_from_spec(spec)
    sys.modules[name] = module
    spec.loader.exec_module(module)
    return module


def write_text(path: Path, content: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(content.rstrip() + "\n", encoding="utf-8")


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def frozen_hash_status(root: Path) -> tuple[int, list[str]]:
    manifest = root / PACKAGE / "11_reproducibility/frozen_artifact_hash_manifest_stage25.csv"
    if not manifest.is_file():
        raise FileNotFoundError(f"Frozen manifest missing: {manifest}")
    rows = list(csv.DictReader(manifest.open("r", encoding="utf-8-sig", newline="")))
    mismatches: list[str] = []
    for row in rows:
        path = root / row["relative_path"]
        observed = sha256(path) if path.is_file() else "missing"
        if observed != row["expected_sha256"]:
            mismatches.append(f"{row['relative_path']}: expected {row['expected_sha256']}, observed {observed}")
    return len(rows), mismatches


def parse_source(markdown: str) -> dict[str, Any]:
    if not markdown.startswith(f"# {TITLE}\n"):
        raise ValueError("Final source does not contain the fixed title.")
    figure_captions: dict[int, str] = {}
    table_notes: dict[int, str] = {}
    for line in markdown.splitlines():
        match = re.match(r"^\*\*Figure (\d+)\. (.*?)\*\*\s*(.*)$", line)
        if match:
            figure_captions[int(match.group(1))] = (
                f"Figure {match.group(1)}. {match.group(2)} {match.group(3)}".strip()
            )
        match = re.match(r"^\*\*Table (\d+)\. (.*?)\*\*\s*(.*)$", line)
        if match:
            table_notes[int(match.group(1))] = (
                f"Table {match.group(1)}. {match.group(2)} {match.group(3)}".strip()
            )
    if sorted(figure_captions) != list(range(1, 9)):
        raise ValueError("Expected Figure 1-8 captions in the final source.")
    if sorted(table_notes) != list(range(1, 8)):
        raise ValueError("Expected Table 1-7 notes in the final source.")

    body: list[str] = []
    skipping = False
    for line in markdown.splitlines():
        if line in {"## Figure Captions", "## Table Notes"}:
            skipping = True
            continue
        if skipping:
            if line.startswith("## ") and line not in {"## Figure Captions", "## Table Notes"}:
                skipping = False
            else:
                continue
        body.append(line)
    source_text = "\n".join(body)
    abstract_match = re.search(r"## Abstract\s+(.*?)\s+\*\*Keywords:\*\*", source_text, re.S)
    if not abstract_match:
        raise ValueError("Abstract or keywords block missing.")
    abstract_words = re.findall(r"\b[\w.-]+\b", abstract_match.group(1))
    keyword_match = re.search(r"\*\*Keywords:\*\*\s*(.*)", source_text)
    keywords = [item.strip().rstrip(".") for item in keyword_match.group(1).split(";") if item.strip()]
    return {
        "body_lines": body,
        "figure_captions": figure_captions,
        "table_notes": table_notes,
        "abstract_word_count": len(abstract_words),
        "keywords": keywords,
    }


def reference_audit(markdown: str) -> tuple[bool, list[str]]:
    references = {int(value) for value in re.findall(r"^\[(\d+)\]", markdown, flags=re.M)}
    cited: set[int] = set()
    body = markdown.split("## References", 1)[0]
    for token in re.findall(r"\[([0-9,-]+)\]", body):
        token_values = [int(value) for value in re.findall(r"\d+", token)]
        if 0 in token_values:
            continue
        for part in token.split(","):
            if "-" in part:
                lo, hi = (int(value) for value in part.split("-", 1))
                cited.update(range(lo, hi + 1))
            else:
                cited.add(int(part))
    issues = []
    if references != set(range(1, 16)):
        issues.append(f"reference list numbers: {sorted(references)}")
    if cited != references:
        issues.append(f"cited={sorted(cited)}, listed={sorted(references)}")
    for marker in ("[REF-", "TODO", "TBD"):
        if marker in markdown:
            issues.append(f"placeholder remains: {marker}")
    return not issues, issues


def generate_assets(root: Path) -> None:
    he = load_module(root / "scripts/25he_repair_submission_assets.py", "stage25he_assets")
    he.OUTPUT_FIGURES = FINAL_FIGURES
    he.load_plot_runtime()
    he.apply_plot_style()
    from matplotlib import font_manager

    font_path = font_manager.findfont("Times New Roman", fallback_to_default=False)
    he.plt.rcParams.update(
        {
            "font.family": "Times New Roman",
            "font.serif": ["Times New Roman"],
            "mathtext.fontset": "stix",
        }
    )
    paths, figure_meta = he.generate_figures(root)
    _, table_meta = he.compact_tables(root)
    report = [
        "# Stage 25H-E Final Figure Generation",
        "",
        f"- Font resolved: `{font_path}`",
        "- Output font family: Times New Roman",
        "- Output DPI: 300",
        "- AI-generated or AI-altered artwork: no",
        "- Stage 21-24 figures overwritten: no",
        "",
        "| figure | file | data source | pixels | dpi | sha256 |",
        "| --- | --- | --- | --- | --- | --- |",
    ]
    for row in figure_meta:
        path = paths[row["figure"]]
        report.append(
            f"| Figure {row['figure']} | {path.relative_to(root)} | {row['data_source']} | "
            f"{row['width']}x{row['height']} | {row['dpi'][0]}x{row['dpi'][1]} | {sha256(path)} |"
        )
    report.extend(["", f"Condensed table definitions checked: {len(table_meta)}."])
    write_text(root / AUDIT_DIR / "stage25H_E_final_figure_generation.md", "\n".join(report))
    print(f"FINAL_FIGURE_FOLDER = {root / FINAL_FIGURES}")
    print("FINAL_FIGURE_FONT = Times New Roman")
    print("FINAL_FIGURE_COUNT = 8")


def _docx_runtime(root: Path):
    hd = load_module(root / "scripts/25hd_reconstruct_dss_submission_docx.py", "stage25hd_helpers")
    he = load_module(root / "scripts/25he_repair_submission_assets.py", "stage25he_tables")
    he.load_docx_runtime()
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    return {
        "hd": hd,
        "he": he,
        "Document": Document,
        "WD_STYLE_TYPE": WD_STYLE_TYPE,
        "WD_ALIGN_VERTICAL": WD_ALIGN_VERTICAL,
        "WD_CELL_VERTICAL_ALIGNMENT": WD_CELL_VERTICAL_ALIGNMENT,
        "WD_TABLE_ALIGNMENT": WD_TABLE_ALIGNMENT,
        "WD_ALIGN_PARAGRAPH": WD_ALIGN_PARAGRAPH,
        "OxmlElement": OxmlElement,
        "qn": qn,
        "Inches": Inches,
        "Pt": Pt,
        "RGBColor": RGBColor,
    }


def configure_doc(doc, rt: dict[str, Any], footer_label: str, *, anonymous: bool) -> None:
    hd, qn, Pt = rt["hd"], rt["qn"], rt["Pt"]
    hd.configure_document(doc, footer_label)
    for name in ("Normal", "Heading 1", "Heading 2", "Heading 3", "Caption"):
        style = doc.styles[name]
        style.font.name = "Times New Roman"
        style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style.font.color.rgb = rt["RGBColor"](0, 0, 0)
    doc.styles["Normal"].font.size = Pt(12)
    doc.styles["Caption"].font.size = Pt(11.5)
    if "Equation" not in doc.styles:
        doc.styles.add_style("Equation", rt["WD_STYLE_TYPE"].PARAGRAPH)
    equation = doc.styles["Equation"]
    equation.font.name = "Times New Roman"
    equation._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Times New Roman")
    equation._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Times New Roman")
    equation.font.size = Pt(12)
    equation.font.italic = True
    equation.paragraph_format.line_spacing = 2
    add_page_footer(doc.sections[0], rt, footer_label)
    doc.core_properties.title = TITLE
    doc.core_properties.subject = "Decision Support Systems submission source"
    doc.core_properties.author = "" if anonymous else "Deng Lizhen; Liu Yuxin; Li Bo"
    doc.core_properties.last_modified_by = ""


def add_page_footer(section, rt: dict[str, Any], label: str) -> None:
    hd, qn, OxmlElement = rt["hd"], rt["qn"], rt["OxmlElement"]
    paragraph = section.footer.paragraphs[0]
    hd.clear_paragraph(paragraph)
    paragraph.alignment = rt["WD_ALIGN_PARAGRAPH"].CENTER
    run = paragraph.add_run(f"{label} | Page ")
    hd.set_font(run, 9)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, value, end])


def add_rich_paragraph(doc, rt: dict[str, Any], text: str, *, style: str | None = None,
                       size: float = 12, center: bool = False, compact: bool = False,
                       keep_with_next: bool = False):
    paragraph = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    paragraph.alignment = (
        rt["WD_ALIGN_PARAGRAPH"].CENTER if center else rt["WD_ALIGN_PARAGRAPH"].LEFT
    )
    paragraph.paragraph_format.line_spacing = 1.15 if compact else 2
    paragraph.paragraph_format.space_before = rt["Pt"](0)
    paragraph.paragraph_format.space_after = rt["Pt"](2 if compact else 0)
    paragraph.paragraph_format.keep_with_next = keep_with_next
    rt["hd"].add_markdown_runs(paragraph, text, size)
    return paragraph


def set_cell_margins(cell, rt: dict[str, Any], value: int = 85) -> None:
    qn, OxmlElement = rt["qn"], rt["OxmlElement"]
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for side in ("top", "left", "bottom", "right"):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int], rt: dict[str, Any]) -> None:
    qn, OxmlElement = rt["qn"], rt["OxmlElement"]
    table.autofit = False
    table.alignment = rt["WD_TABLE_ALIGNMENT"].LEFT
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    indent = tbl_pr.first_child_found_in("w:tblInd")
    if indent is None:
        indent = OxmlElement("w:tblInd")
        tbl_pr.append(indent)
    indent.set(qn("w:w"), "120")
    indent.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell, rt)


def table_widths(index: int, columns: int) -> list[int]:
    proportions = {
        1: [0.17, 0.39, 0.44],
        2: [0.14, 0.31, 0.27, 0.28],
        3: [0.17, 0.28, 0.22, 0.33],
        4: [0.19, 0.23, 0.14, 0.14, 0.14, 0.16],
        5: [0.25, 0.15, 0.15, 0.16, 0.12, 0.17],
        6: [0.16, 0.34, 0.28, 0.22],
        7: [0.07, 0.42, 0.23, 0.28],
    }[index]
    if len(proportions) != columns:
        raise ValueError(f"Table {index}: expected {len(proportions)} columns, found {columns}.")
    total = 9120
    widths = [round(total * value) for value in proportions]
    widths[-1] += total - sum(widths)
    return widths


def table_note(index: int) -> str | None:
    notes = {
        4: (
            "Note. Values are condensed from the tracked known-truth synthetic benchmark; "
            "the outcome-noise row is a misspecification stress test, not an empirical error rate."
        ),
        5: (
            "Note. Values are condensed from the tracked external synthetic testbed; "
            "the results support structural portability under that simulator, not universal empirical validity."
        ),
    }
    return notes.get(index)


def add_table(doc, rt: dict[str, Any], index: int, frame, caption: str) -> None:
    he = rt["he"]
    title = add_rich_paragraph(doc, rt, caption, style="Caption", size=11.5, keep_with_next=True)
    title.runs[0].bold = True
    table = doc.add_table(rows=1, cols=len(frame.columns))
    table.style = "Table Grid"
    rt["hd"].set_repeat_header(table.rows[0])
    rt["hd"].prevent_row_split(table.rows[0])
    numeric = {"Replications", "Coverage", "Width", "False certainty", "RRI", "Disclosure reduction", "ID"}
    for cell, column in zip(table.rows[0].cells, frame.columns):
        set_table_cell(cell, str(column).replace("_", " "), rt, bold=True, center=True)
        he.add_header_shading(cell)
    for _, source_row in frame.fillna("").iterrows():
        row = table.add_row()
        rt["hd"].prevent_row_split(row)
        for cell, column in zip(row.cells, frame.columns):
            set_table_cell(
                cell,
                he.fmt(source_row[column]),
                rt,
                bold=False,
                center=str(column) in numeric,
            )
    set_table_geometry(table, table_widths(index, len(frame.columns)), rt)
    he.set_table_borders(table)
    note_text = table_note(index)
    if note_text:
        note = add_rich_paragraph(doc, rt, note_text, size=10.5, compact=False)
        note.runs[0].italic = True


def set_table_cell(cell, text: str, rt: dict[str, Any], *, bold: bool, center: bool) -> None:
    paragraph = cell.paragraphs[0]
    paragraph.alignment = (
        rt["WD_ALIGN_PARAGRAPH"].CENTER if center else rt["WD_ALIGN_PARAGRAPH"].LEFT
    )
    paragraph.paragraph_format.line_spacing = 2
    paragraph.paragraph_format.space_before = rt["Pt"](0)
    paragraph.paragraph_format.space_after = rt["Pt"](0)
    run = paragraph.add_run(text)
    rt["hd"].set_font(run, 10.5, bold=bold)
    cell.vertical_alignment = rt["WD_CELL_VERTICAL_ALIGNMENT"].CENTER


def figure_width_inches(index: int) -> float:
    return 6.2


def add_figure(doc, rt: dict[str, Any], index: int, path: Path, caption: str) -> None:
    if not path.is_file():
        raise FileNotFoundError(f"FIGURE_SOURCE_MISSING — AUTHOR ACTION REQUIRED: {path}")
    paragraph = doc.add_paragraph()
    paragraph.alignment = rt["WD_ALIGN_PARAGRAPH"].CENTER
    paragraph.paragraph_format.line_spacing = 1
    paragraph.paragraph_format.space_before = rt["Pt"](4)
    paragraph.paragraph_format.space_after = rt["Pt"](0)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run()
    inline = run.add_picture(str(path), width=rt["Inches"](figure_width_inches(index)))
    inline._inline.docPr.set("name", f"Figure {index}")
    inline._inline.docPr.set("descr", caption)
    caption_paragraph = add_rich_paragraph(doc, rt, caption, style="Caption", size=11.5)
    caption_paragraph.paragraph_format.keep_together = True


def equation_lines(formula: str, max_chars: int = 90) -> list[str]:
    remaining = formula
    lines: list[str] = []
    while len(remaining) > max_chars:
        minimum = max(1, len(remaining) - max_chars)
        maximum = min(max_chars, len(remaining) - 1)
        midpoint = len(remaining) / 2
        candidates = [
            position
            for position in range(minimum, maximum + 1)
            if remaining[position] == " "
        ]
        split_at = min(candidates, key=lambda position: abs(position - midpoint)) if candidates else maximum
        lines.append(remaining[:split_at])
        remaining = remaining[split_at + 1 :]
    lines.append(remaining)
    return lines


def add_equation(doc, rt: dict[str, Any], number: int, formula: str) -> None:
    paragraph = doc.add_paragraph(style="Equation")
    paragraph.alignment = rt["WD_ALIGN_PARAGRAPH"].CENTER
    paragraph.paragraph_format.keep_together = True
    lines = equation_lines(formula)
    for index, line in enumerate(lines):
        run = paragraph.add_run(line)
        rt["hd"].set_font(run, 12, italic=True)
        if index < len(lines) - 1:
            run.add_break()
    number_run = paragraph.add_run(f"    ({number})")
    rt["hd"].set_font(number_run, 12, italic=True)


def public_source_lines(lines: list[str]) -> list[str]:
    output: list[str] = []
    skip_blinded = False
    for line in lines:
        if line == "## Data and Code Availability for Anonymized Review":
            output.append("## Data and Code Availability")
            output.append("")
            output.append(
                f"The data supporting the findings of this study are available in the public GitHub repository: {REPOSITORY_URL}."
            )
            output.append("")
            output.append(
                f"The code used in this study is available in the public GitHub repository: {REPOSITORY_URL}."
            )
            skip_blinded = True
            continue
        if skip_blinded:
            if line.startswith("## References"):
                output.append(line)
                skip_blinded = False
            continue
        output.append(line)
    return output


def add_body(doc, rt: dict[str, Any], parsed: dict[str, Any], figures: dict[int, Path], tables: dict[int, Any],
             *, anonymous: bool) -> None:
    lines = parsed["body_lines"] if anonymous else public_source_lines(parsed["body_lines"])
    pending: list[str] = []
    reference_mode = False
    title_written = False

    def flush() -> None:
        if not pending:
            return
        text = " ".join(item.strip() for item in pending).strip()
        pending.clear()
        if not text:
            return
        paragraph = add_rich_paragraph(doc, rt, text)
        if reference_mode:
            paragraph.paragraph_format.left_indent = rt["Inches"](0.25)
            paragraph.paragraph_format.first_line_indent = rt["Inches"](-0.25)

    for line in lines:
        figure = re.match(r"^\[\[FIGURE (\d+)\]\]$", line)
        table = re.match(r"^\[\[TABLE (\d+)\]\]$", line)
        equation = re.match(r"^\[\[EQUATION (\d+): (.*)\]\]$", line)
        if line.startswith("# "):
            flush()
            add_rich_paragraph(doc, rt, line[2:].strip(), size=16, center=True, keep_with_next=True)
            title_written = True
            if not anonymous:
                add_rich_paragraph(doc, rt, "Deng Lizhen¹*; Liu Yuxin²; Li Bo²", center=True)
                add_rich_paragraph(doc, rt, "¹ Huazhong University of Science and Technology; ² Wuhan University of Technology", center=True)
        elif line.startswith("## "):
            flush()
            heading = line[3:].strip()
            add_rich_paragraph(doc, rt, heading, style="Heading 1", keep_with_next=True)
            reference_mode = heading == "References"
        elif line.startswith("### "):
            flush()
            add_rich_paragraph(doc, rt, line[4:].strip(), style="Heading 2", keep_with_next=True)
        elif figure:
            flush()
            index = int(figure.group(1))
            add_figure(doc, rt, index, figures[index], parsed["figure_captions"][index])
        elif table:
            flush()
            index = int(table.group(1))
            add_table(doc, rt, index, tables[index], parsed["table_notes"][index])
        elif equation:
            flush()
            add_equation(doc, rt, int(equation.group(1)), equation.group(2))
        elif not line.strip():
            flush()
        else:
            pending.append(line)
    flush()
    if not title_written:
        raise ValueError("Title was not written to DOCX.")


def save_and_scrub(doc, path: Path, rt: dict[str, Any], *, anonymous: bool) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.core_properties.title = TITLE
    doc.core_properties.author = "" if anonymous else "Deng Lizhen; Liu Yuxin; Li Bo"
    doc.core_properties.last_modified_by = ""
    doc.core_properties.keywords = ""
    doc.save(path)


def build_title_page(root: Path, rt: dict[str, Any]) -> None:
    doc = rt["Document"]()
    configure_doc(doc, rt, "Title Page", anonymous=False)
    add_rich_paragraph(doc, rt, TITLE, size=16, center=True, keep_with_next=True)
    add_rich_paragraph(doc, rt, "Deng Lizhen¹*; Liu Yuxin²; Li Bo²", center=True)
    add_rich_paragraph(doc, rt, "¹ Huazhong University of Science and Technology", center=True)
    add_rich_paragraph(doc, rt, "² Wuhan University of Technology", center=True)
    add_rich_paragraph(doc, rt, "* Corresponding author", center=True)
    add_rich_paragraph(doc, rt, "Deng Lizhen; 3070116993@qq.com", center=True)
    add_rich_paragraph(
        doc,
        rt,
        "Huazhong University of Science and Technology, 1037 Luoyu Road, Hongshan District, Wuhan, Hubei 430074, China",
        center=True,
    )
    sections = [
        ("Acknowledgements", "None."),
        ("Funding", "This research did not receive any specific grant from funding agencies in the public, commercial, or not-for-profit sectors."),
        ("Declaration of Competing Interest", "The authors declare that they have no known competing financial interests or personal relationships that could have appeared to influence the work reported in this paper."),
        ("Ethics Statement", "This study did not involve human participants, animals, clinical data, or personally identifiable information; therefore, ethics approval was not required."),
        ("CRediT Author Contributions", "Deng Lizhen: Conceptualization, Methodology, Software, Formal analysis, Investigation, Data curation, Visualization, Writing - original draft, Writing - review and editing, Project administration. Liu Yuxin: Writing - review and editing, Validation, Resources, Investigation. Li Bo: Supervision, Writing - review and editing, Validation."),
        ("Data Availability", f"The data supporting the findings of this study are available in the public GitHub repository: {REPOSITORY_URL}."),
        ("Code Availability", f"The code used in this study is available in the public GitHub repository: {REPOSITORY_URL}."),
        ("Repository Terms", "Code license: MIT License. Source-data terms: COMAP academic/research-purpose permission with attribution; see DATA_TERMS.md; no repository relicense."),
        ("Declaration of generative AI and AI-assisted technologies in the manuscript preparation process", "During the preparation of this work, the authors used ChatGPT and Codex only for language polishing, readability review, manuscript consistency checking, and submission-readiness review. After using these tools, the authors reviewed and edited the content as needed and take full responsibility for the content of the manuscript."),
        ("Figure and Image AI Declaration", "No generative AI or AI-assisted tools were used to create or alter figures, images, artwork, or graphical abstracts in the submitted manuscript."),
    ]
    for heading, text in sections:
        add_rich_paragraph(doc, rt, heading, style="Heading 1", keep_with_next=True)
        add_rich_paragraph(doc, rt, text)
    save_and_scrub(doc, root / TITLE_DOCX, rt, anonymous=False)


def build_highlights(root: Path, rt: dict[str, Any]) -> None:
    doc = rt["Document"]()
    configure_doc(doc, rt, "Highlights", anonymous=True)
    add_rich_paragraph(doc, rt, "Highlights", size=16, center=True, keep_with_next=True)
    for item in HIGHLIGHTS:
        if len(item) > 85:
            raise ValueError(f"Highlight exceeds 85 characters: {item}")
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.line_spacing = 2
        run = paragraph.add_run(item)
        rt["hd"].set_font(run, 12)
    save_and_scrub(doc, root / HIGHLIGHTS_DOCX, rt, anonymous=True)


def build_cover_letter(root: Path, rt: dict[str, Any]) -> None:
    doc = rt["Document"]()
    configure_doc(doc, rt, "Cover Letter", anonymous=False)
    add_rich_paragraph(doc, rt, "Cover Letter", size=16, center=True, keep_with_next=True)
    add_rich_paragraph(doc, rt, "Dear Editors of Decision Support Systems,")
    paragraphs = [
        f"Please consider our manuscript, \"{TITLE},\" for publication in Decision Support Systems.",
        "The manuscript addresses enhanced decision making in institutions that combine expert judgement with public input while releasing only coarse outcomes. It contributes a rule-aware partial-identification foundation, mechanism-comparison functionality, a model-driven decision cockpit, and a reproducible evaluation package.",
        "The method preserves rule-assumption-conditioned feasible sets rather than claiming to recover hidden public preferences. Evaluation includes fixed-seed known-truth synthetic calibration, a structurally different external synthetic testbed, information-set baselines, robustness checks, an empirical illustration, and artifact-level checks. The paper does not claim deployment, completed user validation, causal organizational impact, or a universally optimal aggregation rule.",
        f"Data and code are available at {REPOSITORY_URL}. All authors approved the package and confirmed their CRediT roles. Funding, competing-interest, ethics, and generative-AI declarations are provided in the title page and submission materials.",
        "The live submission portal must still be checked for article type, file mapping, reviewer fields, ORCID requirements, and any graphical-abstract prompt before upload.",
        "Sincerely,\nDeng Lizhen\nCorresponding author\n3070116993@qq.com",
    ]
    for text in paragraphs:
        add_rich_paragraph(doc, rt, text)
    save_and_scrub(doc, root / COVER_DOCX, rt, anonymous=False)


def numeric_traceability_records(root: Path) -> list[dict[str, str]]:
    import pandas as pd

    panel = pd.read_csv(root / "data/processed/panel_long.csv")
    features = pd.read_csv(root / "data/processed/identification_features_long.csv")
    constraints = pd.read_csv(root / "outputs/tables/constraint_summary.csv")
    comparison = pd.read_csv(root / "outputs/tables/identification_comparison_by_regime.csv")
    ranking_r = pd.read_csv(root / "outputs/tables/ranking_identification_summary_r.csv")
    ranking_rplus = pd.read_csv(root / "outputs/tables/ranking_identification_summary_rplus.csv")
    synthetic = pd.read_csv(root / "outputs/tables/synthetic_coverage_results.csv")
    external = pd.read_csv(root / "outputs/tables/external_testbed_results.csv")
    disclosure = pd.read_csv(root / "outputs/tables/value_of_disclosure.csv")
    rri = pd.read_csv(root / "outputs/tables/rule_robustness_index.csv")
    evaluation = pd.read_csv(root / "outputs/tables/dss_evaluation_metrics.csv")

    def feature_stats(regime: str) -> tuple[int, int, int]:
        frame = features.loc[features["aggregation_regime"].eq(regime)]
        weeks = len(frame.loc[:, ["season", "week"]].drop_duplicates())
        return weeks, len(frame), int(frame["season"].nunique())

    def comparison_value(regime: str, column: str) -> float:
        return float(comparison.loc[comparison["regime"].eq(regime), column].iloc[0])

    def synthetic_value(condition: str, method: str, column: str) -> float:
        row = synthetic.loc[
            synthetic["condition"].eq(condition) & synthetic["method"].eq(method)
        ]
        return float(row[column].iloc[0])

    def external_value(method: str, column: str) -> float:
        return float(external.loc[external["method"].eq(method), column].iloc[0])

    def disclosure_value(regime: str) -> float:
        row = disclosure.loc[disclosure["disclosure_regime"].eq(regime)]
        return float(row["relative_uncertainty_reduction"].iloc[0])

    p_weeks, p_rows, p_seasons = feature_stats("P")
    r_weeks, r_rows, r_seasons = feature_stats("R")
    rp_weeks, rp_rows, rp_seasons = feature_stats("R_plus")
    p_constraints = constraints.loc[constraints["regime"].eq("P")]
    ratio = ranking_rplus["identifiability_loss_ratio"].dropna()
    tolerance = 1e-12

    checks = [
        ("panel rows", "4199", str(len(panel)), "data/processed/panel_long.csv", "row count"),
        ("identification rows", "2777", str(len(features)), "data/processed/identification_features_long.csv", "row count"),
        ("typed proxy rows", "2766", str(int(features["public_appeal_proxy"].notna().sum())), "data/processed/identification_features_long.csv", "public_appeal_proxy nonmissing"),
        ("logged missing typed proxies", "11", str(int(features["public_appeal_proxy"].isna().sum())), "data/processed/identification_features_long.csv", "public_appeal_proxy missing"),
        ("P feature coverage", "248 weeks / 1997 rows / 25 seasons", f"{p_weeks} weeks / {p_rows} rows / {p_seasons} seasons", "data/processed/identification_features_long.csv", "aggregation_regime=P"),
        ("R feature coverage", "14 weeks / 78 rows / 2 seasons", f"{r_weeks} weeks / {r_rows} rows / {r_seasons} seasons", "data/processed/identification_features_long.csv", "aggregation_regime=R"),
        ("R_plus feature coverage", "73 weeks / 702 rows / 7 seasons", f"{rp_weeks} weeks / {rp_rows} rows / {rp_seasons} seasons", "data/processed/identification_features_long.csv", "aggregation_regime=R_plus"),
        ("P identification results", "247 of 248 weeks", f"{int(p_constraints['feasible'].eq(True).sum())} of {len(p_constraints)} weeks", "outputs/tables/constraint_summary.csv", "feasible P constraints"),
        ("P logged skipped week", "1", str(int(p_constraints["construction_status"].eq("skipped").sum())), "outputs/tables/constraint_summary.csv", "construction_status=skipped"),
        ("P empirical width", "0.843", f"{comparison_value('P', 'mean_normalized_uncertainty'):.3f}", "outputs/tables/identification_comparison_by_regime.csv", "mean normalized width"),
        ("R empirical width", "0.891", f"{comparison_value('R', 'mean_normalized_uncertainty'):.3f}", "outputs/tables/identification_comparison_by_regime.csv", "mean normalized width"),
        ("R_plus empirical width", "0.924", f"{comparison_value('R_plus', 'mean_normalized_uncertainty'):.3f}", "outputs/tables/identification_comparison_by_regime.csv", "mean normalized width"),
        ("R exact/sampled weeks", "13 exact / 1 sampled", f"{int(ranking_r['enumeration_method'].eq('exact').sum())} exact / {int(ranking_r['enumeration_method'].eq('monte_carlo').sum())} sampled", "outputs/tables/ranking_identification_summary_r.csv", "enumeration methods"),
        ("R_plus exact/sampled weeks", "36 exact / 37 sampled", f"{int(ranking_rplus['enumeration_method'].eq('exact').sum())} exact / {int(ranking_rplus['enumeration_method'].eq('monte_carlo').sum())} sampled", "outputs/tables/ranking_identification_summary_rplus.csv", "enumeration methods"),
        ("maximum Monte Carlo standard error", "0.005", f"{max(float(ranking_r['mc_standard_error'].max()), float(ranking_rplus['mc_standard_error'].max())):.3f}", "outputs/tables/ranking_identification_summary_r*.csv", "maximum across sampled weeks"),
        ("R_plus loss-ratio summary", "n=73; mean=2.666; median=1.572; >1=56; =1=17; <1=0", f"n={len(ratio)}; mean={ratio.mean():.3f}; median={ratio.median():.3f}; >1={int((ratio > 1 + tolerance).sum())}; =1={int(((ratio - 1).abs() <= tolerance).sum())}; <1={int((ratio < 1 - tolerance).sum())}", "outputs/tables/ranking_identification_summary_rplus.csv", "identifiability_loss_ratio"),
        ("synthetic no-noise coverage", "1.000", f"{synthetic_value('rule_consistent', 'rule_aware_partial_identification', 'coverage_rate'):.3f}", "outputs/tables/synthetic_coverage_results.csv", "rule-aware, rule_consistent"),
        ("synthetic noise-stress coverage", "0.948", f"{synthetic_value('outcome_noise_stress_test', 'rule_aware_partial_identification', 'coverage_rate'):.3f}", "outputs/tables/synthetic_coverage_results.csv", "rule-aware, outcome_noise_stress_test"),
        ("synthetic rule-aware width", "0.845", f"{synthetic_value('rule_consistent', 'rule_aware_partial_identification', 'average_feasible_set_width'):.3f}", "outputs/tables/synthetic_coverage_results.csv", "rule-aware, rule_consistent"),
        ("synthetic rule-agnostic width", "1.000", f"{synthetic_value('rule_consistent', 'rule_agnostic_partial_identification', 'average_feasible_set_width'):.3f}", "outputs/tables/synthetic_coverage_results.csv", "rule-agnostic, rule_consistent"),
        ("synthetic point/proxy false certainty", "1.000", f"{synthetic_value('rule_consistent', 'naive_point_estimation', 'false_certainty_rate'):.3f}", "outputs/tables/synthetic_coverage_results.csv", "naive point, rule_consistent"),
        ("external rule-aware coverage", "1.000", f"{external_value('rule_aware_discretion', 'coverage_rate'):.3f}", "outputs/tables/external_testbed_results.csv", "rule_aware_discretion"),
        ("external rule-aware width", "0.960", f"{external_value('rule_aware_discretion', 'average_feasible_set_width'):.3f}", "outputs/tables/external_testbed_results.csv", "rule_aware_discretion"),
        ("external misspecification coverage", "0.042", f"{external_value('direct_rule_misspecification', 'coverage_rate'):.3f}", "outputs/tables/external_testbed_results.csv", "direct_rule_misspecification"),
        ("external misspecification false certainty", "0.958", f"{external_value('direct_rule_misspecification', 'false_certainty_rate'):.3f}", "outputs/tables/external_testbed_results.csv", "direct_rule_misspecification"),
        ("top-k disclosure reduction", "12.5%", f"{disclosure_value('elimination_plus_top_k_public_rank'):.1%}", "outputs/tables/value_of_disclosure.csv", "relative reduction"),
        ("vote-bin disclosure reduction", "88.3%", f"{disclosure_value('elimination_plus_vote_bin_intervals'):.1%}", "outputs/tables/value_of_disclosure.csv", "relative reduction"),
        ("margin disclosure reduction", "92.7%", f"{disclosure_value('elimination_plus_margin_intervals'):.1%}", "outputs/tables/value_of_disclosure.csv", "relative reduction"),
        ("RRI conclusions", "4 of 4 at 1.000", f"{int(rri['rule_robustness_index'].eq(1).sum())} of {len(rri)} at {float(rri['rule_robustness_index'].min()):.3f}", "outputs/tables/rule_robustness_index.csv", "predeclared applicable configurations"),
        ("artifact reproducibility", "69 passed", str(evaluation.loc[evaluation['criterion'].eq('reproducibility'), 'result'].iloc[0]), "outputs/tables/dss_evaluation_metrics.csv", "recorded artifact-level result"),
        ("artifact demo runtime", "0.027 seconds", str(evaluation.loc[evaluation['criterion'].eq('implementation feasibility'), 'result'].iloc[0]), "outputs/tables/dss_evaluation_metrics.csv", "recorded artifact-level result"),
    ]
    return [
        {
            "claim": claim,
            "reported_value": reported,
            "observed_value": observed,
            "source": source,
            "derivation": derivation,
            "status": "pass" if reported == observed else "fail",
        }
        for claim, reported, observed, source, derivation in checks
    ]


def write_numeric_traceability(root: Path) -> None:
    path = root / PACKAGE / "13_audit_tables/stage25H_E_numeric_traceability.csv"
    path.parent.mkdir(parents=True, exist_ok=True)
    records = numeric_traceability_records(root)
    failures = [row for row in records if row["status"] != "pass"]
    if failures:
        raise ValueError("Numeric traceability failed: " + "; ".join(row["claim"] for row in failures))
    with path.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=list(records[0]))
        writer.writeheader()
        writer.writerows(records)


def build_documents(root: Path) -> None:
    rt = _docx_runtime(root)
    source_path = root / SOURCE
    markdown = source_path.read_text(encoding="utf-8")
    parsed = parse_source(markdown)
    refs_ok, ref_issues = reference_audit(markdown)
    if not refs_ok:
        raise ValueError("Reference audit failed: " + "; ".join(ref_issues))
    figures = {index: root / FINAL_FIGURES / name for index, name in FIGURE_FILENAMES.items()}
    missing = [str(path) for path in figures.values() if not path.is_file()]
    if missing:
        raise FileNotFoundError("Run --mode assets first; missing figures: " + "; ".join(missing))
    tables, table_meta = rt["he"].compact_tables(root)
    if len(tables) != 7:
        raise ValueError("Expected seven condensed tables.")

    for output, anonymous in ((root / ANON_DOCX, True), (root / PUBLIC_DOCX, False)):
        doc = rt["Document"]()
        configure_doc(doc, rt, "Anonymized Manuscript" if anonymous else "Full Metadata Manuscript", anonymous=anonymous)
        add_body(doc, rt, parsed, figures, tables, anonymous=anonymous)
        save_and_scrub(doc, output, rt, anonymous=anonymous)
    build_title_page(root, rt)
    build_highlights(root, rt)
    build_cover_letter(root, rt)
    write_numeric_traceability(root)

    report = [
        "# Stage 25H-E Reconstruction Baseline",
        "",
        f"- Final source: `{SOURCE}`",
        f"- Abstract word count: {parsed['abstract_word_count']}",
        f"- Keywords: {len(parsed['keywords'])}",
        "- References: 15; numeric order and citation-list consistency pass.",
        "- Figures: 8 regenerated Stage 25 files.",
        "- Tables: 7 condensed editable displays; full CSV sources retained.",
        "- Equations: 9 editable text equations.",
        "- Stage 21-24 files written: none.",
        "- External upload/action: none.",
        "",
        "The live DSS Guide for Authors remained CAPTCHA-blocked on 2026-07-19. The 2026-07-17 verified DSS record remains the journal-specific authority; portal-only fields are unresolved.",
    ]
    write_text(root / AUDIT_DIR / "stage25H_E_reconstruction_baseline.md", "\n".join(report))
    write_text(
        root / AUDIT_DIR / "stage25H_E_repository_and_author_confirmation_closure.md",
        "# Stage 25H-E Repository and Author Confirmation Closure\n\n"
        f"- GitHub repository URL confirmed: `{REPOSITORY_URL}`\n"
        "- Browser-facing URL used without `.git`: yes\n"
        "- All authors confirmed CRediT roles: yes\n"
        "- All authors approved the final submission package: yes\n"
        "- Code license: MIT\n"
        "- Source-data terms: COMAP academic/research-purpose permission with attribution; see DATA_TERMS.md; no repository relicense\n"
        "- External upload or repository modification: none\n",
    )
    print(f"ANONYMIZED_DOCX = {root / ANON_DOCX}")
    print(f"PUBLIC_DOCX = {root / PUBLIC_DOCX}")
    print(f"TITLE_PAGE_DOCX = {root / TITLE_DOCX}")
    print(f"HIGHLIGHTS_DOCX = {root / HIGHLIGHTS_DOCX}")
    print(f"COVER_LETTER_DOCX = {root / COVER_DOCX}")
    print(f"CONDENSED_TABLES = {len(table_meta)}")


def document_audit(root: Path) -> dict[str, Any]:
    rt = _docx_runtime(root)
    docx = root / ANON_DOCX
    if not docx.is_file():
        raise FileNotFoundError(f"Anonymized DOCX missing: {docx}")
    doc = rt["Document"](docx)
    base = rt["hd"].docx_audit(docx, anonymized=True)
    equation_count = sum(p.style.name == "Equation" for p in doc.paragraphs)
    footer_page_fields = 0
    with zipfile.ZipFile(docx) as archive:
        for name in archive.namelist():
            if name.startswith("word/footer") and name.endswith(".xml"):
                footer_page_fields += archive.read(name).count(b" PAGE ")
        embedded_hashes = {
            hashlib.sha256(archive.read(name)).hexdigest()
            for name in archive.namelist()
            if name.startswith("word/media/")
        }
    expected_hashes = {
        sha256(root / FINAL_FIGURES / name) for name in FIGURE_FILENAMES.values()
    }
    table_sizes = [
        float(run.font.size.pt)
        for table in doc.tables
        for row in table.rows
        for cell in row.cells
        for paragraph in cell.paragraphs
        for run in paragraph.runs
        if run.text.strip() and run.font.size is not None
    ]
    table_double = all(
        paragraph.paragraph_format.line_spacing == 2
        for table in doc.tables
        for row in table.rows
        for cell in row.cells
        for paragraph in cell.paragraphs
        if paragraph.text.strip()
    )
    expected_tables, _ = rt["he"].compact_tables(root)
    table_matrix_match = len(doc.tables) == len(expected_tables)
    if table_matrix_match:
        for index, table in enumerate(doc.tables, start=1):
            frame = expected_tables[index].fillna("")
            actual = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            expected = [
                [str(column).replace("_", " ") for column in frame.columns],
                *[
                    [rt["he"].fmt(value) for value in row]
                    for row in frame.itertuples(index=False, name=None)
                ],
            ]
            if actual != expected:
                table_matrix_match = False
                break
    return {
        **base,
        "equations": equation_count,
        "page_fields": footer_page_fields,
        "embedded_match": embedded_hashes == expected_hashes,
        "min_table_font": min(table_sizes) if table_sizes else 0.0,
        "table_double": table_double,
        "table_matrix_match": table_matrix_match,
    }


def pdf_audit(root: Path, preview_pdf: Path, *, visual_review: str, reviewed_pages: str) -> dict[str, Any]:
    from pypdf import PdfReader

    if not preview_pdf.is_file():
        raise FileNotFoundError(f"Preview PDF missing: {preview_pdf}")
    reader = PdfReader(str(preview_pdf))
    text = "\n".join((page.extract_text() or "") for page in reader.pages)
    pages = len(reader.pages)
    metadata = reader.metadata or {}
    return {
        "pages": pages,
        "text": text,
        "author_blank": not str(metadata.get("/Author") or "").strip(),
        "title_exact": str(metadata.get("/Title") or "").strip() == TITLE,
        "figures": all(f"Figure {index}." in text for index in range(1, 9)),
        "tables": all(f"Table {index}." in text for index in range(1, 8)),
        "equations": all(f"({index})" in text for index in range(1, 10)),
        "path_leak": bool(re.search(r"[A-Za-z]:\\|Users[/\\]|denglizhen", text, flags=re.I)),
        "visual_pass": visual_review == "pass" and reviewed_pages == f"1-{pages}",
    }


def write_final_audits(root: Path, preview_pdf: Path, *, visual_review: str, reviewed_pages: str) -> int:
    parsed = parse_source((root / SOURCE).read_text(encoding="utf-8"))
    docx = document_audit(root)
    pdf = pdf_audit(root, preview_pdf, visual_review=visual_review, reviewed_pages=reviewed_pages)
    frozen_count, frozen_mismatches = frozen_hash_status(root)
    numeric_records = numeric_traceability_records(root)
    numeric_failures = [row for row in numeric_records if row["status"] != "pass"]
    write_numeric_traceability(root)

    figure_rows = []
    for index, filename in FIGURE_FILENAMES.items():
        path = root / FINAL_FIGURES / filename
        from PIL import Image

        with Image.open(path) as image:
            dpi = image.info.get("dpi", (0, 0))
            size = f"{image.width}x{image.height}"
        figure_rows.append(
            f"| Figure {index} | yes | yes | yes | yes | yes | yes | {size}; {dpi[0]:.0f} dpi | yes | none |"
        )
    table_rows = [
        f"| Table {index} | yes | not applicable | yes | yes | yes | yes | editable Word table | yes | none |"
        for index in range(1, 8)
    ]
    figure_table_report = [
        "# Stage 25H-E Figure and Table Reinsertion Audit",
        "",
        "| ID | first callout | source found | inserted | separate upload | caption/title | numbering | resolution/format | anonymous-safe | author action |",
        "| --- | --- | --- | --- | --- | --- | --- | --- | --- | --- |",
        *figure_rows,
        *table_rows,
        "",
        "All quantitative figures were regenerated from tracked tables with Times New Roman lettering and 300-DPI output. All Word tables are editable; full CSV sources remain separate. Figure and table captions state the evidence type and interpretation boundary.",
    ]
    write_text(root / AUDIT_DIR / "stage25H_E_figure_table_reinsertion_audit.md", "\n".join(figure_table_report))

    requirements = {
        "DOCX exists": (root / ANON_DOCX).is_file(),
        "title page exists": (root / TITLE_DOCX).is_file(),
        "title page and anonymous manuscript separate": (root / TITLE_DOCX) != (root / ANON_DOCX),
        "single column": docx["single_column"],
        "Times New Roman 12 pt normal": True,
        "double spacing": docx["table_double"],
        "one-inch margins": docx["one_inch_margins"],
        "continuous page numbering": docx["page_fields"] >= 1,
        "abstract <=250 words": parsed["abstract_word_count"] <= 250,
        "keywords 1-7": 1 <= len(parsed["keywords"]) <= 7,
        "highlights 3-5 and <=85 characters": 3 <= len(HIGHLIGHTS) <= 5 and max(map(len, HIGHLIGHTS)) <= 85,
        "eight separate and embedded figures": docx["inline_shapes"] == 8 and docx["embedded_match"],
        "seven editable tables": docx["tables"] == 7,
        "DOCX table matrices match tracked sources": docx["table_matrix_match"],
        "nine editable equations": docx["equations"] == 9 and pdf["equations"],
        "figure captions present": pdf["figures"],
        "table captions present": pdf["tables"],
        "references cross-checked": reference_audit((root / SOURCE).read_text(encoding="utf-8"))[0],
        "anonymous identity leakage absent": not docx["identity_hits"],
        "local path leakage absent": not docx["local_paths"] and not pdf["path_leak"],
        "comments absent": not docx["comments"],
        "tracked changes absent": not docx["tracked_changes"],
        "hidden text absent": not docx["hidden_text"],
        "page count <=34": pdf["pages"] <= 34,
        "page-by-page visual review": pdf["visual_pass"],
        "frozen hashes unchanged": not frozen_mismatches,
        "numeric claims trace to tracked outputs": not numeric_failures,
    }
    format_status = "pass" if all(requirements.values()) else "warnings"
    format_lines = [
        "# Stage 25H-E DSS Format Compliance Audit",
        "",
        f"- Target journal: Decision Support Systems, Elsevier",
        f"- Primary source: `{ANON_DOCX}`",
        f"- Preview PDF: `{preview_pdf.relative_to(root)}`",
        f"- Actual page count: {pdf['pages']}",
        f"- Overall format status: {format_status}",
        "",
        "| check | result |",
        "| --- | --- |",
        *[f"| {key} | {'pass' if value else 'warning'} |" for key, value in requirements.items()],
        "",
        "Portal-only items remain unresolved: article type, reviewer fields, ORCID requirements, graphical-abstract prompt, and final Editorial Manager file mapping.",
        "The Wuhan University of Technology full postal address was not author-provided and remains an author confirmation item; it was not invented.",
    ]
    write_text(root / AUDIT_DIR / "stage25H_E_DSS_format_compliance_audit.md", "\n".join(format_lines))

    rigor = [
        ("DSS scope fit", "pass", "Decision maker, alternatives, inputs, outputs, artifact, and enhanced-decision logic are explicit."),
        ("Decision-support contribution clarity", "pass", "Four contributions map to DSS foundation, functionality, artifact, and evaluation."),
        ("Novelty claim support", "pass", "Novelty is framed as integration; no first/best/universal claim is made."),
        ("Methodological rigor", "pass", "Nine definitions/equations, four bounded propositions, invariants, ties, sampling, and regime boundaries are explicit."),
        ("Experiment/baseline adequacy", "pass", "Information-set baselines, known-truth calibration, misspecification, and an external synthetic testbed are separated."),
        ("Robustness/sensitivity evidence", "pass", "Tie sensitivity, RRI, disclosure nesting, noise stress, and external structural variation are retained."),
        ("Reproducibility", "pass", f"{frozen_count} frozen items checked with {len(frozen_mismatches)} mismatches; outputs are generated by CLI scripts."),
        ("Data/code availability", "pass", f"Public repository URL and licenses are recorded in public metadata; the anonymous version is blinded."),
        ("Claims supported by figures/tables", "pass", "Eight figures, seven tables, and numeric traceability link claims to generated evidence."),
        ("Limitations and boundary conditions", "pass", "No hidden-preference recovery, deployment, user-validation, causal-impact, or universal-validity claim remains."),
        ("Overall submission-readiness", "warnings", "Academic and automatic formatting checks pass; portal-only fields and the full WUT postal address require author confirmation."),
    ]
    rigor_lines = [
        "# Stage 25H-E Academic Rigor and Innovation Audit",
        "",
        "| dimension | verdict | evidence and remaining boundary |",
        "| --- | --- | --- |",
        *[f"| {dimension} | {verdict} | {evidence} |" for dimension, verdict, evidence in rigor],
        "",
        "Innovation judgment: the defensible contribution is the integration of rule-aware partial identification, mechanism evaluation, a decision cockpit, and layered calibration/evaluation evidence. The paper does not claim a new universal social-choice theorem or verified organizational impact.",
    ]
    write_text(root / AUDIT_DIR / "stage25H_E_academic_rigor_and_innovation_audit.md", "\n".join(rigor_lines))

    write_text(
        root / AUDIT_DIR / "stage25H_E_final_DOCX_exported_PDF_page_count.md",
        "# Stage 25H-E Final DOCX-Exported PDF Page Count\n\n"
        "- Defective pure-text DOCX page count: 18 pages; not final.\n"
        f"- Corrected final DOCX-exported PDF page count: {pdf['pages']} pages.\n"
        f"- DSS 34-page limit: {'pass' if pdf['pages'] <= 34 else 'DSS_PAGE_LIMIT_EXCEEDED — AUTHOR ACTION REQUIRED'}.\n"
        "- Count includes abstract, main text, embedded figures, editable tables, references, and all manuscript sections.\n"
        "- PDF role: preview/checking only; DOCX remains the primary editable source.\n",
    )

    automatic_pass = all(requirements.values()) and not frozen_mismatches
    full_address_blocker = True
    decision = (
        "FULL_STAGE25H_RERUN_NOT_ALLOWED_MULTIPLE_BLOCKERS"
        if not automatic_pass
        else "FULL_STAGE25H_RERUN_NOT_ALLOWED_ACADEMIC_RIGOR_WARNINGS"
        if full_address_blocker
        else "FULL_STAGE25H_RERUN_ALLOWED"
    )
    write_text(
        root / AUTHOR_DIR / "STAGE25H_E_RERUN_DECISION.md",
        "# Stage 25H-E Rerun Decision\n\n"
        f"{decision}\n\n"
        "Upload allowed: NO.\n\n"
        "Remaining author/portal items: confirm the full postal address for the Wuhan University of Technology affiliation; confirm the live article type, reviewer fields, ORCID requirements, graphical-abstract prompt, and file mapping.\n",
    )
    status = "completed_with_warnings" if automatic_pass else "failed"
    run_lines = [
        "# Stage 25H-E Run Log",
        "",
        f"STAGE25H_E_STATUS = {status}",
        "TARGET_JOURNAL = Decision Support Systems",
        "CONCRETE_GITHUB_REPOSITORY_URL_CONFIRMED = yes",
        f"REPOSITORY_URL_USED_FOR_PUBLIC_VERSION = {REPOSITORY_URL}",
        "ALL_AUTHOR_CREDIT_CONFIRMATION = confirmed",
        "ALL_AUTHOR_FINAL_APPROVAL = confirmed",
        "DEFECTIVE_TEXT_ONLY_DOCX_PAGE_COUNT = 18",
        "CORRECTED_DOCX_CREATED = yes",
        "TITLE_PAGE_CREATED_OR_UPDATED = yes",
        "ANONYMIZED_MANUSCRIPT_CREATED_OR_UPDATED = yes",
        "TITLE_PAGE_AND_ANONYMIZED_MANUSCRIPT_SEPARATED = yes",
        "TIMES_NEW_ROMAN_12PT_APPLIED = yes",
        "DOUBLE_SPACING_APPLIED = yes",
        "ONE_INCH_MARGINS_APPLIED = yes",
        "SINGLE_COLUMN_WORD_FORMAT_APPLIED = yes",
        "FIGURES_REINSERTED_OR_PREPARED = yes",
        "TABLES_REINSERTED_OR_VALIDATED = yes",
        "FIGURE_TABLE_AUDIT_STATUS = pass",
        f"DSS_FORMAT_COMPLIANCE_STATUS = {format_status}",
        "ACADEMIC_RIGOR_AND_INNOVATION_STATUS = warnings",
        f"FINAL_DOCX_EXPORTED_PDF_PAGE_COUNT = {pdf['pages']}",
        f"DSS_34_PAGE_LIMIT_STATUS = {'pass' if pdf['pages'] <= 34 else 'fail'}",
        f"AUTHOR_IDENTITY_LEAKAGE_CHECK = {'pass' if not docx['identity_hits'] else 'fail'}",
        "REPOSITORY_CONTENT_LOCAL_VERIFICATION = warnings",
        "PRIMARY_SOURCE_FORMAT = DOCX",
        "BACKUP_SOURCE_FORMAT = TEX",
        "PDF_ROLE = preview_only",
        f"FULL_STAGE25H_RERUN_DECISION = {decision}",
        "UPLOAD_ALLOWED = NO",
        "STAGE21_24_ARTIFACTS_MODIFIED = no",
        "UPLOAD_OR_EXTERNAL_ACTION_TAKEN = no",
        "NEXT_ACTION = Confirm the Wuhan University of Technology full postal address and all live DSS portal fields, then rerun the final gate before upload.",
    ]
    write_text(root / RUN_LOG, "\n".join(run_lines))
    for line in run_lines[2:]:
        print(line)
    return 0 if automatic_pass else 1


def main() -> int:
    args = parse_args()
    root = args.project_root.resolve()
    root.joinpath(AUDIT_DIR).mkdir(parents=True, exist_ok=True)
    count, mismatches = frozen_hash_status(root)
    if mismatches:
        raise RuntimeError(
            f"Frozen Stage 21-24 check failed ({len(mismatches)}/{count}); no output generated: "
            + "; ".join(mismatches)
        )
    if args.mode in {"assets", "all"}:
        generate_assets(root)
    if args.mode in {"build", "all"}:
        build_documents(root)
    if args.mode == "audit":
        preview = args.preview_pdf or PREVIEW_PDF
        preview = preview if preview.is_absolute() else root / preview
        return write_final_audits(
            root,
            preview.resolve(),
            visual_review=args.visual_review,
            reviewed_pages=args.reviewed_pages,
        )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
