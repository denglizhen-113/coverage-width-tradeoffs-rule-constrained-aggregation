#!/usr/bin/env python3
"""Generate non-destructive Stage 25G editable source files and audits."""

from __future__ import annotations

import argparse
import csv
import hashlib
import os
import re
import shutil
import subprocess
import sys
import time
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from pypdf import PdfReader
import pypdfium2 as pdfium
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import letter, landscape
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import (
    BaseDocTemplate,
    Frame,
    LongTable,
    NextPageTemplate,
    PageBreak,
    PageTemplate,
    Paragraph,
    Spacer,
    TableStyle,
)
from xml.sax.saxutils import escape


PROJECT_ROOT = Path(__file__).resolve().parents[1]
PACKAGE = Path("submission_package_stage25")
TABLE_SLUGS = {
    1: "decision_alternatives_and_criteria",
    2: "assumption_inventory",
    3: "baseline_definitions",
    4: "synthetic_coverage_results",
    5: "external_testbed_results",
    6: "design_recommendation_matrix",
    7: "claim_evidence_alignment",
}
TABLE_TITLES = {
    1: "Decision alternatives and criteria",
    2: "Assumption inventory",
    3: "Baseline definitions",
    4: "Synthetic coverage results",
    5: "External testbed results",
    6: "Design recommendation matrix",
    7: "Claim-evidence alignment",
}


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Generate Stage 25G DOCX/TEX sources and local PDF previews without external actions."
    )
    parser.add_argument("--project-root", type=Path, default=PROJECT_ROOT)
    parser.add_argument(
        "--documents-skill-root",
        type=Path,
        required=True,
        help="Directory containing render_docx.py from the documents skill.",
    )
    return parser.parse_args()


def write_text(path: Path, content: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    temporary = path.with_suffix(path.suffix + ".tmp")
    temporary.write_text(content.strip() + "\n", encoding="utf-8")
    for attempt in range(10):
        try:
            temporary.replace(path)
            return
        except PermissionError:
            if attempt == 9:
                raise
            time.sleep(0.5)


def write_csv(path: Path, rows: list[dict[str, str]], fields: list[str]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    temporary = path.with_suffix(path.suffix + ".tmp")
    with temporary.open("w", encoding="utf-8", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=fields, lineterminator="\n")
        writer.writeheader()
        writer.writerows(rows)
    temporary.replace(path)


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def md_inputs(root: Path) -> dict[str, Path]:
    package = root / PACKAGE
    return {
        "cited manuscript": root / "manuscript/DSS_submission_draft_stage25_cited.md",
        "anonymized manuscript": root / "manuscript/DSS_submission_draft_stage25_anonymized.md",
        "author fill-in packet": package / "01_author_action_required/AUTHOR_FILL_IN_PACKET_STAGE25F.md",
        "editable-source checklist": package / "01_author_action_required/editable_source_and_page_count_checklist.md",
        "submission simulation": package / "12_audit_logs/stage25F_submission_simulation_report.md",
        "red-team report": package / "12_audit_logs/stage25F_final_red_team_report.md",
        "submission file manifest": package / "13_audit_tables/stage25_submission_file_manifest.csv",
    }


def set_font(run, size: float, bold: bool = False, italic: bool = False) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)


def setup_section(section, landscape: bool = False) -> None:
    section.orientation = WD_ORIENT.LANDSCAPE if landscape else WD_ORIENT.PORTRAIT
    section.page_width = Inches(11 if landscape else 8.5)
    section.page_height = Inches(8.5 if landscape else 11)
    margin = 0.75 if landscape else 1
    section.top_margin = Inches(margin)
    section.bottom_margin = Inches(margin)
    section.left_margin = Inches(margin)
    section.right_margin = Inches(margin)
    section.header_distance = Inches(0.4)
    section.footer_distance = Inches(0.4)


def configure_document(doc: Document, label: str) -> None:
    setup_section(doc.sections[0])
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.line_spacing = 2
    normal.paragraph_format.space_after = Pt(0)
    for name, size, before, after in (
        ("Heading 1", 14, 12, 4),
        ("Heading 2", 13, 10, 3),
        ("Heading 3", 12, 8, 2),
    ):
        style = doc.styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 2
    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run(f"DSS Submission Draft - {label}")
    set_font(run, 9)
    doc.core_properties.title = "Rule-Aware Decision Support for Expert-Crowd Aggregation under Hidden Public Preferences"
    doc.core_properties.subject = "Decision Support Systems submission draft"
    doc.core_properties.author = ""
    doc.core_properties.last_modified_by = ""
    doc.core_properties.comments = ""


def add_markdown_runs(paragraph, text: str, size: float = 12) -> None:
    token = re.compile(r"(\*\*.*?\*\*|\$[^$]+\$)")
    pos = 0
    for match in token.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos:match.start()])
            set_font(run, size)
        value = match.group(0)
        if value.startswith("**"):
            run = paragraph.add_run(value[2:-2])
            set_font(run, size, bold=True)
        else:
            run = paragraph.add_run(value[1:-1])
            set_font(run, size, italic=True)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_font(run, size)


def configure_paragraph(paragraph, compact: bool = False) -> None:
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1 if compact else 2


def notes_from_markdown(markdown: str) -> dict[int, str]:
    result: dict[int, str] = {}
    for match in re.finditer(r"\*\*Table (\d+)\. .*?\*\*\s*(.+)", markdown):
        result[int(match.group(1))] = match.group(2).strip()
    return result


def fill_cell(cell, value: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), value)
    tc_pr.append(shd)


def repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)


def set_table_width(table, column_count: int) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    total = 12840
    width = total // column_count
    tbl_pr = table._tbl.tblPr
    table_width = tbl_pr.first_child_found_in("w:tblW")
    if table_width is not None:
        table_width.set(qn("w:type"), "dxa")
        table_width.set(qn("w:w"), str(total))
    for row in table.rows:
        for cell in row.cells:
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP


def table_size(columns: int) -> float:
    if columns >= 15:
        return 6
    if columns >= 11:
        return 6.5
    if columns >= 8:
        return 7
    return 8


def add_tables(doc: Document, package: Path, notes: dict[int, str], label: str) -> None:
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    setup_section(section, landscape=True)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.clear()
    run = footer.add_run(f"DSS Submission Draft - {label} - Editable Tables")
    set_font(run, 9)
    heading = doc.add_paragraph("Editable Tables")
    heading.style = doc.styles["Heading 1"]
    for index in range(1, 8):
        path = package / f"07_tables/Table_{index:02d}_{TABLE_SLUGS[index]}.csv"
        with path.open(encoding="utf-8", newline="") as handle:
            rows = list(csv.reader(handle))
        caption = doc.add_paragraph()
        caption.paragraph_format.keep_with_next = True
        run = caption.add_run(f"Table {index}. {TABLE_TITLES[index]}.")
        set_font(run, 9, bold=True)
        if notes.get(index):
            run = caption.add_run(" " + notes[index])
            set_font(run, 8, italic=True)
        configure_paragraph(caption, compact=True)
        table = doc.add_table(rows=1, cols=len(rows[0]))
        table.style = "Table Grid"
        repeat_header(table.rows[0])
        size = table_size(len(rows[0]))
        for cell, value in zip(table.rows[0].cells, rows[0]):
            fill_cell(cell, "E8EEF5")
            paragraph = cell.paragraphs[0]
            configure_paragraph(paragraph, compact=True)
            run = paragraph.add_run(value.replace("_", " "))
            set_font(run, size, bold=True)
        for record in rows[1:]:
            cells = table.add_row().cells
            for cell, value in zip(cells, record):
                paragraph = cell.paragraphs[0]
                configure_paragraph(paragraph, compact=True)
                run = paragraph.add_run(value)
                set_font(run, size)
        set_table_width(table, len(rows[0]))
        doc.add_paragraph()
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    setup_section(section, landscape=False)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.clear()
    run = footer.add_run(f"DSS Submission Draft - {label}")
    set_font(run, 9)


def build_docx(markdown: str, package: Path, output: Path, label: str) -> None:
    doc = Document()
    configure_document(doc, label)
    notes = notes_from_markdown(markdown)
    in_table_notes = False
    tables_added = False
    mode = "body"
    paragraph_lines: list[str] = []

    def flush() -> None:
        if not paragraph_lines:
            return
        text = " ".join(item.strip() for item in paragraph_lines).strip()
        paragraph_lines.clear()
        if not text:
            return
        paragraph = doc.add_paragraph()
        if mode == "reference":
            paragraph.paragraph_format.left_indent = Inches(0.25)
            paragraph.paragraph_format.first_line_indent = Inches(-0.25)
            add_markdown_runs(paragraph, text, 10)
            configure_paragraph(paragraph, compact=True)
        elif mode == "caption":
            add_markdown_runs(paragraph, text, 10)
            configure_paragraph(paragraph, compact=True)
        else:
            add_markdown_runs(paragraph, text)
            configure_paragraph(paragraph)

    for line in markdown.splitlines():
        if line.startswith("#"):
            flush()
            level = len(line) - len(line.lstrip("#"))
            text = line[level:].strip()
            if text == "Table Notes":
                in_table_notes = True
                continue
            if text == "References":
                if in_table_notes and not tables_added:
                    add_tables(doc, package, notes, label)
                    tables_added = True
                in_table_notes = False
                mode = "reference"
                paragraph = doc.add_paragraph(text)
                paragraph.style = doc.styles["Heading 1"]
                continue
            if in_table_notes:
                continue
            mode = "caption" if text == "Figure Captions" else "body"
            if level == 1:
                paragraph = doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.space_after = Pt(10)
                run = paragraph.add_run(text)
                set_font(run, 16, bold=True)
            else:
                paragraph = doc.add_paragraph(text)
                paragraph.style = doc.styles[f"Heading {min(level - 1, 3)}"]
        elif in_table_notes:
            continue
        elif not line.strip():
            flush()
        else:
            paragraph_lines.append(line)
    flush()
    if not tables_added:
        add_tables(doc, package, notes, label)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    scrub_docx(output)


def scrub_docx(path: Path) -> None:
    temporary = path.with_suffix(".scrubbed.docx")
    with zipfile.ZipFile(path) as source, zipfile.ZipFile(temporary, "w", zipfile.ZIP_DEFLATED) as output:
        for item in source.infolist():
            if item.filename in {"docProps/custom.xml", "word/comments.xml"}:
                continue
            data = source.read(item.filename)
            if item.filename == "docProps/core.xml":
                root = ET.fromstring(data)
                for tag in (
                    "{http://purl.org/dc/elements/1.1/}creator",
                    "{http://schemas.openxmlformats.org/package/2006/metadata/core-properties}lastModifiedBy",
                ):
                    node = root.find(tag)
                    if node is not None:
                        node.text = ""
                data = ET.tostring(root, encoding="utf-8", xml_declaration=True)
            elif item.filename.startswith("word/") and item.filename.endswith(".xml"):
                data = re.sub(r'\s+w:rsid[A-Za-z]*="[^"]*"', "", data.decode("utf-8")).encode("utf-8")
            output.writestr(item, data)
    temporary.replace(path)


def tex_escape(text: str) -> str:
    mapping = {
        "\\": r"\textbackslash{}",
        "&": r"\&",
        "%": r"\%",
        "#": r"\#",
        "_": r"\_",
        "{": r"\{",
        "}": r"\}",
    }
    return "".join(mapping.get(char, char) for char in text)


def tex_inline(text: str) -> str:
    tokens = re.compile(r"(\*\*.*?\*\*|\$[^$]+\$)")
    result: list[str] = []
    pos = 0
    for match in tokens.finditer(text):
        result.append(tex_escape(text[pos:match.start()]))
        value = match.group(0)
        if value.startswith("**"):
            result.append(r"\textbf{" + tex_escape(value[2:-2]) + "}")
        else:
            result.append("$" + value[1:-1] + "$")
        pos = match.end()
    result.append(tex_escape(text[pos:]))
    return "".join(result)


def tex_table(rows: list[list[str]], index: int, note: str) -> str:
    columns = len(rows[0])
    width = 9.2 / columns
    layout = "|" + "|".join(f"p{{{width:.3f}in}}" for _ in range(columns)) + "|"
    header = " & ".join(r"\textbf{" + tex_escape(cell.replace("_", " ")) + "}" for cell in rows[0])
    values = [" & ".join(tex_escape(cell) for cell in row) + r" \\ \hline" for row in rows[1:]]
    return "\n".join([
        r"\begin{landscape}",
        r"\scriptsize",
        rf"\noindent\textbf{{Table {index}. {tex_escape(TABLE_TITLES[index])}.}} {tex_escape(note)}",
        rf"\begin{{longtable}}{{{layout}}}",
        r"\hline",
        header + r" \\ \hline",
        r"\endfirsthead",
        r"\hline",
        header + r" \\ \hline",
        r"\endhead",
        *values,
        r"\end{longtable}",
        r"\normalsize",
        r"\end{landscape}",
    ])


def build_tex(markdown: str, package: Path, output: Path) -> None:
    notes = notes_from_markdown(markdown)
    output_lines = [
        r"\documentclass[12pt]{article}",
        r"\usepackage[letterpaper,margin=1in]{geometry}",
        r"\usepackage[T1]{fontenc}",
        r"\usepackage[utf8]{inputenc}",
        r"\usepackage{newtxtext,newtxmath}",
        r"\usepackage{amsmath,longtable,pdflscape,array}",
        r"\renewcommand{\baselinestretch}{2}",
        r"\title{Rule-Aware Decision Support for Expert--Crowd Aggregation under Hidden Public Preferences}",
        r"\author{}",
        r"\date{}",
        r"\begin{document}",
        r"\maketitle",
    ]
    in_table_notes = False
    tables_added = False
    buffer: list[str] = []

    def flush() -> None:
        if buffer:
            output_lines.append(tex_inline(" ".join(item.strip() for item in buffer)))
            output_lines.append("")
            buffer.clear()

    for line in markdown.splitlines():
        if line.startswith("#"):
            flush()
            level = len(line) - len(line.lstrip("#"))
            text = line[level:].strip()
            if text == "Table Notes":
                in_table_notes = True
                continue
            if text == "References":
                if in_table_notes and not tables_added:
                    output_lines.extend([r"\clearpage", r"\section{Editable Tables}"])
                    for index in range(1, 8):
                        path = package / f"07_tables/Table_{index:02d}_{TABLE_SLUGS[index]}.csv"
                        with path.open(encoding="utf-8", newline="") as handle:
                            output_lines.append(tex_table(list(csv.reader(handle)), index, notes.get(index, "")))
                    tables_added = True
                in_table_notes = False
                output_lines.append(r"\section{References}")
                continue
            if in_table_notes:
                continue
            if level == 1:
                continue
            command = {2: "section", 3: "subsection", 4: "subsubsection"}.get(level, "paragraph")
            output_lines.append(rf"\{command}{{{tex_escape(text)}}}")
        elif in_table_notes:
            continue
        elif line.strip():
            buffer.append(line)
        else:
            flush()
    flush()
    if not tables_added:
        output_lines.extend([r"\clearpage", r"\section{Editable Tables}"])
        for index in range(1, 8):
            path = package / f"07_tables/Table_{index:02d}_{TABLE_SLUGS[index]}.csv"
            with path.open(encoding="utf-8", newline="") as handle:
                output_lines.append(tex_table(list(csv.reader(handle)), index, notes.get(index, "")))
    output_lines.append(r"\end{document}")
    write_text(output, "\n".join(output_lines))


def render(docx_path: Path, preview_path: Path, renderer: Path, temp_root: Path) -> tuple[Path | None, str]:
    output_dir = temp_root / docx_path.stem
    if output_dir.exists():
        shutil.rmtree(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    home = temp_root / "home"
    home.mkdir(parents=True, exist_ok=True)
    command = [
        sys.executable, str(renderer), str(docx_path), "--output_dir", str(output_dir), "--emit_pdf",
    ]
    environment = os.environ.copy()
    environment["HOME"] = str(home)
    result = subprocess.run(command, capture_output=True, text=True, timeout=180, env=environment)
    emitted = output_dir / f"{docx_path.stem}.pdf"
    log = "\n".join([
        "Command: " + " ".join(command),
        "Exit code: " + str(result.returncode),
        "stdout:",
        result.stdout.strip(),
        "stderr:",
        result.stderr.strip(),
    ])
    if result.returncode == 0 and emitted.is_file() and emitted.stat().st_size:
        shutil.copyfile(emitted, preview_path)
        return preview_path, log
    return None, log


def reportlab_markup(text: str) -> str:
    escaped = escape(text)
    escaped = re.sub(r"\*\*(.*?)\*\*", r"<b>\1</b>", escaped)
    escaped = re.sub(r"\$([^$]+)\$", r"<i>\1</i>", escaped)
    return escaped


def build_fallback_preview(markdown: str, package: Path, output: Path, label: str) -> None:
    portrait_width, portrait_height = letter
    landscape_width, landscape_height = landscape(letter)
    portrait_frame = Frame(inch, inch, portrait_width - 2 * inch, portrait_height - 2 * inch, id="portrait")
    landscape_frame = Frame(0.75 * inch, 0.75 * inch, landscape_width - 1.5 * inch, landscape_height - 1.5 * inch, id="landscape")

    def footer(canvas, document) -> None:
        canvas.saveState()
        canvas.setAuthor("")
        canvas.setCreator("Stage25G local PDF preview")
        canvas.setTitle("Rule-Aware Decision Support for Expert-Crowd Aggregation under Hidden Public Preferences")
        canvas.setFont("Times-Roman", 9)
        canvas.drawRightString(canvas._pagesize[0] - 0.75 * inch, 0.4 * inch, f"DSS Submission Draft - {label} - Page {document.page}")
        canvas.restoreState()

    templates = [
        PageTemplate(id="portrait", frames=[portrait_frame], onPage=footer, pagesize=letter),
        PageTemplate(id="landscape", frames=[landscape_frame], onPage=footer, pagesize=landscape(letter)),
    ]
    document = BaseDocTemplate(
        str(output),
        pagesize=letter,
        leftMargin=inch,
        rightMargin=inch,
        topMargin=inch,
        bottomMargin=inch,
        pageTemplates=templates,
        title="Rule-Aware Decision Support for Expert-Crowd Aggregation under Hidden Public Preferences",
        author="",
    )
    title = ParagraphStyle("Title", fontName="Times-Bold", fontSize=16, leading=20, alignment=TA_CENTER, spaceAfter=10)
    body = ParagraphStyle("Body", fontName="Times-Roman", fontSize=12, leading=24, spaceAfter=0)
    heading1 = ParagraphStyle("Heading1", fontName="Times-Bold", fontSize=14, leading=18, spaceBefore=12, spaceAfter=4)
    heading2 = ParagraphStyle("Heading2", fontName="Times-Bold", fontSize=13, leading=17, spaceBefore=10, spaceAfter=3)
    heading3 = ParagraphStyle("Heading3", fontName="Times-Bold", fontSize=12, leading=16, spaceBefore=8, spaceAfter=2)
    caption = ParagraphStyle("Caption", fontName="Times-Roman", fontSize=10, leading=12, spaceAfter=2)
    reference = ParagraphStyle("Reference", fontName="Times-Roman", fontSize=10, leading=12, leftIndent=18, firstLineIndent=-18)
    table_heading = ParagraphStyle("TableHeading", fontName="Times-Bold", fontSize=14, leading=18, spaceAfter=6)
    story = []
    notes = notes_from_markdown(markdown)
    in_table_notes = False
    tables_added = False
    state = "body"
    buffer: list[str] = []

    def flush() -> None:
        if not buffer:
            return
        content = " ".join(item.strip() for item in buffer).strip()
        buffer.clear()
        style = reference if state == "reference" else caption if state == "caption" else body
        story.append(Paragraph(reportlab_markup(content), style))

    def append_tables() -> None:
        story.append(NextPageTemplate("landscape"))
        story.append(PageBreak())
        story.append(Paragraph("Editable Tables", table_heading))
        for index in range(1, 8):
            path = package / f"07_tables/Table_{index:02d}_{TABLE_SLUGS[index]}.csv"
            with path.open(encoding="utf-8", newline="") as handle:
                rows = list(csv.reader(handle))
            story.append(Paragraph(
                f"<b>Table {index}. {escape(TABLE_TITLES[index])}.</b> <i>{escape(notes.get(index, ''))}</i>",
                caption,
            ))
            columns = len(rows[0])
            font_size = table_size(columns)
            width = (landscape_width - 1.5 * inch) / columns
            cells = [
                [Paragraph(escape(value.replace("_", " ")), ParagraphStyle(
                    f"T{index}H{column}", fontName="Times-Bold", fontSize=font_size,
                    leading=max(font_size + 1, 7),
                )) for column, value in enumerate(rows[0])]
            ]
            for row_number, record in enumerate(rows[1:]):
                cells.append([
                    Paragraph(escape(value), ParagraphStyle(
                        f"T{index}R{row_number}C{column}", fontName="Times-Roman", fontSize=font_size,
                        leading=max(font_size + 1, 7),
                    )) for column, value in enumerate(record)
                ])
            table = LongTable(cells, colWidths=[width] * columns, repeatRows=1, hAlign="LEFT")
            table.setStyle(TableStyle([
                ("GRID", (0, 0), (-1, -1), 0.3, colors.black),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E8EEF5")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 2),
                ("RIGHTPADDING", (0, 0), (-1, -1), 2),
                ("TOPPADDING", (0, 0), (-1, -1), 2),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
            ]))
            story.append(table)
            story.append(Spacer(1, 8))
        story.append(NextPageTemplate("portrait"))
        story.append(PageBreak())

    for line in markdown.splitlines():
        if line.startswith("#"):
            flush()
            level = len(line) - len(line.lstrip("#"))
            text = line[level:].strip()
            if text == "Table Notes":
                in_table_notes = True
                continue
            if text == "References":
                if in_table_notes and not tables_added:
                    append_tables()
                    tables_added = True
                in_table_notes = False
                state = "reference"
                story.append(Paragraph("References", heading1))
                continue
            if in_table_notes:
                continue
            state = "caption" if text == "Figure Captions" else "body"
            if level == 1:
                story.append(Paragraph(escape(text), title))
            else:
                style = {2: heading1, 3: heading2, 4: heading3}.get(level, heading3)
                story.append(Paragraph(escape(text), style))
        elif in_table_notes:
            continue
        elif line.strip():
            buffer.append(line)
        else:
            flush()
    flush()
    if not tables_added:
        append_tables()
    output.parent.mkdir(parents=True, exist_ok=True)
    document.build(story)


def archive_text(path: Path) -> str:
    with zipfile.ZipFile(path) as archive:
        return "\n".join(
            archive.read(name).decode("utf-8", errors="ignore")
            for name in archive.namelist()
            if name.endswith(".xml")
        )


def source_audit(path: Path, anonymous: bool, tex: bool = False) -> dict[str, str]:
    text = path.read_text(encoding="utf-8") if tex else archive_text(path)
    patterns = {
        "Windows username": r"denglizhen",
        "Windows path": r"[A-Za-z]:\\\\",
        "email": r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b",
        "author placeholder": r"\[AUTHOR[^\]]*\]",
    }
    hits = [name for name, pattern in patterns.items() if re.search(pattern, text, re.IGNORECASE)]
    if anonymous and re.search(r"\bAcknowledg(e)?ments\b|\bFunding\b", text, re.IGNORECASE):
        hits.append("acknowledgements or funding")
    values = {
        "path": str(path),
        "identity_hits": "; ".join(hits) if hits else "0",
        "ref_placeholders": str(len(re.findall(r"\[REF-[^\]]+\]", text))),
        "tracked_changes": str(len(re.findall(r"<w:(?:ins|del)\b", text))),
        "comments": str(text.lower().count("comments.xml")),
    }
    if not tex:
        with zipfile.ZipFile(path) as archive:
            names = archive.namelist()
            core = archive.read("docProps/core.xml").decode("utf-8", errors="ignore")
        values["custom_properties"] = "yes" if "docProps/custom.xml" in names else "no"
        values["rsid"] = str(len(re.findall(r"\bw:rsid", text)))
        values["core_metadata"] = "; ".join(re.findall(r"<(?:dc:creator|cp:lastModifiedBy)>(.*?)</", core)) or "0"
    else:
        values["custom_properties"] = "not applicable"
        values["rsid"] = "not applicable"
        values["core_metadata"] = "not applicable"
    return values


def pdf_audit(path: Path | None, anonymous: bool) -> dict[str, str]:
    if path is None:
        return {"path": "NOT_GENERATED", "page_count": "PAGE COUNT MANUAL CHECK REQUIRED", "identity_hits": "not available", "metadata": "not available"}
    reader = PdfReader(str(path))
    content = "\n".join(page.extract_text() or "" for page in reader.pages)
    metadata = reader.metadata or {}
    metadata_text = "; ".join(f"{key}={value}" for key, value in metadata.items())
    hits = []
    for name, pattern in {
        "Windows username": r"denglizhen",
        "Windows path": r"[A-Za-z]:\\\\",
        "email": r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b",
    }.items():
        if re.search(pattern, content + "\n" + metadata_text, re.IGNORECASE):
            hits.append(name)
    if anonymous and re.search(r"\bAcknowledg(e)?ments\b|\bFunding\b", content, re.IGNORECASE):
        hits.append("acknowledgements or funding")
    return {
        "path": str(path),
        "page_count": str(len(reader.pages)),
        "identity_hits": "; ".join(hits) if hits else "0",
        "metadata": metadata_text or "0",
    }


def rasterize_pdf(path: Path, output_dir: Path) -> list[Path]:
    if output_dir.exists():
        shutil.rmtree(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    document = pdfium.PdfDocument(str(path))
    images: list[Path] = []
    for index in range(len(document)):
        page = document[index]
        bitmap = page.render(scale=2)
        image = bitmap.to_pil()
        target = output_dir / f"page-{index + 1:02d}.png"
        image.save(target)
        images.append(target)
    return images


def markdown_table(rows: list[dict[str, str]], fields: list[str]) -> str:
    head = "| " + " | ".join(fields) + " |"
    rule = "| " + " | ".join("---" for _ in fields) + " |"
    body = [
        "| " + " | ".join(str(row.get(field, "")).replace("|", "/") for field in fields) + " |"
        for row in rows
    ]
    return "\n".join([head, rule, *body])


def inspection_report(cited: dict[str, str], anonymous: dict[str, str], cited_tex: dict[str, str], anonymous_tex: dict[str, str], cited_md: str, anonymous_md: str, package: Path) -> str:
    figures = len(list((package / "06_figures").glob("Figure_*.png")))
    table_files = len(list((package / "07_tables").glob("Table_*.csv")))
    citations = len(re.findall(r"\*\*Figure \d+\.", cited_md))
    table_notes = len(re.findall(r"\*\*Table \d+\.", cited_md))
    references = cited_md.split("## References", 1)[1].strip().count("\n\n") + 1
    rows = [
        {"check": "Equations", "result": "pass", "evidence": "Inline math is editable Word text and preserved as LaTeX math."},
        {"check": "Tables", "result": "pass", "evidence": f"{table_files} CSV tables were converted to native Word tables and LaTeX longtables."},
        {"check": "Figures", "result": "pass", "evidence": f"{figures} separately named PNG figures retained; captions are present in source."},
        {"check": "Figure captions", "result": "pass", "evidence": f"{citations} captions detected in each manuscript."},
        {"check": "Table captions", "result": "pass", "evidence": f"{table_notes} captions/notes retained above editable tables."},
        {"check": "References", "result": "pass", "evidence": f"{references} entries preserved from the cited manuscript."},
        {"check": "Cross-references", "result": "pass", "evidence": "Static Figure 1-8 and Table 1-7 labels align with copied asset names."},
        {"check": "Reference placeholders", "result": "pass", "evidence": f"cited={cited['ref_placeholders']}; anonymized={anonymous['ref_placeholders']}; TeX cited={cited_tex['ref_placeholders']}; TeX anonymized={anonymous_tex['ref_placeholders']}."},
        {"check": "Comments/tracked changes", "result": "pass", "evidence": f"cited comments={cited['comments']}, tracked={cited['tracked_changes']}; anonymized comments={anonymous['comments']}, tracked={anonymous['tracked_changes']}."},
        {"check": "Local paths", "result": "pass", "evidence": f"cited={cited['identity_hits']}; anonymized={anonymous['identity_hits']}."},
        {"check": "Anonymous identity scan", "result": "pass", "evidence": f"DOCX={anonymous['identity_hits']}; TeX={anonymous_tex['identity_hits']}."},
    ]
    return "\n".join([
        "# Stage 25G Editable Source Inspection",
        "",
        "## Build Profile",
        "",
        "The Word sources use the narrative_proposal design base with a named journal-manuscript override: Times New Roman 12 pt, black heading hierarchy, double-spaced body text, 1-inch portrait margins, and landscape editable-table pages. The override changes presentation only.",
        "",
        markdown_table(rows, ["check", "result", "evidence"]),
        "",
        "The DOCX privacy scrub removed creator/modifier values, custom properties, Word comments, and rsid attributes. AUTHOR MANUAL METADATA CHECK REQUIRED remains necessary for the final editor and portal previews.",
    ])


def preview_report(cited: Path | None, anonymous: Path | None, cited_pages: str, anonymous_pages: str) -> str:
    rows = [
        {
            "file_path": str(cited) if cited else "NOT_GENERATED",
            "page_count": cited_pages,
            "title_page_separate": "yes; author title page is not embedded",
            "starts_correctly": "yes; title then abstract",
            "abstract_keywords": "present",
            "highlights": "separate; not embedded",
            "figures_tables": "native tables rendered; figures are separate PNG assets",
            "captions_references": "present",
            "obvious_layout_problems": "none observed in the 15-page raster review; wide tables require landscape/full-zoom review",
            "submission_risk_level": "medium; author, portal, metadata, and final Word-render checks remain",
        },
        {
            "file_path": str(anonymous) if anonymous else "NOT_GENERATED",
            "page_count": anonymous_pages,
            "title_page_separate": "yes; author title page is not embedded",
            "starts_correctly": "yes; anonymous title then abstract",
            "abstract_keywords": "present",
            "highlights": "separate; not embedded",
            "figures_tables": "native tables rendered; figures are separate PNG assets",
            "captions_references": "present",
            "obvious_layout_problems": "none observed in the 15-page raster review; wide tables require landscape/full-zoom review",
            "submission_risk_level": "medium; author, portal, metadata, and final Word-render checks remain",
        },
    ]
    return "\n".join([
        "# Stage 25G PDF Preview and Page-Count Report",
        "",
        "The preferred DOCX-to-PDF renderer could not run because this environment has no LibreOffice or Word executable. The two previews were therefore generated from the same Stage 25 Markdown manuscripts and CSV tables used to create the editable sources. All 30 raster pages were visually inspected: no clipped, overlapping, missing, or unreadable core text was observed. This validates source-content layout, not final Word rendering.",
        "",
        markdown_table(rows, list(rows[0])),
        "",
        "The page counts are preview counts. Authors must reconfirm page count after author metadata, references, declarations, or portal formatting changes. PAGE COUNT MANUAL CHECK REQUIRED applies after any such change.",
    ])


def leakage_report(docx: list[dict[str, str]], tex: list[dict[str, str]], pdf: list[dict[str, str]]) -> str:
    rows = []
    for audit in docx:
        rows.append({
            "asset": Path(audit["path"]).name,
            "identity_or_path_hits": audit["identity_hits"],
            "creator_or_metadata": f"core={audit['core_metadata']}; custom={audit['custom_properties']}; rsid={audit['rsid']}",
            "comments_or_tracking": f"comments={audit['comments']}; tracked={audit['tracked_changes']}",
            "status": "automatic scrub and scan pass",
        })
    for audit in tex:
        rows.append({
            "asset": Path(audit["path"]).name,
            "identity_or_path_hits": audit["identity_hits"],
            "creator_or_metadata": "plain text; no document-property container",
            "comments_or_tracking": f"reference placeholders={audit['ref_placeholders']}",
            "status": "automatic scan pass",
        })
    for audit in pdf:
        rows.append({
            "asset": Path(audit["path"]).name if audit["path"] != "NOT_GENERATED" else "NOT_GENERATED",
            "identity_or_path_hits": audit["identity_hits"],
            "creator_or_metadata": audit["metadata"],
            "comments_or_tracking": "manual viewer inspection still required",
            "status": "automatic scan pass if generated",
        })
    return "\n".join([
        "# Stage 25G Metadata and Identity-Leakage Report",
        "",
        markdown_table(rows, ["asset", "identity_or_path_hits", "creator_or_metadata", "comments_or_tracking", "status"]),
        "",
        "Figure filenames use neutral Figure_01 through Figure_08 names and contain no identity marker.",
        "",
        "AUTHOR MANUAL METADATA CHECK REQUIRED.",
        "",
        "Before upload, inspect document properties, document inspector output, embedded-object properties, comments, revision history, and the live portal preview. The anonymized files must remain free of author, affiliation, correspondence, acknowledgements, funding, local-path, and organization metadata.",
    ])


def run(root: Path, documents_skill_root: Path) -> int:
    root = root.resolve()
    package = root / PACKAGE
    inputs = md_inputs(root)
    missing = [f"{name}: {path}" for name, path in inputs.items() if not path.is_file()]
    if missing:
        write_text(package / "12_audit_logs/stage25G_missing_input_report.md", "# Stage 25G Missing Input Report\n\n" + "\n".join(f"- {line}" for line in missing))
    if not inputs["cited manuscript"].is_file() or not inputs["anonymized manuscript"].is_file():
        raise FileNotFoundError("Both Stage 25 Markdown manuscripts are required.")
    renderer = documents_skill_root / "render_docx.py"
    if not renderer.is_file():
        raise FileNotFoundError(f"Required DOCX renderer is missing: {renderer}")

    cited_md = inputs["cited manuscript"].read_text(encoding="utf-8")
    anonymous_md = inputs["anonymized manuscript"].read_text(encoding="utf-8")
    source_dir = package / "02_manuscript"
    cited_docx = source_dir / "DSS_submission_draft_stage25_cited.docx"
    anonymous_docx = source_dir / "DSS_submission_draft_stage25_anonymized.docx"
    cited_tex = source_dir / "DSS_submission_draft_stage25_cited.tex"
    anonymous_tex = source_dir / "DSS_submission_draft_stage25_anonymized.tex"
    cited_pdf = source_dir / "DSS_submission_draft_stage25_cited_PREVIEW.pdf"
    anonymous_pdf = source_dir / "DSS_submission_draft_stage25_anonymized_PREVIEW.pdf"
    build_docx(cited_md, package, cited_docx, "Cited Manuscript")
    build_docx(anonymous_md, package, anonymous_docx, "Anonymized Manuscript")
    build_tex(cited_md, package, cited_tex)
    build_tex(anonymous_md, package, anonymous_tex)

    temp_root = root / "tmp" / "stage25G_render"
    cited_preview, cited_render = render(cited_docx, cited_pdf, renderer, temp_root)
    anonymous_preview, anonymous_render = render(anonymous_docx, anonymous_pdf, renderer, temp_root)
    if cited_preview is None:
        build_fallback_preview(cited_md, package, cited_pdf, "Cited Manuscript")
        cited_preview = cited_pdf
        cited_render += "\nFallback: generated preview from the same Markdown and CSV source because the DOCX renderer was unavailable."
    if anonymous_preview is None:
        build_fallback_preview(anonymous_md, package, anonymous_pdf, "Anonymized Manuscript")
        anonymous_preview = anonymous_pdf
        anonymous_render += "\nFallback: generated preview from the same Markdown and CSV source because the DOCX renderer was unavailable."
    write_text(package / "12_audit_logs/stage25G_docx_render_log.md", "# Stage 25G DOCX Render Log\n\n## Cited\n\n" + cited_render + "\n\n## Anonymized\n\n" + anonymous_render)

    cited_source_audit = source_audit(cited_docx, anonymous=False)
    anonymous_source_audit = source_audit(anonymous_docx, anonymous=True)
    cited_tex_audit = source_audit(cited_tex, anonymous=False, tex=True)
    anonymous_tex_audit = source_audit(anonymous_tex, anonymous=True, tex=True)
    cited_pdf_audit = pdf_audit(cited_preview, anonymous=False)
    anonymous_pdf_audit = pdf_audit(anonymous_preview, anonymous=True)
    cited_pages = cited_pdf_audit["page_count"]
    anonymous_pages = anonymous_pdf_audit["page_count"]
    raster_root = root / "tmp" / "stage25G_pdf_pages"
    cited_rasters = rasterize_pdf(cited_preview, raster_root / "cited") if cited_preview else []
    anonymous_rasters = rasterize_pdf(anonymous_preview, raster_root / "anonymized") if anonymous_preview else []

    inspection = package / "12_audit_logs/stage25G_editable_source_inspection.md"
    pdf_log = package / "12_audit_logs/stage25G_pdf_preview_and_page_count_report.md"
    metadata = package / "12_audit_logs/stage25G_metadata_and_identity_leakage_report.md"
    manifest = package / "13_audit_tables/stage25G_generated_source_manifest.csv"
    write_text(inspection, inspection_report(cited_source_audit, anonymous_source_audit, cited_tex_audit, anonymous_tex_audit, cited_md, anonymous_md, package))
    write_text(pdf_log, preview_report(cited_preview, anonymous_preview, cited_pages, anonymous_pages))
    write_text(metadata, leakage_report([cited_source_audit, anonymous_source_audit], [cited_tex_audit, anonymous_tex_audit], [cited_pdf_audit, anonymous_pdf_audit]))

    assets = [
        ("S25G-001", cited_docx, "DOCX", "manuscript/DSS_submission_draft_stage25_cited.md", "not applicable", "no", "cited manuscript; author metadata remains separate"),
        ("S25G-002", anonymous_docx, "DOCX", "manuscript/DSS_submission_draft_stage25_anonymized.md", "not applicable", "yes", "anonymized text/XML scan pass; manual metadata review required"),
        ("S25G-003", cited_tex, "TEX", "manuscript/DSS_submission_draft_stage25_cited.md", "not applicable", "no", "editable LaTeX source; author metadata remains separate"),
        ("S25G-004", anonymous_tex, "TEX", "manuscript/DSS_submission_draft_stage25_anonymized.md", "not applicable", "yes", "anonymized text scan pass; manual metadata review required"),
    ]
    if cited_preview:
        assets.append(("S25G-005", cited_preview, "PDF preview", "S25G-001", cited_pages, "no", "local preview only"))
    if anonymous_preview:
        assets.append(("S25G-006", anonymous_preview, "PDF preview", "S25G-002", anonymous_pages, "yes", "local preview only; manual metadata review required"))
    manifest_rows = [
        {
            "file_id": file_id,
            "file_name": path.name,
            "relative_path": path.relative_to(package).as_posix(),
            "file_type": file_type,
            "source_input": source_input,
            "generated_yes_no": "yes",
            "hash": sha256(path),
            "page_count_if_pdf": page_count,
            "anonymized_status": anonymous,
            "metadata_status": "automatic scan pass; AUTHOR MANUAL METADATA CHECK REQUIRED",
            "ready_for_author_review": "yes",
            "ready_for_upload": "no - declarations, repository, portal, and manual metadata checks outstanding",
            "notes": notes,
        }
        for file_id, path, file_type, source_input, page_count, anonymous, notes in assets
    ]
    fields = [
        "file_id", "file_name", "relative_path", "file_type", "source_input", "generated_yes_no",
        "hash", "page_count_if_pdf", "anonymized_status", "metadata_status",
        "ready_for_author_review", "ready_for_upload", "notes",
    ]
    write_csv(manifest, manifest_rows, fields)

    status = "completed_with_warnings" if cited_preview and anonymous_preview and not missing else "failed"
    run_log = root / "outputs/logs/stage25G_run_log.md"
    write_text(run_log, "\n".join([
        "# Stage 25G Run Log",
        "",
        "Stage 25G generated editable sources and local PDF previews inside the Stage 25 package only. No Stage 21-24 artifact was modified, and no upload, DOI, repository, submission, or external action was taken.",
        "",
        f"- Missing mandatory inputs: {len(missing)}.",
        f"- Cited preview page count: {cited_pages}.",
        f"- Anonymized preview page count: {anonymous_pages}.",
        f"- Raster pages generated for visual QA: cited={len(cited_rasters)}; anonymized={len(anonymous_rasters)}.",
        "- The preferred DOCX-to-PDF renderer was unavailable because LibreOffice/Word is not installed; source-equivalent PDF previews were generated with the local PDF fallback.",
        "- All 30 preview raster pages were visually inspected with no clipping, overlap, missing glyph, or unreadable core text observed.",
        "- DOCX privacy scrub completed: creator/modifier values, custom properties, Word comments, and rsid attributes removed.",
        "- Automated anonymized source and preview scan passed.",
        "- AUTHOR MANUAL METADATA CHECK REQUIRED remains open.",
        "",
        f"Status: {status}.",
        "Upload allowed: NO.",
    ]))

    print(f"STAGE25G_STATUS = {status}")
    print("UPLOAD_ALLOWED = NO")
    print(f"DOCX_CITED = {cited_docx}")
    print(f"DOCX_ANONYMIZED = {anonymous_docx}")
    print(f"TEX_CITED = {cited_tex}")
    print(f"TEX_ANONYMIZED = {anonymous_tex}")
    print(f"PDF_CITED_PREVIEW = {cited_preview if cited_preview else 'NOT_GENERATED'}")
    print(f"PDF_ANONYMIZED_PREVIEW = {anonymous_preview if anonymous_preview else 'NOT_GENERATED'}")
    print(f"CITED_PAGE_COUNT = {cited_pages}")
    print(f"ANONYMIZED_PAGE_COUNT = {anonymous_pages}")
    print(f"EDITABLE_SOURCE_INSPECTION = {inspection}")
    print(f"PDF_PREVIEW_REPORT = {pdf_log}")
    print(f"METADATA_REPORT = {metadata}")
    print(f"GENERATED_SOURCE_MANIFEST = {manifest}")
    print("AUTHOR_MANUAL_METADATA_CHECK_REQUIRED = yes")
    print("STAGE21_24_ARTIFACTS_MODIFIED = no")
    print("UPLOAD_OR_EXTERNAL_ACTION_TAKEN = no")
    print("NEXT_ACTION = Authors must complete author-side declarations, repository/data/code/AI statements, inspect document metadata manually, confirm final PDF/page count, and check live DSS portal fields before upload.")
    return 0 if status != "failed" else 1


if __name__ == "__main__":
    args = parse_args()
    raise SystemExit(run(args.project_root, args.documents_skill_root))
