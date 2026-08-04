#!/usr/bin/env python3
"""Create a non-anonymized Stage 25 author-review DOCX from the cited source."""

from __future__ import annotations

import argparse
import csv
import hashlib
import re
import shutil
import zipfile
from datetime import datetime, timezone
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Pt


PROJECT_ROOT = Path(__file__).resolve().parents[1]
PACKAGE = Path("submission_package_stage25")
TITLE = "Rule-Aware Decision Support for Expert-Crowd Aggregation under Hidden Public Preferences"
CORRESPONDING_ADDRESS = (
    "Huazhong University of Science and Technology, 1037 Luoyu Road, "
    "Hongshan District, Wuhan, Hubei 430074, China"
)
OUTPUT_NAME = "DSS_submission_draft_stage25_non_anonymized_author_review.docx"


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Create a non-anonymized Stage 25 author-review DOCX with the confirmed title page."
    )
    parser.add_argument("--project-root", type=Path, default=PROJECT_ROOT)
    return parser.parse_args()


def write_text(path: Path, content: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    temporary = path.with_suffix(path.suffix + ".tmp")
    temporary.write_text(content.strip() + "\n", encoding="utf-8")
    temporary.replace(path)


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def verify_frozen_artifacts(root: Path) -> int:
    manifest = root / PACKAGE / "11_reproducibility/frozen_artifact_hash_manifest_stage25.csv"
    if not manifest.is_file():
        raise FileNotFoundError(f"Frozen-artifact manifest is missing: {manifest}")
    with manifest.open("r", encoding="utf-8", newline="") as handle:
        rows = list(csv.DictReader(handle))
    if not rows:
        raise ValueError(f"Frozen-artifact manifest has no entries: {manifest}")

    mismatches: list[str] = []
    for row in rows:
        relative_path = row.get("relative_path", "")
        expected = row.get("expected_sha256", "")
        target = root / relative_path
        if not relative_path or not expected or not target.is_file():
            mismatches.append(f"{relative_path or '<missing path>'}: malformed or missing")
            continue
        observed = sha256(target)
        if observed != expected:
            mismatches.append(f"{relative_path}: expected {expected}, observed {observed}")
    if mismatches:
        raise RuntimeError(
            "Frozen Stage 21--24 artifact verification failed; no author-review DOCX was written. "
            + "; ".join(mismatches)
        )
    return len(rows)


def set_font(run, size: float, *, bold: bool = False, italic: bool = False) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_title_page(doc: Document) -> None:
    """Build title-page paragraphs at the end, then move them before the source body."""
    created = []

    def paragraph(text: str = "", *, size: float = 10.5, bold: bool = False,
                  italic: bool = False, center: bool = False, before: float = 0,
                  after: float = 4) -> None:
        item = doc.add_paragraph()
        item.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
        item.paragraph_format.space_before = Pt(before)
        item.paragraph_format.space_after = Pt(after)
        item.paragraph_format.line_spacing = 1
        item.paragraph_format.keep_with_next = bool(bold and text.endswith(":"))
        if text:
            run = item.add_run(text)
            set_font(run, size, bold=bold, italic=italic)
        created.append(item._p)

    paragraph(TITLE, size=16, bold=True, center=True, after=14)
    paragraph("Deng Lizhen; Liu Yuxin; Li Bo", size=12, center=True, after=12)
    paragraph("Affiliations", size=11, bold=True, after=3)
    paragraph("1. Huazhong University of Science and Technology", after=1)
    paragraph("2. Wuhan University of Technology", after=1)
    paragraph("3. Wuhan University of Technology", after=9)
    paragraph("Corresponding Author", size=11, bold=True, after=3)
    paragraph("Deng Lizhen", after=1)
    paragraph("Email: 3070116993@qq.com", after=1)
    paragraph(f"Address: {CORRESPONDING_ADDRESS}", after=9)
    paragraph("Acknowledgements", size=11, bold=True, after=3)
    paragraph("None.", after=7)
    paragraph("Funding", size=11, bold=True, after=3)
    paragraph(
        "This research received no specific grant from any funding agency in the public, commercial, or not-for-profit sectors.",
        after=7,
    )
    paragraph("Declaration of Competing Interest", size=11, bold=True, after=3)
    paragraph(
        "The authors declare that they have no known competing financial interests or personal relationships that could have appeared to influence the work reported in this paper.",
        after=10,
    )
    paragraph("Author-review version. Final CRediT confirmation, concrete repository URL, and all-author approval remain pending before upload.", size=9, italic=True, after=0)

    page_break = doc.add_paragraph()
    page_break.paragraph_format.space_after = Pt(0)
    page_break.add_run().add_break(WD_BREAK.PAGE)
    created.append(page_break._p)

    body = doc._element.body
    for index, element in enumerate(created):
        body.insert(index, element)

    first_section = doc.sections[0]
    first_section.different_first_page_header_footer = True
    for paragraph_item in first_section.first_page_footer.paragraphs:
        paragraph_item.clear()


def scrub_metadata(path: Path) -> dict[str, int]:
    """Clear document metadata, custom properties, comments, and revision session IDs."""
    temporary = path.with_suffix(".scrubbed.docx")
    stats = {"core": 0, "custom": 0, "comments": 0, "rsid": 0}
    namespaces = {
        "dc": "http://purl.org/dc/elements/1.1/",
        "cp": "http://schemas.openxmlformats.org/package/2006/metadata/core-properties",
    }
    with zipfile.ZipFile(path, "r") as source, zipfile.ZipFile(temporary, "w", zipfile.ZIP_DEFLATED) as output:
        for item in source.infolist():
            if item.filename == "docProps/custom.xml":
                stats["custom"] += 1
                continue
            if item.filename == "word/comments.xml":
                stats["comments"] += 1
                continue
            data = source.read(item.filename)
            if item.filename == "docProps/core.xml":
                root = ET.fromstring(data)
                for tag in (
                    f"{{{namespaces['dc']}}}creator",
                    f"{{{namespaces['cp']}}}lastModifiedBy",
                ):
                    node = root.find(tag)
                    if node is not None and (node.text or ""):
                        node.text = ""
                        stats["core"] += 1
                data = ET.tostring(root, encoding="utf-8", xml_declaration=True)
            elif item.filename.startswith("word/") and item.filename.endswith(".xml"):
                decoded = data.decode("utf-8")
                cleaned, count = re.subn(r'\s+w:rsid[A-Za-z]*="[^"]*"', "", decoded)
                stats["rsid"] += count
                data = cleaned.encode("utf-8")
            output.writestr(item, data)
    temporary.replace(path)
    return stats


def audit_docx(path: Path) -> dict[str, object]:
    document = Document(path)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    with zipfile.ZipFile(path) as archive:
        names = set(archive.namelist())
        core = archive.read("docProps/core.xml").decode("utf-8") if "docProps/core.xml" in names else ""
        revision_hits = sum(
            len(re.findall(r"rsid[A-Za-z]*=", archive.read(name).decode("utf-8")))
            for name in names
            if name.startswith("word/") and name.endswith(".xml")
        )
    return {
        "paragraphs": len(document.paragraphs),
        "tables": len(document.tables),
        "sections": len(document.sections),
        "title_present": TITLE in text,
        "author_present": "Deng Lizhen; Liu Yuxin; Li Bo" in text,
        "address_present": CORRESPONDING_ADDRESS in text,
        "pending_boundary_present": "Final CRediT confirmation, concrete repository URL" in text,
        "comments_present": "word/comments.xml" in names,
        "custom_properties_present": "docProps/custom.xml" in names,
        "revision_session_ids": revision_hits,
        "creator_present": "<dc:creator>" in core and "<dc:creator></dc:creator>" not in core,
        "last_modified_by_present": "<cp:lastModifiedBy>" in core and "<cp:lastModifiedBy></cp:lastModifiedBy>" not in core,
    }


def audit_report(
    source: Path,
    output: Path,
    frozen_count: int,
    scrub: dict[str, int],
    audit: dict[str, object],
    render_status: str,
) -> str:
    return f"""# Stage 25H-C Non-Anonymized Author-Review DOCX Generation

## Scope

The output was created by copying the Stage 25 cited DOCX and prepending a non-anonymized title page with author-confirmed names, affiliations, corresponding-author details, funding, and competing-interest declarations. The manuscript body, editable tables, and cited figure references were retained from the Stage 25 cited source. No Stage 21-24 artifact was modified.

## Files

- Source: `{source.relative_to(source.parents[3]).as_posix()}`
- Output: `{output.relative_to(output.parents[3]).as_posix()}`
- Output SHA-256: `{sha256(output)}`

## Structural Checks

| check | result |
| --- | --- |
| frozen Stage 21-24 manifest | {frozen_count}/{frozen_count} SHA-256 entries matched |
| paragraphs | {audit['paragraphs']} |
| editable tables | {audit['tables']} |
| sections | {audit['sections']} |
| confirmed title page fields | title={audit['title_present']}; authors={audit['author_present']}; address={audit['address_present']} |
| pending submission boundary retained | {audit['pending_boundary_present']} |
| Word comments | {audit['comments_present']} |
| custom properties | {audit['custom_properties_present']} |
| revision session IDs | {audit['revision_session_ids']} |
| creator metadata present | {audit['creator_present']} |
| last-modified-by metadata present | {audit['last_modified_by_present']} |
| visual render QA | {render_status} |

## Metadata Scrub

Core properties cleared: {scrub['core']}; custom properties removed: {scrub['custom']}; comments removed: {scrub['comments']}; revision session IDs removed: {scrub['rsid']}.

## Submission Boundary

This is a non-anonymized author-review DOCX, not an upload-authorized source. A concrete public repository URL, final DOCX-exported PDF page count, all-author CRediT confirmation, and final all-author approval remain required.
"""


def run(root: Path) -> int:
    root = root.resolve()
    package = root / PACKAGE
    source = package / "02_manuscript/DSS_submission_draft_stage25_cited.docx"
    author_packet = package / "01_author_action_required/AUTHOR_FILL_IN_PACKET_STAGE25F.md"
    required = [
        source,
        author_packet,
        package / "01_author_action_required/STAGE25H_C_RERUN_DECISION.md",
        package / "11_reproducibility/frozen_artifact_hash_manifest_stage25.csv",
    ]
    missing = [str(path) for path in required if not path.is_file()]
    if missing:
        raise FileNotFoundError("Missing required author-review DOCX inputs: " + "; ".join(missing))

    frozen_count = verify_frozen_artifacts(root)
    output = package / "02_manuscript" / OUTPUT_NAME
    shutil.copy2(source, output)
    document = Document(output)
    add_title_page(document)
    document.core_properties.title = TITLE
    document.core_properties.subject = "Decision Support Systems non-anonymized author-review manuscript"
    document.core_properties.author = ""
    document.core_properties.last_modified_by = ""
    document.core_properties.keywords = ""
    document.core_properties.created = datetime.now(timezone.utc)
    document.save(output)

    scrub = scrub_metadata(output)
    audit = audit_docx(output)
    failed = [key for key in (
        "title_present", "author_present", "address_present", "pending_boundary_present"
    ) if not audit[key]]
    if audit["comments_present"] or audit["custom_properties_present"] or audit["revision_session_ids"]:
        failed.append("metadata scrub")
    if failed:
        raise RuntimeError("Author-review DOCX audit failed: " + ", ".join(failed))

    render_status = (
        "not completed: LibreOffice/soffice is unavailable in the current environment; "
        "structural and metadata checks completed"
        if shutil.which("soffice") is None
        else "not run by the generation script; run the document renderer before final upload"
    )
    report = package / "12_audit_logs/stage25H_C_non_anonymized_author_review_docx.md"
    write_text(report, audit_report(source, output, frozen_count, scrub, audit, render_status))
    write_text(
        root / "outputs/logs/stage25H_C_final_docx_run_log.md",
        "\n".join([
            "# Stage 25H-C Final DOCX Run Log",
            "",
            f"Generated: {output.relative_to(root).as_posix()}",
            f"Frozen Stage 21-24 verification: {frozen_count}/{frozen_count} SHA-256 entries matched.",
            "The file is non-anonymized and prepared for author review only.",
            "No repository, upload, DOI, deposit, or submission record was created.",
            f"Visual render QA: {render_status}.",
            "Final page-count and all-author approval blockers remain open.",
        ]),
    )
    print("STAGE25H_C_FINAL_DOCX_STATUS = completed_with_warnings")
    print(f"FINAL_DOCX = {output}")
    print("DOCX_ROLE = non_anonymized_author_review")
    print("UPLOAD_ALLOWED = NO")
    print("STAGE21_24_ARTIFACTS_MODIFIED = no")
    print("UPLOAD_OR_EXTERNAL_ACTION_TAKEN = no")
    return 0


if __name__ == "__main__":
    args = parse_args()
    raise SystemExit(run(args.project_root))
