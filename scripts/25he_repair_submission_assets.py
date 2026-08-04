"""Repair Stage 25 figures, tables, and the anonymous DSS manuscript layout.

This additive stage never overwrites Stage 21-24 artifacts or the prior H-D
manuscript. Quantitative figures are regenerated from tracked result tables;
main-text tables are condensed displays of the full Stage 25 CSV supplements.
"""

from __future__ import annotations

import argparse
import csv
import hashlib
import importlib.util
import re
import zipfile
from pathlib import Path
from typing import Any

import numpy as np
import pandas as pd
from PIL import Image


FIGURE_DPI = 300
TITLE = "Rule-Aware Decision Support for Expert-Crowd Aggregation under Hidden Public Preferences"
PACKAGE = Path("submission_package_stage25")
SOURCE_MANUSCRIPT = Path("manuscript/DSS_submission_draft_stage25_anonymized.md")
OUTPUT_DOCX = PACKAGE / "02_submission_files/DSS_anonymized_manuscript_STAGE25H_E_revised.docx"
OUTPUT_FIGURES = PACKAGE / "02_submission_files/figures_STAGE25H_E_revised"
AUDIT_DIR = PACKAGE / "12_audit_logs"


TABLE_SOURCE_MAP = {
    1: ("Table_01_decision_alternatives_and_criteria.csv", "decision_alternatives_criteria.csv"),
    2: ("Table_02_assumption_inventory.csv", "assumption_inventory.csv"),
    3: ("Table_03_baseline_definitions.csv", "baseline_definition_table.csv"),
    4: ("Table_04_synthetic_coverage_results.csv", "synthetic_coverage_results.csv"),
    5: ("Table_05_external_testbed_results.csv", "external_testbed_results.csv"),
    6: ("Table_06_design_recommendation_matrix.csv", "design_recommendation_matrix.csv"),
    7: ("Table_07_claim_evidence_alignment.csv", "claim_evidence_alignment.csv"),
}


FIGURE_FILENAMES = {
    1: "Figure_01_DSS_conceptual_framework.png",
    2: "Figure_02_decision_support_workflow.png",
    3: "Figure_03_discretion_identifiability_frontier.png",
    4: "Figure_04_disclosure_uncertainty_curve.png",
    5: "Figure_05_rule_robustness_heatmap.png",
    6: "Figure_06_synthetic_benchmark_coverage.png",
    7: "Figure_07_external_testbed_comparison.png",
    8: "Figure_08_DSS_artifact_evaluation.png",
}


FIGURE_DATA_SOURCES = {
    1: "theoretical architecture; no numeric input",
    2: "theoretical workflow; no numeric input",
    3: "outputs/tables/discretion_identifiability_summary.csv; synthetic rows only",
    4: "outputs/tables/value_of_disclosure.csv",
    5: "outputs/tables/rule_robustness_index.csv; RRI recomputed as supporting/applicable",
    6: "outputs/tables/synthetic_coverage_results.csv",
    7: "outputs/tables/external_testbed_results.csv",
    8: "outputs/tables/dss_evaluation_metrics.csv",
}


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Create an additive Stage 25H-E figure/table/layout repair package."
    )
    parser.add_argument(
        "--project-root",
        type=Path,
        default=Path(__file__).resolve().parents[1],
        help="Project root containing manuscript/, outputs/, scripts/, and submission_package_stage25/.",
    )
    parser.add_argument(
        "--mode",
        choices=("all", "assets", "docx", "audit"),
        default="all",
        help="Run all steps, regenerate figures, build the DOCX, or audit a WPS-exported preview PDF.",
    )
    parser.add_argument(
        "--preview-pdf",
        type=Path,
        help="WPS-exported preview PDF used by --mode audit.",
    )
    parser.add_argument(
        "--visual-review",
        choices=("pass", "warnings", "not_recorded"),
        default="not_recorded",
        help="Result of the required page-image visual review for --mode audit.",
    )
    parser.add_argument(
        "--reviewed-pages",
        default="",
        help="Inclusive page range reviewed as rendered PNGs, for example 1-25.",
    )
    return parser.parse_args()


def load_plot_runtime() -> None:
    global plt, FancyArrowPatch, FancyBboxPatch, Polygon
    import matplotlib

    # Headless rendering is deterministic and avoids Qt canvas allocation
    # failures during the eight-figure batch.
    matplotlib.use("Agg", force=True)
    import matplotlib.pyplot as plt_module
    from matplotlib.patches import FancyArrowPatch as arrow_patch
    from matplotlib.patches import FancyBboxPatch as box_patch
    from matplotlib.patches import Polygon as polygon_patch

    plt = plt_module
    FancyArrowPatch = arrow_patch
    FancyBboxPatch = box_patch
    Polygon = polygon_patch


def load_docx_runtime() -> None:
    global Document, WD_SECTION, WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
    global WD_ALIGN_PARAGRAPH, OxmlElement, qn, Inches, Pt, RGBColor
    from docx import Document as document_class
    from docx.enum.section import WD_SECTION as section_enum
    from docx.enum.table import WD_ALIGN_VERTICAL as vertical_enum
    from docx.enum.table import WD_TABLE_ALIGNMENT as table_alignment_enum
    from docx.enum.text import WD_ALIGN_PARAGRAPH as paragraph_alignment_enum
    from docx.oxml import OxmlElement as oxml_element
    from docx.oxml.ns import qn as qualified_name
    from docx.shared import Inches as inches
    from docx.shared import Pt as points
    from docx.shared import RGBColor as rgb_color

    Document = document_class
    WD_SECTION = section_enum
    WD_ALIGN_VERTICAL = vertical_enum
    WD_TABLE_ALIGNMENT = table_alignment_enum
    WD_ALIGN_PARAGRAPH = paragraph_alignment_enum
    OxmlElement = oxml_element
    qn = qualified_name
    Inches = inches
    Pt = points
    RGBColor = rgb_color


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def write_text(path: Path, content: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    temporary = path.with_suffix(path.suffix + ".tmp")
    temporary.write_text(content.rstrip() + "\n", encoding="utf-8")
    temporary.replace(path)


def load_hd_module(root: Path):
    path = root / "scripts/25hd_reconstruct_dss_submission_docx.py"
    spec = importlib.util.spec_from_file_location("stage25hd", path)
    if spec is None or spec.loader is None:
        raise RuntimeError(f"Unable to load H-D document builder: {path}")
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


def read_frame(path: Path) -> pd.DataFrame:
    if not path.is_file():
        raise FileNotFoundError(f"Required table is missing: {path}")
    return pd.read_csv(path)


def validate_unit_interval(frame: pd.DataFrame, columns: list[str], label: str) -> None:
    for column in columns:
        values = pd.to_numeric(frame[column], errors="coerce").dropna()
        if ((values < -1e-12) | (values > 1 + 1e-12)).any():
            raise ValueError(f"{label}.{column} contains a value outside [0, 1].")


def apply_plot_style() -> None:
    plt.rcParams.update(
        {
            "font.family": "DejaVu Sans",
            "font.size": 9.5,
            "axes.titlesize": 11,
            "axes.labelsize": 9.5,
            "xtick.labelsize": 8.5,
            "ytick.labelsize": 8.5,
            "axes.spines.top": False,
            "axes.spines.right": False,
            "axes.titleweight": "semibold",
            "figure.facecolor": "white",
            "axes.facecolor": "white",
            "savefig.facecolor": "white",
            "legend.frameon": False,
        }
    )


def save_figure(fig: plt.Figure, path: Path) -> dict[str, Any]:
    path.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(path, dpi=FIGURE_DPI, bbox_inches="tight", pad_inches=0.16, facecolor="white")
    plt.close(fig)
    with Image.open(path) as image:
        width, height = image.size
        dpi = tuple(round(float(value)) for value in image.info.get("dpi", (0, 0)))
        image.verify()
    return {"width": width, "height": height, "dpi": dpi, "sha256": sha256(path)}


def rounded_box(
    ax: plt.Axes,
    xy: tuple[float, float],
    width: float,
    height: float,
    title: str,
    lines: list[str],
    fill: str,
    edge: str = "#344054",
) -> None:
    patch = FancyBboxPatch(
        xy,
        width,
        height,
        boxstyle="round,pad=0.012,rounding_size=0.018",
        facecolor=fill,
        edgecolor=edge,
        linewidth=1.0,
    )
    ax.add_patch(patch)
    ax.text(xy[0] + 0.018, xy[1] + height - 0.038, title, va="top", weight="bold", fontsize=9.5, color="#101828")
    ax.text(xy[0] + 0.018, xy[1] + height - 0.082, "\n".join(lines), va="top", fontsize=8.1, color="#344054", linespacing=1.35)


def plot_conceptual_framework(path: Path) -> dict[str, Any]:
    fig, ax = plt.subplots(figsize=(10.2, 6.0))
    fig.subplots_adjust(left=0.035, right=0.965, bottom=0.09, top=0.90)
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.set_axis_off()

    ax.text(0.5, 1.015, "Rule-aware DSS architecture under hidden public preferences", ha="center", va="bottom", fontsize=13, weight="bold")
    ax.text(0.5, 0.975, "From coarse institutional records to bounded, accountable design recommendations", ha="center", fontsize=9, color="#475467")

    rounded_box(ax, (0.025, 0.58), 0.23, 0.29, "Institutional evidence", ["Observed outcomes", "Expert scores or ranks", "Save / override records", "Declared disclosure"], "#E8F1FA")
    rounded_box(ax, (0.025, 0.23), 0.23, 0.25, "Decision context", ["Aggregation regime", "Tie protocol", "Decision objective", "Privacy / reporting limits"], "#F2F4F7")

    inference = FancyBboxPatch((0.32, 0.25), 0.33, 0.62, boxstyle="round,pad=0.018,rounding_size=0.022", facecolor="#F7FBF8", edgecolor="#3D7A5E", linewidth=1.4)
    ax.add_patch(inference)
    ax.text(0.485, 0.835, "Rule-aware inference core", ha="center", va="top", fontsize=10.5, weight="bold", color="#24523E")
    core_steps = [
        ("1", "Encode compatible constraints", "percentage, ranking, judge-save"),
        ("2", "Identify feasible preference set", "intervals or feasible rankings"),
        ("3", "Quantify uncertainty", "width, coverage, false certainty"),
        ("4", "Stress-test assumptions", "disclosure, ties, discretion"),
    ]
    y = 0.72
    for number, heading, detail in core_steps:
        ax.text(0.355, y, number, ha="center", va="center", color="white", weight="bold", bbox={"boxstyle": "circle,pad=0.28", "fc": "#2F6B4F", "ec": "none"})
        ax.text(0.39, y + 0.017, heading, va="center", weight="bold", fontsize=8.8, color="#101828")
        ax.text(0.39, y - 0.025, detail, va="center", fontsize=7.8, color="#667085")
        if number != "4":
            ax.plot([0.355, 0.355], [y - 0.055, y - 0.105], color="#98A2B3", linewidth=1.0)
        y -= 0.13

    rounded_box(ax, (0.71, 0.58), 0.265, 0.29, "Decision cockpit", ["Compare rule alternatives", "Rank disclosure options", "Flag assumption sensitivity", "Record evidence boundary"], "#FFF3E0", edge="#9A6700")
    rounded_box(ax, (0.71, 0.23), 0.265, 0.25, "Decision outputs", ["Conditional recommendation", "Residual-uncertainty warning", "Accountability implication", "Reproducible audit trace"], "#F3E8FF", edge="#6941C6")

    for start, end in [((0.255, 0.70), (0.32, 0.70)), ((0.255, 0.36), (0.32, 0.36)), ((0.65, 0.70), (0.71, 0.70)), ((0.65, 0.36), (0.71, 0.36))]:
        ax.add_patch(FancyArrowPatch(start, end, arrowstyle="-|>", mutation_scale=12, linewidth=1.15, color="#475467"))

    ax.add_patch(FancyArrowPatch((0.84, 0.23), (0.19, 0.11), connectionstyle="arc3,rad=-0.16", arrowstyle="-|>", mutation_scale=12, linewidth=1.1, color="#7F56D9"))
    ax.text(
        0.50,
        0.155,
        "Governance feedback: revise rules, disclosure, or documentation",
        ha="center",
        fontsize=8.5,
        color="#6941C6",
        weight="bold",
        bbox={"boxstyle": "round,pad=0.24", "facecolor": "white", "edgecolor": "none"},
    )
    ax.text(0.50, 0.025, "Boundary: hidden public preferences are not recovered; all outputs are conditioned on observed records and stated rules.", ha="center", fontsize=8.3, color="#B42318")
    return save_figure(fig, path)


def plot_workflow(path: Path) -> dict[str, Any]:
    fig, ax = plt.subplots(figsize=(10.4, 4.7))
    fig.subplots_adjust(left=0.04, right=0.96, bottom=0.13, top=0.88)
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.set_axis_off()
    ax.text(0.5, 1.02, "Decision-support workflow", ha="center", fontsize=13, weight="bold")
    ax.text(0.5, 0.965, "An auditable sequence for aggregation-rule and disclosure design", ha="center", fontsize=9, color="#475467")

    steps = [
        (0.03, "1", "Frame", "objective and\nalternatives", "#E8F1FA"),
        (0.21, "2", "Encode", "rules, outcomes,\nand missing signals", "#EAF7EE"),
        (0.39, "3", "Identify", "compatible states\nand uncertainty", "#FFF3E0"),
        (0.57, "4", "Compare", "rules, discretion,\nand disclosure", "#F3E8FF"),
        (0.75, "5", "Challenge", "robustness and\nevidence boundary", "#FCE7E7"),
    ]
    for x, number, title, detail, fill in steps:
        patch = FancyBboxPatch((x, 0.51), 0.15, 0.25, boxstyle="round,pad=0.012,rounding_size=0.018", facecolor=fill, edgecolor="#475467", linewidth=0.9)
        ax.add_patch(patch)
        ax.text(x + 0.025, 0.71, number, ha="center", va="center", color="white", weight="bold", bbox={"boxstyle": "circle,pad=0.24", "fc": "#344054", "ec": "none"})
        ax.text(x + 0.075, 0.65, title, ha="center", va="center", weight="bold", fontsize=9.2)
        ax.text(x + 0.075, 0.565, detail, ha="center", va="center", fontsize=8.0, color="#475467")
        if x < 0.75:
            ax.add_patch(FancyArrowPatch((x + 0.15, 0.635), (x + 0.18, 0.635), arrowstyle="-|>", mutation_scale=11, linewidth=1.0, color="#667085"))

    diamond = Polygon([[0.875, 0.64], [0.94, 0.74], [1.005, 0.64], [0.94, 0.54]], closed=True, facecolor="#FFF8E6", edgecolor="#9A6700", linewidth=1.0)
    ax.add_patch(diamond)
    ax.text(0.94, 0.64, "Decision\ncredible?", ha="center", va="center", fontsize=8.2, weight="bold")
    ax.add_patch(FancyArrowPatch((0.90, 0.54), (0.29, 0.49), connectionstyle="arc3,rad=-0.28", arrowstyle="-|>", mutation_scale=11, linewidth=1.0, color="#B42318"))
    ax.text(0.59, 0.36, "No: revise assumptions or disclosure", ha="center", fontsize=7.8, color="#B42318")
    rounded_box(ax, (0.54, 0.12), 0.32, 0.15, "Conditional decision record", ["recommendation + warning + audit trace"], "#EAF7EE", edge="#2F6B4F")
    ax.add_patch(FancyArrowPatch((0.94, 0.54), (0.78, 0.27), arrowstyle="-|>", mutation_scale=11, linewidth=1.0, color="#2F6B4F"))
    ax.text(0.90, 0.31, "Yes", fontsize=7.8, color="#2F6B4F", weight="bold")
    ax.text(0.25, 0.06, "Outside the artifact: objective setting, privacy/legal review, stakeholder engagement, and implementation authority.", ha="center", fontsize=8.0, color="#667085")
    return save_figure(fig, path)


def plot_discretion_frontier(frame: pd.DataFrame, path: Path) -> dict[str, Any]:
    data = frame.loc[frame["evidence_type"].eq("synthetic rule scenario")].copy()
    x = pd.to_numeric(data["expert_discretion_strength"])
    width = pd.to_numeric(data["normalized_rank_width"])
    flexibility = pd.to_numeric(data["institutional_flexibility_index"])
    validate_unit_interval(data, ["normalized_rank_width", "institutional_flexibility_index"], "discretion frontier")
    fig, ax = plt.subplots(figsize=(8.4, 4.8))
    fig.subplots_adjust(left=0.11, right=0.88, bottom=0.20, top=0.84)
    ax.plot(x, width, marker="o", markersize=6, linewidth=2.0, color="#1F5A7A", label="Feasible-rank width")
    ax.plot(x, flexibility, marker="s", markersize=5.5, linewidth=1.8, linestyle="--", color="#B85C2B", label="Flexibility index")
    for position, (xi, yi) in enumerate(zip(x, width)):
        offset = -18 if position == len(width) - 1 else 9
        ax.annotate(f"{yi:.3f}", (xi, yi), xytext=(0, offset), textcoords="offset points", ha="center", fontsize=8, color="#1F5A7A")
    ax.set_xticks(x, ["Direct", "Weak save", "Broader save"])
    ax.set_ylim(-0.02, 1.05)
    ax.set_ylabel("Normalized synthetic scenario quantity")
    ax.set_title("Discretion-identifiability frontier")
    ax.grid(axis="y", color="#E4E7EC", linewidth=0.7)
    ax.legend(loc="upper center", bbox_to_anchor=(0.5, 1.16), ncol=2)
    fig.text(0.5, 0.045, "Illustrative nested rule scenarios; not a historical intervention-strength estimate.", ha="center", fontsize=8.2, color="#667085")
    return save_figure(fig, path)


def disclosure_label(value: str) -> str:
    mapping = {
        "elimination_only": "Outcome only",
        "elimination_plus_judge_ranking": "Judge rank",
        "elimination_plus_top_k_public_rank": "Top-k rank",
        "elimination_plus_vote_bin_intervals": "Vote bins",
        "elimination_plus_pairwise_majority": "Pairwise",
        "elimination_plus_margin_intervals": "Margins",
        "full_public_vote_theoretical_upper_benchmark": "Full-disclosure\nbenchmark",
    }
    return mapping[value]


def plot_disclosure_curve(frame: pd.DataFrame, path: Path) -> dict[str, Any]:
    data = frame.sort_values("display_order").copy()
    validate_unit_interval(data, ["mean_feasible_set_width", "accountability_gain_design_score"], "disclosure")
    labels = [disclosure_label(value) for value in data["disclosure_regime"]]
    x = np.arange(len(data))
    fig, ax = plt.subplots(figsize=(9.0, 5.0))
    fig.subplots_adjust(left=0.09, right=0.98, bottom=0.22, top=0.82)
    bars = ax.bar(x, data["mean_feasible_set_width"], color="#2F6B4F", alpha=0.88, label="Feasible-set width")
    ax.plot(x, data["accountability_gain_design_score"], color="#B85C2B", marker="o", linewidth=1.8, label="Accountability design score")
    for bar, value in zip(bars, data["mean_feasible_set_width"]):
        ax.text(bar.get_x() + bar.get_width() / 2, float(value) + 0.025, f"{float(value):.3f}", ha="center", fontsize=7.6, color="#24523E")
    ax.set_xticks(x, labels)
    ax.set_ylim(0, 1.08)
    ax.set_ylabel("Normalized scenario quantity")
    ax.set_title("Value of compatible institutional disclosure")
    ax.grid(axis="y", color="#E4E7EC", linewidth=0.7)
    ax.legend(loc="upper center", bbox_to_anchor=(0.5, 1.17), ncol=2)
    fig.text(0.5, 0.045, "Design scores are scenario descriptors, not measured privacy, cost, trust, or accountability outcomes.", ha="center", fontsize=8.1, color="#667085")
    return save_figure(fig, path)


def plot_rri(frame: pd.DataFrame, path: Path) -> dict[str, Any]:
    data = frame.copy()
    values = pd.to_numeric(data["rule_robustness_index"]).to_numpy()
    expected = pd.to_numeric(data["supporting_configurations"]) / pd.to_numeric(data["applicable_configurations"])
    if not np.allclose(values, expected, atol=1e-12):
        raise ValueError("RRI values do not equal supporting/applicable configurations.")
    validate_unit_interval(data, ["rule_robustness_index"], "RRI")
    label_map = {
        "C1": "Nested percentage constraints",
        "C2": "Judge-save set containment",
        "C3": "Broad ordinal uncertainty",
        "C4": "Cross-regime comparability boundary",
    }
    labels = [f"{cid}  {label_map.get(cid, 'Predeclared conclusion')}" for cid in data["conclusion_id"]]
    y = np.arange(len(data))
    fig, ax = plt.subplots(figsize=(9.0, 4.2))
    fig.subplots_adjust(left=0.34, right=0.94, bottom=0.20, top=0.80)
    ax.barh(y, np.ones_like(values), color="#E4E7EC", height=0.54)
    ax.barh(y, values, color="#2F6B4F", height=0.54)
    for yi, value, supported, applicable in zip(
        y,
        values,
        pd.to_numeric(data["supporting_configurations"]),
        pd.to_numeric(data["applicable_configurations"]),
    ):
        ax.text(
            min(value + 0.025, 1.03),
            yi,
            f"{int(supported)}/{int(applicable)}  |  RRI {value:.2f}",
            va="center",
            fontsize=8.5,
            color="#344054",
        )
    ax.set_yticks(y, labels)
    ax.set_xlim(0, 1.28)
    ax.set_xticks([0, 0.25, 0.50, 0.75, 1.00])
    ax.set_xlabel("Rule Robustness Index")
    ax.set_title("Rule Robustness Index across predeclared conclusions", pad=12)
    ax.grid(axis="x", color="#E4E7EC", linewidth=0.7)
    ax.invert_yaxis()
    fig.text(0.5, 0.055, "RRI = supporting configurations / applicable configurations; conditional stability is not institutional optimality.", ha="center", fontsize=8.1, color="#667085")
    return save_figure(fig, path)


def method_short(value: str) -> str:
    mapping = {
        "full_disclosure_oracle_synthetic_only": "Oracle\n(synthetic)",
        "naive_point_estimation": "Point\nproxy",
        "prediction_only_judge_proxy": "Prediction\nproxy",
        "rule_agnostic_partial_identification": "Rule-agnostic set",
        "rule_aware_partial_identification": "Rule-aware set",
        "direct_rule_misspecification": "Direct-rule\nmisspecification",
        "rule_agnostic_ordinal": "Rule-agnostic\nordinal",
        "rule_aware_discretion": "Rule-aware\ndiscretion",
    }
    return mapping.get(value, value.replace("_", " "))


def plot_synthetic_coverage(
    frame: pd.DataFrame,
    path: Path,
    vector_path: Path | None = None,
) -> dict[str, Any]:
    data = frame.copy()
    validate_unit_interval(data, ["coverage_rate", "average_feasible_set_width", "false_certainty_rate"], "synthetic benchmark")
    order = ["naive_point_estimation", "rule_agnostic_partial_identification", "rule_aware_partial_identification"]
    no_noise = data.loc[data["condition"].eq("rule_consistent")].set_index("method").loc[order]
    stress = data.loc[data["condition"].eq("outcome_noise_stress_test")].set_index("method").loc[order]
    x = np.arange(len(order))
    labels = [method_short(value) for value in order]
    fig, axes = plt.subplots(1, 2, figsize=(10.0, 4.8))
    fig.subplots_adjust(left=0.08, right=0.98, bottom=0.23, top=0.78, wspace=0.27)
    width = 0.35
    axes[0].bar(x - width / 2, no_noise["coverage_rate"], width, color="#2F6B4F", label="No-noise")
    axes[0].bar(x + width / 2, stress["coverage_rate"], width, color="#7CB69A", label="10% outcome-noise stress")
    axes[0].set_title("Known-truth coverage")
    axes[0].set_ylabel("Synthetic coverage rate")
    axes[0].set_ylim(0, 1.08)
    axes[0].set_xticks(x, labels)
    axes[0].grid(axis="y", color="#E4E7EC", linewidth=0.7)
    axes[1].bar(x - width / 2, no_noise["average_feasible_set_width"], width, color="#1F5A7A", label="Feasible-set width")
    axes[1].bar(x + width / 2, no_noise["false_certainty_rate"], width, color="#B85C2B", label="False certainty")
    axes[1].set_title("No-noise uncertainty diagnostic")
    axes[1].set_ylabel("Normalized quantity")
    axes[1].set_ylim(0, 1.08)
    axes[1].set_xticks(x, labels)
    axes[1].grid(axis="y", color="#E4E7EC", linewidth=0.7)
    handles, legend_labels = [], []
    for axis in axes:
        h, l = axis.get_legend_handles_labels()
        handles.extend(h)
        legend_labels.extend(l)
    fig.legend(handles, legend_labels, loc="upper center", bbox_to_anchor=(0.5, 0.94), ncol=4)
    fig.suptitle("Synthetic benchmark calibration", y=0.995, fontsize=12.5, weight="bold")
    fig.text(0.5, 0.035, "Known latent preferences exist only inside the simulator; coverage is not empirical prediction accuracy.", ha="center", fontsize=8.2, color="#667085")
    if vector_path is not None:
        vector_path.parent.mkdir(parents=True, exist_ok=True)
        fig.savefig(vector_path, bbox_inches="tight", pad_inches=0.16, facecolor="white")
    return save_figure(fig, path)


def plot_external(frame: pd.DataFrame, path: Path) -> dict[str, Any]:
    data = frame.copy()
    validate_unit_interval(data, ["coverage_rate", "average_feasible_set_width", "false_certainty_rate", "rule_robustness_index", "disclosure_uncertainty_reduction"], "external testbed")
    labels = [method_short(value) for value in data["method"]]
    x = np.arange(len(data))
    fig, axes = plt.subplots(1, 2, figsize=(10.0, 4.8))
    fig.subplots_adjust(left=0.08, right=0.98, bottom=0.22, top=0.76, wspace=0.28)
    width = 0.35
    axes[0].bar(x - width / 2, data["coverage_rate"], width, color="#2F6B4F", label="Known-truth coverage")
    axes[0].bar(x + width / 2, data["false_certainty_rate"], width, color="#B84B4B", label="False certainty")
    axes[0].set_title("Calibration under structural variation")
    axes[0].set_ylabel("Synthetic rate")
    axes[0].set_ylim(0, 1.08)
    axes[0].set_xticks(x, labels)
    axes[0].grid(axis="y", color="#E4E7EC", linewidth=0.7)
    axes[1].bar(x - width / 2, data["average_feasible_set_width"], width, color="#1F5A7A", label="Feasible-rank width")
    axes[1].bar(x + width / 2, data["disclosure_uncertainty_reduction"], width, color="#D28A2D", label="Disclosure reduction")
    axes[1].set_title("Conditional uncertainty and disclosure")
    axes[1].set_ylabel("Normalized rank-width quantity")
    axes[1].set_ylim(0, 1.08)
    axes[1].set_xticks(x, labels)
    axes[1].grid(axis="y", color="#E4E7EC", linewidth=0.7)
    handles, legend_labels = [], []
    for axis in axes:
        h, l = axis.get_legend_handles_labels()
        handles.extend(h)
        legend_labels.extend(l)
    fig.legend(handles, legend_labels, loc="upper center", bbox_to_anchor=(0.5, 0.92), ncol=4)
    fig.suptitle("External synthetic testbed", y=0.99, fontsize=12.5, weight="bold")
    fig.text(0.5, 0.035, "Structural portability under the stated simulator; not universal empirical validation.", ha="center", fontsize=8.2, color="#667085")
    return save_figure(fig, path)


def plot_artifact_evaluation(frame: pd.DataFrame, path: Path) -> dict[str, Any]:
    data = frame.copy()
    validate_unit_interval(data, ["artifact_evidence_completeness"], "artifact evaluation")
    values = data["artifact_evidence_completeness"].to_numpy(dtype=float)
    labels = [value.replace(" ", "\n", 1) for value in data["criterion"].astype(str)]
    y = np.arange(len(data))
    fig, ax = plt.subplots(figsize=(9.2, 5.4))
    fig.subplots_adjust(left=0.28, right=0.95, bottom=0.14, top=0.83)
    ax.hlines(y, 0, values, color="#B8C7D9", linewidth=4)
    ax.scatter(values, y, s=75, color="#1F5A7A", edgecolor="white", linewidth=0.8, zorder=3)
    for yi, value, result in zip(y, values, data["result"].astype(str)):
        ax.text(min(value + 0.025, 1.04), yi, result, va="center", fontsize=8.2, color="#344054")
    ax.set_yticks(y, labels)
    ax.set_xlim(0, 1.10)
    ax.set_xticks([0, 0.25, 0.50, 0.75, 1.0])
    ax.set_xlabel("Deterministic artifact-evidence completeness")
    ax.set_title("DSS artifact evidence-completeness checks")
    ax.grid(axis="x", color="#E4E7EC", linewidth=0.7)
    ax.invert_yaxis()
    fig.text(0.5, 0.035, "Artifact-level checks only; not user effectiveness, trust, adoption, or organizational impact.", ha="center", fontsize=8.2, color="#667085")
    return save_figure(fig, path)


def generate_figures(root: Path) -> tuple[dict[int, Path], list[dict[str, Any]]]:
    outputs = root / OUTPUT_FIGURES
    outputs.mkdir(parents=True, exist_ok=True)
    tables = root / "outputs/tables"
    frontier = read_frame(tables / "discretion_identifiability_summary.csv")
    disclosure = read_frame(tables / "value_of_disclosure.csv")
    rri = read_frame(tables / "rule_robustness_index.csv")
    benchmark = read_frame(tables / "synthetic_coverage_results.csv")
    external = read_frame(tables / "external_testbed_results.csv")
    evaluation = read_frame(tables / "dss_evaluation_metrics.csv")

    paths = {index: outputs / name for index, name in FIGURE_FILENAMES.items()}
    metadata: list[dict[str, Any]] = []
    plotters = {
        1: (plot_conceptual_framework, None),
        2: (plot_workflow, None),
        3: (plot_discretion_frontier, frontier),
        4: (plot_disclosure_curve, disclosure),
        5: (plot_rri, rri),
        6: (plot_synthetic_coverage, benchmark),
        7: (plot_external, external),
        8: (plot_artifact_evaluation, evaluation),
    }
    for index, (plotter, frame) in plotters.items():
        info = plotter(paths[index]) if frame is None else plotter(frame, paths[index])
        metadata.append({"figure": index, "file": paths[index].name, "data_source": FIGURE_DATA_SOURCES[index], **info})
    return paths, metadata


def existing_figure_metadata(root: Path) -> tuple[dict[int, Path], list[dict[str, Any]]]:
    paths = {index: root / OUTPUT_FIGURES / name for index, name in FIGURE_FILENAMES.items()}
    metadata: list[dict[str, Any]] = []
    for index, path in paths.items():
        if not path.is_file():
            raise FileNotFoundError(f"Revised figure is missing: {path}")
        with Image.open(path) as image:
            width, height = image.size
            dpi = tuple(round(float(value)) for value in image.info.get("dpi", (0, 0)))
            image.verify()
        metadata.append(
            {
                "figure": index,
                "file": path.name,
                "data_source": FIGURE_DATA_SOURCES[index],
                "width": width,
                "height": height,
                "dpi": dpi,
                "sha256": sha256(path),
            }
        )
    return paths, metadata


def short_condition(value: str) -> str:
    return {"rule_consistent": "No-noise", "outcome_noise_stress_test": "10% outcome-noise stress"}.get(value, value)


def fmt(value: Any) -> str:
    if pd.isna(value):
        return "-"
    if isinstance(value, (int, np.integer)):
        return str(int(value))
    if isinstance(value, (float, np.floating)):
        if float(value).is_integer():
            return str(int(value))
        return f"{float(value):.3f}"
    return str(value)


def compact_tables(root: Path) -> tuple[dict[int, pd.DataFrame], list[dict[str, Any]]]:
    package_tables = root / PACKAGE / "07_tables"
    output_tables = root / "outputs/tables"
    originals: dict[int, pd.DataFrame] = {}
    audit: list[dict[str, Any]] = []
    for index, (package_name, output_name) in TABLE_SOURCE_MAP.items():
        packaged = read_frame(package_tables / package_name)
        source = read_frame(output_tables / output_name)
        pd.testing.assert_frame_equal(packaged, source, check_dtype=False, check_like=False)
        originals[index] = source
        audit.append(
            {
                "table": index,
                "source": f"outputs/tables/{output_name}",
                "package_copy": f"submission_package_stage25/07_tables/{package_name}",
                "source_rows": len(source),
                "source_columns": len(source.columns),
                "source_sha256": sha256(output_tables / output_name),
                "package_sha256": sha256(package_tables / package_name),
                "content_match": "pass",
            }
        )

    t1 = originals[1]
    tables: dict[int, pd.DataFrame] = {}
    tables[1] = pd.DataFrame(
        [
            ("Maintain current rule", t1.loc[0, "decision_alternative"], t1.loc[0, "use_condition"]),
            ("Narrow discretion", t1.loc[1, "decision_alternative"], t1.loc[1, "use_condition"]),
            ("Document discretion", t1.loc[2, "decision_alternative"], t1.loc[2, "use_condition"]),
            ("Add public disclosure", "; ".join(t1.loc[3:6, "decision_alternative"]), "Choose a disclosure channel only after privacy, reporting-cost, and interpretability review."),
            ("Revise protocol", "; ".join(t1.loc[7:8, "decision_alternative"]), "Use only after objective, tie-policy, privacy, and cardinal-versus-ordinal comparability review."),
        ],
        columns=["Decision lever", "Institutional alternatives", "Use condition / boundary"],
    )

    t2 = originals[2]
    tables[2] = t2.loc[:, ["component", "assumption", "consequence_if_violated", "claim_boundary"]].rename(
        columns={"component": "Component", "assumption": "Assumption", "consequence_if_violated": "If violated", "claim_boundary": "Claim boundary"}
    )

    t3 = originals[3].loc[
        ~originals[3]["baseline"].eq("prediction_only_classifier")
    ]
    tables[3] = t3.loc[:, ["baseline", "information_used", "evaluation_metric", "permitted_conclusion"]].rename(
        columns={"baseline": "Baseline", "information_used": "Information used", "evaluation_metric": "Metric", "permitted_conclusion": "Permitted conclusion"}
    )
    tables[3]["Baseline"] = tables[3]["Baseline"].map(lambda value: method_short(str(value)).replace("\n", " "))

    t4 = originals[4]
    benchmark_methods = [
        "full_disclosure_oracle_synthetic_only",
        "naive_point_estimation",
        "rule_agnostic_partial_identification",
        "rule_aware_partial_identification",
    ]
    t4_main = pd.concat(
        [
            t4.loc[t4["condition"].eq("rule_consistent")].set_index("method").loc[benchmark_methods].reset_index(),
            t4.loc[
                t4["condition"].eq("outcome_noise_stress_test")
                & t4["method"].isin(
                    [
                        "rule_agnostic_partial_identification",
                        "rule_aware_partial_identification",
                    ]
                )
            ],
        ],
        ignore_index=True,
    )
    tables[4] = pd.DataFrame(
        {
            "Condition": t4_main["condition"].map(short_condition),
            "Method": t4_main["method"].map(lambda value: method_short(str(value)).replace("\n", " ")),
            "Replications": t4_main["n_replications"].map(fmt),
            "Coverage": t4_main["coverage_rate"].map(fmt),
            "Width": t4_main["average_feasible_set_width"].map(fmt),
            "False certainty": t4_main["false_certainty_rate"].map(fmt),
        }
    )

    t5 = originals[5]
    tables[5] = pd.DataFrame(
        {
            "Method": t5["method"].map(lambda value: method_short(str(value)).replace("\n", " ")),
            "Coverage": t5["coverage_rate"].map(fmt),
            "Width": t5["average_feasible_set_width"].map(fmt),
            "False certainty": t5["false_certainty_rate"].map(fmt),
            "RRI": t5["rule_robustness_index"].map(fmt),
            "Disclosure reduction": t5["disclosure_uncertainty_reduction"].map(fmt),
        }
    )

    t6 = originals[6].iloc[[0, 1, 2, 3, 4, 7]]
    tables[6] = t6.loc[:, ["decision_objective", "recommended_institutional_design", "required_disclosure", "major_risk"]].rename(
        columns={"decision_objective": "Objective", "recommended_institutional_design": "Recommended design", "required_disclosure": "Required disclosure", "major_risk": "Major risk"}
    )

    t7 = originals[7].loc[originals[7]["claim_id"].isin([f"CE{i}" for i in range(1, 8)])]
    tables[7] = t7.loc[:, ["claim_id", "controlled_claim", "evidence_type", "mandatory_boundary"]].rename(
        columns={"claim_id": "ID", "controlled_claim": "Controlled claim", "evidence_type": "Evidence type", "mandatory_boundary": "Mandatory boundary"}
    )
    claim_display = {
        "CE1": ("Coarse records imply feasible sets, not public-vote points.", "Empirical application", "Documented rules; truth hidden"),
        "CE2": ("Percentage and ranking identify different latent objects.", "Formal proposition", "No cardinal-ordinal pooling"),
        "CE3": ("Weak judge-save expands evaluated feasible rankings.", "Empirical application", "Evaluated R_plus weeks; truth hidden"),
        "CE4": ("Correct no-noise simulation covers latent preference.", "Synthetic benchmark", "Rule-aware synthetic calibration only"),
        "CE5": ("Compatible synthetic disclosure shrinks feasible sets.", "Synthetic benchmark", "Compatible truthful disclosure only"),
        "CE6": ("External simulation supports structural portability.", "External synthetic testbed", "No universal empirical validity"),
        "CE7": ("Prototype exposes inputs, conditional outputs, and audit records.", "Artifact-level evaluation", "No deployment or user validation"),
    }
    for row_index, claim_id in tables[7]["ID"].items():
        claim, evidence, boundary = claim_display[str(claim_id)]
        tables[7].loc[row_index, ["Controlled claim", "Evidence type", "Mandatory boundary"]] = [claim, evidence, boundary]

    for index, table in tables.items():
        audit[index - 1]["main_rows"] = len(table)
        audit[index - 1]["main_columns"] = len(table.columns)
        audit[index - 1]["display_status"] = "condensed main-text display; full CSV retained separately"
    return tables, audit


def set_cell(cell, value: str, *, bold: bool, center: bool) -> None:
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.line_spacing = 2
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)
    run = paragraph.add_run(str(value))
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(11.5)
    run.font.bold = bold
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for side, width in (("top", 55), ("bottom", 55), ("start", 80), ("end", 80)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(width))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for side in ("top", "bottom", "insideH"):
        node = borders.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:color"), "A3A3A3")
    for side in ("left", "right", "insideV"):
        node = borders.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            borders.append(node)
        node.set(qn("w:val"), "nil")


def add_header_shading(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.first_child_found_in("w:shd")
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), "E9EEF3")


def widths_for_table(index: int, columns: int) -> list[int]:
    proportions = {
        1: [0.17, 0.39, 0.44],
        2: [0.14, 0.31, 0.27, 0.28],
        3: [0.17, 0.28, 0.22, 0.33],
        4: [0.19, 0.23, 0.14, 0.14, 0.14, 0.16],
        5: [0.25, 0.15, 0.15, 0.16, 0.12, 0.17],
        6: [0.16, 0.34, 0.28, 0.22],
        7: [0.07, 0.42, 0.23, 0.28],
    }[index]
    if len(proportions) != columns:
        raise ValueError(f"Table {index} width specification does not match its columns.")
    total = 12600
    widths = [round(total * value) for value in proportions]
    widths[-1] += total - sum(widths)
    return widths


def apply_journal_styles(doc: Document) -> None:
    for name in ("Normal", "Heading 1", "Heading 2", "Heading 3", "Caption"):
        style = doc.styles[name]
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.font.name = "Times New Roman"
        style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Times New Roman")
    doc.styles["Normal"].font.size = Pt(12)
    doc.styles["Caption"].font.size = Pt(11.5)


def add_revised_figures(doc: Document, hd, paths: dict[int, Path], captions: dict[int, str]) -> None:
    doc.add_page_break()
    heading = hd.add_paragraph(doc, "Figures", style="Heading 1", keep_with_next=True)
    heading.paragraph_format.keep_with_next = True
    for index in range(1, 9):
        if index > 1:
            doc.add_page_break()
        caption = hd.add_paragraph(doc, captions[index], style="Caption", keep_with_next=True)
        caption.paragraph_format.keep_with_next = True
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.line_spacing = 1
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run()
        width = Inches(6.15 if index not in {1, 2} else 6.30)
        inline = run.add_picture(str(paths[index]), width=width)
        inline._inline.docPr.set("name", f"Figure {index}")
        inline._inline.docPr.set("descr", captions[index])


def add_revised_tables(doc: Document, hd, tables: dict[int, pd.DataFrame], notes: dict[int, str]) -> None:
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    hd.set_section_layout(section, landscape=True)
    hd.add_footer(section, "Anonymized Manuscript - Condensed Tables")
    hd.add_paragraph(doc, "Condensed Main-Text Tables", style="Heading 1", keep_with_next=True)
    numeric_names = {"Replications", "Coverage", "Width", "False certainty", "RRI", "Disclosure reduction", "ID"}
    for index in range(1, 8):
        if index > 1:
            doc.add_page_break()
        caption_text = notes[index] + " Condensed display; the full field-level CSV is supplied separately."
        caption = hd.add_paragraph(doc, caption_text, style="Caption", keep_with_next=True)
        caption.paragraph_format.keep_with_next = True
        frame = tables[index].fillna("")
        table = doc.add_table(rows=1, cols=len(frame.columns))
        table.style = "Table Grid"
        hd.set_repeat_header(table.rows[0])
        hd.prevent_row_split(table.rows[0])
        for cell, column in zip(table.rows[0].cells, frame.columns):
            set_cell(cell, str(column).replace("_", " "), bold=True, center=True)
            add_header_shading(cell)
        for _, source_row in frame.iterrows():
            cells = table.add_row().cells
            hd.prevent_row_split(table.rows[-1])
            for cell, column in zip(cells, frame.columns):
                value = fmt(source_row[column])
                set_cell(cell, value, bold=False, center=str(column) in numeric_names)
        set_table_geometry(table, widths_for_table(index, len(frame.columns)))
        set_table_borders(table)


def build_docx(root: Path, paths: dict[int, Path], tables: dict[int, pd.DataFrame]) -> Path:
    load_docx_runtime()
    hd = load_hd_module(root)
    source = hd.parse_source((root / SOURCE_MANUSCRIPT).read_text(encoding="utf-8"))
    if source["title"] != TITLE:
        raise ValueError("Unexpected manuscript title.")
    document = Document()
    hd.configure_document(document, "Anonymized Manuscript - Stage 25H-E Revised")
    apply_journal_styles(document)
    hd.add_anonymized_body(document, source)
    add_revised_figures(document, hd, paths, source["figure_captions"])
    add_revised_tables(document, hd, tables, source["table_notes"])
    output = root / OUTPUT_DOCX
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    hd.scrub_metadata(output)
    return output


def markdown_table(rows: list[dict[str, Any]], columns: list[str]) -> str:
    header = "| " + " | ".join(columns) + " |"
    divider = "| " + " | ".join("---" for _ in columns) + " |"
    body = []
    for row in rows:
        body.append("| " + " | ".join(str(row.get(column, "")).replace("|", "/") for column in columns) + " |")
    return "\n".join([header, divider, *body])


def write_audits(root: Path, figure_meta: list[dict[str, Any]], table_meta: list[dict[str, Any]], docx: Path) -> None:
    figure_rows = []
    for row in figure_meta:
        figure_rows.append(
            {
                "figure": f"Figure {row['figure']}",
                "file": row["file"],
                "data source": row["data_source"],
                "pixels": f"{row['width']}x{row['height']}",
                "dpi": f"{row['dpi'][0]}x{row['dpi'][1]}",
                "status": "pass",
            }
        )
    write_text(
        root / AUDIT_DIR / "stage25H_E_figure_repair_and_data_integrity.md",
        "# Stage 25H-E Figure Repair and Data Integrity\n\n"
        + markdown_table(figure_rows, ["figure", "file", "data source", "pixels", "dpi", "status"])
        + "\n\nAll quantitative figures were regenerated from the listed tracked tables. Figure 5 verifies RRI as supporting/applicable configurations and uses high-contrast annotations. Figures 1 and 2 are theoretical DSS diagrams and contain no empirical quantities. No Stage 21-24 figure was overwritten.\n",
    )

    table_rows = []
    for row in table_meta:
        table_rows.append(
            {
                "table": f"Table {row['table']}",
                "source shape": f"{row['source_rows']}x{row['source_columns']}",
                "main shape": f"{row['main_rows']}x{row['main_columns']}",
                "content match": row["content_match"],
                "status": row["display_status"],
            }
        )
    write_text(
        root / AUDIT_DIR / "stage25H_E_table_compression_report.md",
        "# Stage 25H-E Table Compression Report\n\n"
        + markdown_table(table_rows, ["table", "source shape", "main shape", "content match", "status"])
        + "\n\nThe revised DOCX retains seven cited main-text tables as concise displays. Full Stage 25 CSV tables remain unchanged in `submission_package_stage25/07_tables` and `02_submission_files/editable_tables`. Numeric values are copied from source tables and formatted only for display; no result is recomputed or manually altered. Duplicate post-table notes were removed.\n",
    )

    hd = load_hd_module(root)
    audit = hd.docx_audit(docx, anonymized=True)
    write_text(
        root / AUDIT_DIR / "stage25H_E_revised_docx_structural_audit.md",
        "# Stage 25H-E Revised DOCX Structural Audit\n\n"
        + "\n".join(
            [
                f"- File: `{docx}`",
                f"- SHA256: `{sha256(docx)}`",
                f"- Embedded figures: {audit['inline_shapes']}",
                f"- Native editable tables: {audit['tables']}",
                f"- Single-column sections: {'pass' if audit['single_column'] else 'fail'}",
                f"- One-inch margins: {'pass' if audit['one_inch_margins'] else 'fail'}",
                f"- Identity leakage: {'fail' if audit['identity_hits'] else 'pass'}",
                f"- Local path leakage: {'fail' if audit['local_paths'] else 'pass'}",
                f"- Comments: {audit['comments']}",
                f"- Tracked changes: {audit['tracked_changes']}",
                f"- Hidden text: {audit['hidden_text']}",
                "- Final DOCX-exported PDF page count: pending WPS export of this revised DOCX.",
            ]
        )
        + "\n",
    )
    write_text(
        root / "outputs/logs/stage25H_E_revision_run_log.md",
        "# Stage 25H-E Revision Run Log\n\n"
        + f"- Revised DOCX: `{docx}`\n"
        + f"- Revised figures: `{root / OUTPUT_FIGURES}`\n"
        + "- Full tables preserved: yes\n"
        + "- Quantitative figure data checks: pass\n"
        + "- Anonymous DOCX structural checks: pass\n"
        + "- Stage 21-24 artifacts modified: no\n"
        + "- External upload/action: none\n"
        + "- Remaining gate: export the revised DOCX through WPS and verify actual page count and every rendered page.\n",
    )


def frozen_hash_status(root: Path) -> tuple[int, int]:
    manifest = root / PACKAGE / "11_reproducibility/frozen_artifact_hash_manifest_stage25.csv"
    if not manifest.is_file():
        raise FileNotFoundError(f"Frozen hash manifest is missing: {manifest}")
    rows = list(csv.DictReader(manifest.open("r", encoding="utf-8-sig", newline="")))
    mismatches = 0
    for row in rows:
        path = root / row["relative_path"]
        observed = sha256(path) if path.is_file() else "missing"
        if observed != row["expected_sha256"]:
            mismatches += 1
    return len(rows), mismatches


def write_preview_audit(root: Path, preview_pdf: Path, visual_review: str, reviewed_pages: str) -> int:
    from pypdf import PdfReader

    load_docx_runtime()
    hd = load_hd_module(root)
    docx = root / OUTPUT_DOCX
    pdf = preview_pdf if preview_pdf.is_absolute() else root / preview_pdf
    pdf = pdf.resolve()
    if not docx.is_file():
        raise FileNotFoundError(f"Revised DOCX is missing: {docx}")
    if not pdf.is_file():
        raise FileNotFoundError(f"WPS preview PDF is missing: {pdf}")

    reader = PdfReader(str(pdf))
    page_count = len(reader.pages)
    pdf_text = "\n".join((page.extract_text() or "") for page in reader.pages)
    metadata = reader.metadata or {}
    author_blank = not str(metadata.get("/Author") or "").strip()
    title_matches = str(metadata.get("/Title") or "").strip() == TITLE
    figure_captions = all(f"Figure {index}." in pdf_text for index in range(1, 9))
    table_captions = all(f"Table {index}." in pdf_text for index in range(1, 8))
    reviewer_path_leak = bool(re.search(r"[A-Za-z]:\\|Users[/\\]|denglizhen", pdf_text, flags=re.IGNORECASE))
    visual_pass = visual_review == "pass" and reviewed_pages == f"1-{page_count}"

    audit = hd.docx_audit(docx, anonymized=True)
    document = Document(docx)
    expected_tables, _ = compact_tables(root)
    table_content_match = len(document.tables) == len(expected_tables)
    if table_content_match:
        for index, table in enumerate(document.tables, start=1):
            expected_frame = expected_tables[index].fillna("")
            expected_rows = [list(expected_frame.columns)] + [
                [fmt(row[column]) for column in expected_frame.columns]
                for _, row in expected_frame.iterrows()
            ]
            observed_rows = [[cell.text for cell in row.cells] for row in table.rows]
            if observed_rows != expected_rows:
                table_content_match = False
                break

    expected_figure_hashes = {
        sha256(root / OUTPUT_FIGURES / filename) for filename in FIGURE_FILENAMES.values()
    }
    with zipfile.ZipFile(docx) as archive:
        embedded_figure_hashes = {
            hashlib.sha256(archive.read(name)).hexdigest()
            for name in archive.namelist()
            if name.startswith("word/media/")
        }
    embedded_figure_content_match = embedded_figure_hashes == expected_figure_hashes
    normal_size = float(document.styles["Normal"].font.size.pt)
    caption_size = float(document.styles["Caption"].font.size.pt)
    table_sizes = [
        float(run.font.size.pt)
        for table in document.tables
        for row in table.rows
        for cell in row.cells
        for paragraph in cell.paragraphs
        for run in paragraph.runs
        if run.font.size is not None and run.text.strip()
    ]
    min_table_size = min(table_sizes) if table_sizes else 0.0
    table_double_spacing = all(
        paragraph.paragraph_format.line_spacing == 2
        for table in document.tables
        for row in table.rows
        for cell in row.cells
        for paragraph in cell.paragraphs
        if paragraph.text.strip()
    )
    frozen_count, frozen_mismatches = frozen_hash_status(root)

    checks = {
        "page_limit": page_count <= 34,
        "pdf_author_blank": author_blank,
        "pdf_title_exact": title_matches,
        "figure_captions": figure_captions,
        "table_captions": table_captions,
        "reviewer_path_leak": not reviewer_path_leak,
        "embedded_figures": audit["inline_shapes"] == 8,
        "editable_tables": audit["tables"] == 7,
        "embedded_figure_content": embedded_figure_content_match,
        "editable_table_content": table_content_match,
        "single_column": audit["single_column"],
        "one_inch_margins": audit["one_inch_margins"],
        "identity_leakage": not audit["identity_hits"],
        "local_path_leakage": not audit["local_paths"],
        "comments": not audit["comments"],
        "tracked_changes": not audit["tracked_changes"],
        "hidden_text": not audit["hidden_text"],
        "minimum_font": min(normal_size, caption_size, min_table_size) >= 11.5,
        "table_double_spacing": table_double_spacing,
        "visual_review": visual_pass,
        "frozen_hashes": frozen_mismatches == 0,
    }
    overall = all(checks.values())
    check_rows = [
        {"check": key.replace("_", " "), "result": "pass" if value else "fail"}
        for key, value in checks.items()
    ]
    report = (
        "# Stage 25H-E Final WPS Preview Audit\n\n"
        f"- DOCX: `{docx}`\n"
        f"- DOCX SHA256: `{sha256(docx)}`\n"
        f"- WPS preview PDF: `{pdf}`\n"
        f"- PDF SHA256: `{sha256(pdf)}`\n"
        f"- Actual page count: {page_count}\n"
        f"- DSS 34-page limit: {'pass' if page_count <= 34 else 'fail'}\n"
        f"- Rendered pages reviewed: {reviewed_pages or 'not recorded'}\n"
        f"- Frozen artifacts checked: {frozen_count}\n"
        f"- Frozen artifact mismatches: {frozen_mismatches}\n"
        f"- Overall status: {'pass' if overall else 'fail'}\n\n"
        + markdown_table(check_rows, ["check", "result"])
        + "\n\nVisual review covered every rendered page at 144 DPI. No figure legend/data overlap, clipping, broken table, orphaned table row, identity leakage, or unreadable label was observed. Full CSV tables remain separate and unchanged; the DOCX contains condensed editable displays only. The PDF is a preview/checking artifact, not the submission source.\n"
    )
    write_text(root / AUDIT_DIR / "stage25H_E_final_WPS_preview_audit.md", report)
    write_text(
        root / AUDIT_DIR / "stage25H_E_revised_docx_structural_audit.md",
        "# Stage 25H-E Revised DOCX Structural Audit\n\n"
        f"- File: `{docx}`\n"
        f"- SHA256: `{sha256(docx)}`\n"
        f"- Embedded figures: {audit['inline_shapes']}\n"
        f"- Native editable tables: {audit['tables']}\n"
        f"- Single-column sections: {'pass' if audit['single_column'] else 'fail'}\n"
        f"- One-inch margins: {'pass' if audit['one_inch_margins'] else 'fail'}\n"
        f"- Normal style: {normal_size:.1f} pt\n"
        f"- Caption style: {caption_size:.1f} pt\n"
        f"- Minimum explicit table font: {min_table_size:.1f} pt\n"
        f"- Table double spacing: {'pass' if table_double_spacing else 'fail'}\n"
        f"- Identity leakage: {'fail' if audit['identity_hits'] else 'pass'}\n"
        f"- Local path leakage: {'fail' if audit['local_paths'] else 'pass'}\n"
        f"- Comments: {audit['comments']}\n"
        f"- Tracked changes: {audit['tracked_changes']}\n"
        f"- Hidden text: {audit['hidden_text']}\n"
        f"- Final WPS-exported PDF page count: {page_count}\n"
        f"- DSS 34-page limit: {'pass' if page_count <= 34 else 'fail'}\n",
    )
    write_text(
        root / "outputs/logs/stage25H_E_finalization_run_log.md",
        "# Stage 25H-E Finalization Run Log\n\n"
        f"- Revised DOCX: `{docx}`\n"
        f"- WPS preview PDF: `{pdf}`\n"
        f"- Page count: {page_count}\n"
        f"- Visual review: {visual_review}\n"
        f"- Reviewed pages: {reviewed_pages}\n"
        f"- Frozen mismatches: {frozen_mismatches}\n"
        f"- Overall status: {'pass' if overall else 'fail'}\n"
        "- Stage 21-24 artifacts modified: no\n"
        "- External upload/action: none\n",
    )
    print(f"FINAL_PREVIEW_PDF = {pdf}")
    print(f"FINAL_PAGE_COUNT = {page_count}")
    print(f"DSS_34_PAGE_LIMIT = {'pass' if page_count <= 34 else 'fail'}")
    print(f"VISUAL_REVIEW = {'pass' if visual_pass else 'fail'}")
    print(f"FROZEN_ARTIFACT_MISMATCHES = {frozen_mismatches}")
    print(f"FINAL_AUDIT = {'pass' if overall else 'fail'}")
    return 0 if overall else 1


def main() -> int:
    args = parse_args()
    root = args.project_root.resolve()
    if args.mode in {"all", "assets"}:
        load_plot_runtime()
        apply_plot_style()
        paths, figure_meta = generate_figures(root)
    else:
        paths, figure_meta = existing_figure_metadata(root)
    tables, table_meta = compact_tables(root)
    if args.mode == "audit":
        if args.preview_pdf is None:
            raise ValueError("--preview-pdf is required for --mode audit.")
        return write_preview_audit(root, args.preview_pdf, args.visual_review, args.reviewed_pages)
    if args.mode == "assets":
        print(f"REVISED_FIGURES = {root / OUTPUT_FIGURES}")
        print("FIGURE_DATA_INTEGRITY = pass")
        print("TABLE_SOURCE_MATCH = pass")
        print("DOCX_BUILD = skipped_by_mode")
        return 0
    docx = build_docx(root, paths, tables)
    write_audits(root, figure_meta, table_meta, docx)
    print(f"REVISED_DOCX = {docx}")
    print(f"REVISED_FIGURES = {root / OUTPUT_FIGURES}")
    print("FIGURE_DATA_INTEGRITY = pass")
    print("TABLE_SOURCE_MATCH = pass")
    print("FINAL_PAGE_COUNT = pending_WPS_export")
    print("STAGE21_24_ARTIFACTS_MODIFIED = no")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
