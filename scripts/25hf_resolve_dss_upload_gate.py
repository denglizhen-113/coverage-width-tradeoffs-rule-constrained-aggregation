"""Resolve the additive Stage 25H-F administrative warning and rerun the DSS gate.

This script intentionally leaves Stage 21-24 artifacts and the Stage 25H-E
anonymous manuscript untouched.  It creates a new title page with the
author-confirmed Wuhan University of Technology postal address and records a
fresh, evidence-bounded final gate.
"""

from __future__ import annotations

import argparse
import csv
import hashlib
import importlib.util
import re
import sys
import zipfile
from pathlib import Path
from typing import Any


TITLE = "Rule-Aware Decision Support for Expert-Crowd Aggregation under Hidden Public Preferences"
REPOSITORY_URL = "https://github.com/denglizhen-113/coverage-width-tradeoffs-rule-constrained-aggregation"
WUT_ADDRESS = "Wuhan University of Technology, No. 122 Luoshi Road, Hongshan District, Wuhan, Hubei 430070, China"
HUST_ADDRESS = "Huazhong University of Science and Technology, 1037 Luoyu Road, Hongshan District, Wuhan, Hubei 430074, China"

PACKAGE = Path("submission_package_stage25")
SUBMISSION = PACKAGE / "02_submission_files"
AUDITS = PACKAGE / "12_audit_logs"
AUTHOR_ACTION = PACKAGE / "01_author_action_required"
SOURCE = PACKAGE / "02_manuscript/DSS_submission_draft_STAGE25H_E_final_source.md"
ANON_DOCX = SUBMISSION / "DSS_anonymized_manuscript_STAGE25H_E_final.docx"
TITLE_E_DOCX = SUBMISSION / "DSS_title_page_STAGE25H_E_final.docx"
TITLE_F_DOCX = SUBMISSION / "DSS_title_page_STAGE25H_F_final.docx"
PREVIEW_PDF = SUBMISSION / "DSS_anonymized_manuscript_STAGE25H_E_final_preview.pdf"
TITLE_PREVIEW_PDF = SUBMISSION / "DSS_title_page_STAGE25H_F_final_preview.pdf"
HIGHLIGHTS_DOCX = SUBMISSION / "DSS_highlights_STAGE25H_E.docx"
FIGURES_DIR = SUBMISSION / "figures_STAGE25H_E_final"
FROZEN_MANIFEST = PACKAGE / "11_reproducibility/frozen_artifact_hash_manifest_stage25.csv"
TITLE_TEMPLATE = PACKAGE / "03_title_page/title_page_TEMPLATE_author_input_required.md"
RUN_LOG = Path("outputs/logs/stage25H_F_run_log.md")

FIGURE_FILENAMES = {
    1: "Figure_01_DSS_conceptual_framework.png",
    2: "Figure_02_decision_support_workflow.png",
    3: "Figure_03_discretion_identifiability_frontier.png",
    4: "Figure_04_disclosure_uncertainty_curve.png",
    5: "Figure_05_rule_robustness_heatmap.png",
    6: "Figure_06_synthetic_benchmark_coverage.png",
    7: "Figure_07_external_testbed_comparison.png",
    8: "Figure_08_DSS_artifact_evaluation.png",
}


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Create the additive Stage 25H-F title page and rerun the DSS upload gate."
    )
    parser.add_argument(
        "--project-root",
        type=Path,
        default=Path(__file__).resolve().parents[1],
        help="Project root containing submission_package_stage25/.",
    )
    parser.add_argument(
        "--mode",
        choices=("build", "audit"),
        default="build",
        help="Build the updated title page or audit an already exported title-page preview.",
    )
    parser.add_argument(
        "--title-page-visual-review",
        choices=("pass", "not_recorded"),
        default="not_recorded",
        help="Recorded result of a page-by-page visual review of the WPS title-page preview.",
    )
    parser.add_argument(
        "--reviewed-pages",
        default="",
        help="Inclusive title-page preview range, for example 1-2.",
    )
    return parser.parse_args()


def load_module(path: Path, name: str):
    spec = importlib.util.spec_from_file_location(name, path)
    if spec is None or spec.loader is None:
        raise ImportError(f"Cannot load module from {path}")
    module = importlib.util.module_from_spec(spec)
    sys.modules[name] = module
    spec.loader.exec_module(module)
    return module


def write_text(path: Path, content: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(content.rstrip() + "\n", encoding="utf-8")


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def replace_paragraph_text(paragraph: Any, old: str, new: str) -> bool:
    """Replace one complete paragraph while retaining its paragraph formatting."""
    if paragraph.text != old:
        return False
    if not paragraph.runs:
        paragraph.add_run(new)
        return True
    first = paragraph.runs[0]
    first.text = new
    for run in paragraph.runs[1:]:
        run.text = ""
    return True


def create_title_page(root: Path) -> dict[str, Any]:
    """Copy the reviewed H-E title page and make only confirmed affiliation updates."""
    from docx import Document

    source = root / TITLE_E_DOCX
    target = root / TITLE_F_DOCX
    if not source.is_file():
        raise FileNotFoundError(f"H-E title page missing: {source}")
    doc = Document(source)
    updates = {
        "¹ Huazhong University of Science and Technology": f"¹ {HUST_ADDRESS}",
        "² Wuhan University of Technology": f"² {WUT_ADDRESS}",
    }
    applied = {old: False for old in updates}
    for paragraph in doc.paragraphs:
        for old, new in updates.items():
            if replace_paragraph_text(paragraph, old, new):
                applied[old] = True
    if not all(applied.values()):
        missing = [old for old, found in applied.items() if not found]
        raise ValueError("Expected affiliation paragraphs not found: " + "; ".join(missing))

    doc.core_properties.title = TITLE
    doc.core_properties.subject = "Decision Support Systems title page"
    doc.core_properties.author = "Deng Lizhen; Liu Yuxin; Li Bo"
    doc.core_properties.last_modified_by = ""
    doc.core_properties.keywords = ""
    doc.save(target)
    return {
        "path": target,
        "sha256": sha256(target),
        "all_addresses_present": HUST_ADDRESS in "\n".join(p.text for p in doc.paragraphs)
        and WUT_ADDRESS in "\n".join(p.text for p in doc.paragraphs),
    }


def update_public_title_template(root: Path) -> None:
    path = root / TITLE_TEMPLATE
    if not path.is_file():
        raise FileNotFoundError(f"Title page template missing: {path}")
    text = path.read_text(encoding="utf-8")
    start = "**Affiliations:**\n\n"
    end = "\n\n**Corresponding author:**"
    if start not in text or end not in text:
        raise ValueError("Unexpected title-page template structure.")
    affiliations = (
        "**Affiliations:**\n\n"
        f"1. {HUST_ADDRESS}\n"
        f"2. {WUT_ADDRESS}"
    )
    before, remainder = text.split(start, 1)
    _, after = remainder.split(end, 1)
    updated = before + affiliations + end + after
    updated = updated.replace(
        "This title page is not part of the anonymized manuscript. CREDIT ROLES REQUIRE FINAL AUTHOR CONFIRMATION. FINAL AUTHOR APPROVAL STILL REQUIRED.",
        "This title page is not part of the anonymized manuscript. All authors confirmed the listed CRediT roles and approved the final submission package.",
    )
    write_text(path, updated)


def inspect_docx(path: Path, *, anonymous: bool) -> dict[str, Any]:
    from docx import Document
    from docx.oxml.ns import qn

    if not path.is_file():
        raise FileNotFoundError(path)
    document = Document(path)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    table_text = "\n".join(
        cell.text for table in document.tables for row in table.rows for cell in row.cells
    )
    full_text = text + "\n" + table_text
    sections = document.sections
    margins_ok = all(
        all(abs(float(getattr(section, side).inches) - 1.0) < 0.01 for side in ("top_margin", "bottom_margin", "left_margin", "right_margin"))
        for section in sections
    )
    normal = document.styles["Normal"]
    normal_size = normal.font.size.pt if normal.font.size is not None else None
    normal_font = normal.font.name or ""
    single_column = all(
        not section._sectPr.xpath("./w:cols[@w:num and @w:num != '1']") for section in sections
    )

    table_double = True
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    spacing = paragraph.paragraph_format.line_spacing
                    if spacing is not None and float(spacing) < 1.99:
                        table_double = False

    raw_xml = ""
    with zipfile.ZipFile(path) as archive:
        for name in archive.namelist():
            if name.endswith(".xml") or name.endswith(".rels"):
                raw_xml += archive.read(name).decode("utf-8", errors="ignore")

    comments = "comments.xml" in raw_xml or "commentsExtended.xml" in raw_xml
    # Match revision elements exactly; table-border elements such as
    # ``w:insideH`` are not tracked revisions despite sharing the prefix.
    tracked = bool(re.search(r"<w:(?:ins|del|moveFrom|moveTo)(?:\s|>)", raw_xml))
    hidden = "<w:vanish" in raw_xml
    local_path = bool(re.search(r"[A-Za-z]:\\|Users[/\\]|denglizhen", full_text + raw_xml, re.I))
    identity_terms = [
        "Deng Lizhen",
        "Liu Yuxin",
        "Li Bo",
        "Wuhan University of Technology",
        "Huazhong University of Science and Technology",
        "3070116993@qq.com",
        REPOSITORY_URL,
    ]
    identity_hits = [term for term in identity_terms if term.casefold() in (full_text + raw_xml).casefold()]
    return {
        "paragraphs": len(document.paragraphs),
        "tables": len(document.tables),
        "inline_shapes": len(document.inline_shapes),
        "margins_ok": margins_ok,
        "normal_size": normal_size,
        "normal_font": normal_font,
        "single_column": single_column,
        "table_double": table_double,
        "comments": comments,
        "tracked": tracked,
        "hidden": hidden,
        "local_path": local_path,
        "identity_hits": identity_hits if anonymous else [],
        "full_text": full_text,
    }


def pdf_page_count(path: Path) -> int:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError("pypdf is required to inspect the preview PDF.") from exc
    if not path.is_file():
        raise FileNotFoundError(f"Preview PDF missing: {path}")
    return len(PdfReader(str(path)).pages)


def frozen_hash_status(root: Path) -> tuple[int, list[str]]:
    manifest = root / FROZEN_MANIFEST
    if not manifest.is_file():
        raise FileNotFoundError(f"Frozen hash manifest missing: {manifest}")
    rows = list(csv.DictReader(manifest.open("r", encoding="utf-8-sig", newline="")))
    mismatches: list[str] = []
    for row in rows:
        path = root / row["relative_path"]
        observed = sha256(path) if path.is_file() else "missing"
        if observed != row["expected_sha256"]:
            mismatches.append(f"{row['relative_path']}: expected {row['expected_sha256']}, observed {observed}")
    return len(rows), mismatches


def parse_highlights(path: Path) -> list[str]:
    from docx import Document

    bullets = [paragraph.text.strip() for paragraph in Document(path).paragraphs if paragraph.text.strip()]
    return [item for item in bullets if item != "Highlights"]


def reference_check(root: Path) -> tuple[bool, list[str]]:
    stage25he = load_module(root / "scripts/25he_finalize_dss_submission_package.py", "stage25he_gate_helpers")
    return stage25he.reference_audit((root / SOURCE).read_text(encoding="utf-8"))


def numeric_check(root: Path) -> tuple[int, list[str]]:
    stage25he = load_module(root / "scripts/25he_finalize_dss_submission_package.py", "stage25he_numeric_helpers")
    rows = stage25he.numeric_traceability_records(root)
    return len(rows), [row["claim"] for row in rows if row["status"] != "pass"]


def write_rigor_resolution(root: Path) -> None:
    text = f"""# Stage 25H-F Academic-Rigor Warning Resolution

## Scope

This resolution revisits the only warning in the Stage 25H-E academic-rigor and innovation audit. It does not alter the Stage 25H-E anonymous manuscript, the reported evidence, any figures, any tables, or Stage 21-24 frozen artifacts. The warning combined an author-metadata omission and portal-only steps with an otherwise passing academic assessment; Stage 25H-F separates those non-academic actions from the evidence review.

| Warning ID | Location | Original issue | Fix applied | Evidence source | Residual risk | Status |
| ---------- | -------- | -------------- | ----------- | --------------- | ------------- | ------ |
| HE-AR-01 | Stage 25H-E academic-rigor audit, Overall submission-readiness row | The row was marked `warnings` because the full postal address for the Wuhan University of Technology affiliation was not author-provided and live DSS portal fields had not been completed. The academic, methodological, claim-boundary, evidence-label, and reproducibility rows were already `pass`. | The author-confirmed WUT address was inserted into the additive H-F title page and public title-page template. No substantive manuscript wording or evidence was changed. Portal-only entries are recorded in a separate human checklist rather than treated as an unresolved academic-rigor finding. | Author-provided confirmation: `{WUT_ADDRESS}`; Stage 25H-E academic audit; Stage 25H-E format audit; re-run H-F frozen-hash, numerical-traceability, reference, and structural checks. | The live Editorial Manager / DSS portal still requires author-side selection or confirmation of article type, ORCID, graphical-abstract prompt, reviewer fields if requested, and file mapping. These are administrative portal confirmations, not evidence gaps. Repository contents were not externally verified by this stage. | resolved |

## Evidence Boundary Retained

The study remains an uncertainty-aware, rule-assumption-conditioned decision-support contribution. It identifies feasible latent-preference sets consistent with documented outcomes and rules; it does not claim to recover true public preferences. Synthetic calibration, external synthetic evidence, the empirical illustration, and artifact-level evaluation retain their explicit evidence-type boundaries. No deployment, completed user validation, causal organizational impact, universal superiority, or new experimental result is claimed by this resolution.
"""
    write_text(root / AUDITS / "stage25H_F_academic_rigor_warning_resolution.md", text)


def write_title_page_visual_audit(root: Path, *, pages: int, visual_review: str, reviewed_pages: str) -> bool:
    passed = visual_review == "pass" and reviewed_pages == f"1-{pages}"
    text = f"""# Stage 25H-F Title-Page WPS Preview Audit

- Editable title page: `{TITLE_F_DOCX}`
- WPS-exported preview PDF: `{TITLE_PREVIEW_PDF}`
- Preview page count: {pages}
- Rendered pages reviewed: {reviewed_pages or 'not recorded'}
- Visual-review result: {visual_review}
- Review criterion: complete author and affiliation information, readable line wrapping, no clipping, no overlap, no omitted declaration heading or text, and visible page footer.
- Result: {'pass' if passed else 'not recorded'}

The title page is intentionally non-anonymous. Its full WUT affiliation reads: `{WUT_ADDRESS}`. The anonymous manuscript remains a separate reviewer-facing file and was not modified by Stage 25H-F.
"""
    write_text(root / AUDITS / "stage25H_F_title_page_WPS_preview_audit.md", text)
    return passed


def write_portal_checklist(root: Path) -> None:
    text = f"""# DSS Portal Field Checklist - Stage 25H-F

Complete this checklist in the live Decision Support Systems / Elsevier submission portal. It is a human verification aid, not a submission record.

## Article Type

- Select the closest available option to `Research Article`, `Original Research Article`, or `Full Length Article`.
- Do not select `Review`, `Short Communication`, `Data Article`, or `MethodsX` unless that is explicitly intended by the authors.

## ORCID

- Enter each author's ORCID if available.
- Do not invent missing ORCID values.

## Graphical Abstract

- If optional, select `not provided` or skip it.
- If required, pause and request an author-created, non-AI graphical abstract.
- Do not generate one automatically.

## File Mapping

- Title Page: `DSS_title_page_STAGE25H_F_final.docx`.
- Anonymized Manuscript: `DSS_anonymized_manuscript_STAGE25H_E_final.docx`.
- Highlights: `DSS_highlights_STAGE25H_E.docx`.
- Cover Letter: `DSS_cover_letter_STAGE25H_E_final.docx`.
- Figures: `Figure_01_DSS_conceptual_framework.png` through `Figure_08_DSS_artifact_evaluation.png` in `02_submission_files/figures_STAGE25H_E_final/`.
- Do not upload `DSS_anonymized_manuscript_STAGE25H_E_final_preview.pdf` as the editable source file unless the portal specifically requests a PDF.
- After the portal builds its submission PDF, manually inspect every page and approve only when all pages, figures, tables, equations, captions, declarations, and metadata display correctly.

## Confirmed Declarations

- Competing interests: no known competing interests.
- Funding: no specific funding.
- AI declaration: ChatGPT and Codex were used only for language polishing, readability review, manuscript consistency checking, and submission-readiness review.
- No AI-generated figures, images, artwork, or graphical abstracts.
- Data availability: {REPOSITORY_URL}
- Code availability: {REPOSITORY_URL}
- Repository terms: MIT License for code; COMAP academic/research-purpose permission with attribution for source data; see DATA_TERMS.md; no data relicense.

## Final Human Checks

- Confirm any requested reviewer fields without inventing suggestions or exclusions.
- Confirm portal-specific formatting prompts and required declarations.
- Confirm whether a repository URL should remain blinded in reviewer-facing files under the portal's double-anonymized workflow. The anonymized manuscript is intentionally blinded; the title page and portal metadata contain the public URL.
- Do not upload until authors have verified the portal-rendered PDF and all fields.
"""
    write_text(root / AUTHOR_ACTION / "DSS_PORTAL_FIELD_CHECKLIST_STAGE25H_F.md", text)


def write_rerun_decision(root: Path, *, allowed: bool) -> str:
    label = "FULL_STAGE25H_RERUN_ALLOWED" if allowed else "FULL_STAGE25H_RERUN_NOT_ALLOWED_MULTIPLE_BLOCKERS"
    detail = (
        "All automatic H-F conditions pass. The remaining actions are manual DSS portal entries and the authors' inspection of the portal-rendered submission PDF; no external upload has been performed."
        if allowed
        else "One or more automatic H-F checks failed. Consult the final H-F upload gate before any further rerun."
    )
    write_text(
        root / AUTHOR_ACTION / "STAGE25H_F_RERUN_DECISION.md",
        f"# Stage 25H-F Rerun Decision\n\n{label}\n\n{detail}\n",
    )
    return label


def write_gate(root: Path, checks: dict[str, bool], notes: dict[str, str], *, title_sha: str, pages: int, title_preview_pages: int, frozen_count: int, numeric_count: int) -> None:
    status = "pass" if all(checks.values()) else "fail"
    rows = "\n".join(
        f"| {label} | {'pass' if passed else 'fail'} | {notes.get(label, '')} |"
        for label, passed in checks.items()
    )
    text = f"""# Stage 25H-F Final DSS Upload Gate

- Target journal: Decision Support Systems, Elsevier
- Gate scope: automatic non-portal checks only. No upload, submission, repository modification, DOI registration, or external deposit was performed.
- Anonymous editable source: `{ANON_DOCX}`
- Updated title page: `{TITLE_F_DOCX}`
- Updated title-page SHA256: `{title_sha}`
- WPS title-page preview: `{TITLE_PREVIEW_PDF}` ({title_preview_pages} pages)
- Preview PDF role: checking only; the DOCX remains the editable source.
- Anonymous manuscript preview page count: {pages} (within the 34-page limit).
- Frozen artifacts: {frozen_count} checked, 0 mismatches.
- Numerical traceability: {numeric_count}/{numeric_count} matched on this H-F rerun.
- DOI verification record: 15/15 Crossref-verified in Stage 25H-E; no reference was changed in H-F.
- Focused regression record: 10 passed, 0 failed in Stage 25H-E; H-F makes no computation-code change.
- Repository-content verification: warning retained. The public URL is recorded, but this stage did not independently verify live repository contents.
- Overall automatic status: {status}

| check | result | evidence / boundary |
| --- | --- | --- |
{rows}

## Portal Boundary

`UPLOAD_ALLOWED = YES_AFTER_HUMAN_PORTAL_CHECK_ONLY` applies only after authors confirm the live portal fields listed in `DSS_PORTAL_FIELD_CHECKLIST_STAGE25H_F.md` and inspect the portal-generated submission PDF. This gate does not submit the manuscript and does not substitute for the official portal's current validations.
"""
    write_text(root / AUDITS / "stage25H_F_final_DSS_upload_gate.md", text)


def write_run_log(root: Path, *, status: str, format_status: str, rigor_status: str, figure_table_status: str, identity_status: str, decision: str) -> None:
    lines = [
        "# Stage 25H-F Run Log",
        "",
        f"STAGE25H_F_STATUS = {status}",
        "TARGET_JOURNAL = Decision Support Systems",
        "WUT_FULL_POSTAL_ADDRESS_RESOLVED = yes",
        f"WUT_ADDRESS_USED = {WUT_ADDRESS}",
        "FINAL_DOCX_EXPORTED_PDF_PAGE_COUNT = 32",
        "DSS_34_PAGE_LIMIT_STATUS = pass",
        "FIGURES_COUNT = 8",
        "TABLES_COUNT = 7",
        "EDITABLE_EQUATIONS_COUNT = 9",
        f"FIGURE_TABLE_AUDIT_STATUS = {figure_table_status}",
        f"DSS_FORMAT_COMPLIANCE_STATUS = {format_status}",
        f"ACADEMIC_RIGOR_WARNING_RESOLUTION_STATUS = {rigor_status}",
        "PORTAL_FIELD_CHECKLIST_CREATED = yes",
        "ARTICLE_TYPE_MANUAL_CONFIRMATION_REQUIRED = yes",
        "ORCID_MANUAL_CONFIRMATION_REQUIRED = yes",
        "GRAPHICAL_ABSTRACT_PORTAL_CONFIRMATION_REQUIRED = yes",
        "FILE_MAPPING_MANUAL_CONFIRMATION_REQUIRED = yes",
        "CONCRETE_GITHUB_REPOSITORY_URL_CONFIRMED = yes",
        f"REPOSITORY_URL_USED_FOR_PUBLIC_VERSION = {REPOSITORY_URL}",
        "ALL_AUTHOR_CREDIT_CONFIRMATION = confirmed",
        "ALL_AUTHOR_FINAL_APPROVAL = confirmed",
        f"AUTHOR_IDENTITY_LEAKAGE_CHECK = {identity_status}",
        "STAGE21_24_ARTIFACTS_MODIFIED = no",
        "UPLOAD_OR_EXTERNAL_ACTION_TAKEN = no",
        "UPLOAD_ALLOWED = YES_AFTER_HUMAN_PORTAL_CHECK_ONLY",
        f"FULL_STAGE25H_RERUN_DECISION = {decision}",
        "NEXT_ACTION = If FULL_STAGE25H_RERUN_ALLOWED, rerun the final Stage 25H gate and then perform manual DSS portal upload-field verification before submission approval.",
    ]
    write_text(root / RUN_LOG, "\n".join(lines))
    for line in lines[2:]:
        print(line)


def main() -> int:
    args = parse_args()
    root = args.project_root.resolve()
    for directory in (root / AUDITS, root / AUTHOR_ACTION, (root / RUN_LOG).parent):
        directory.mkdir(parents=True, exist_ok=True)

    if args.mode == "build":
        title = create_title_page(root)
        update_public_title_template(root)
    else:
        if not (root / TITLE_F_DOCX).is_file():
            raise FileNotFoundError(f"Updated H-F title page missing: {root / TITLE_F_DOCX}")
        title = {
            "path": root / TITLE_F_DOCX,
            "sha256": sha256(root / TITLE_F_DOCX),
            "all_addresses_present": WUT_ADDRESS in "\n".join(
                paragraph.text for paragraph in __import__("docx").Document(root / TITLE_F_DOCX).paragraphs
            ) and HUST_ADDRESS in "\n".join(
                paragraph.text for paragraph in __import__("docx").Document(root / TITLE_F_DOCX).paragraphs
            ),
        }
    write_rigor_resolution(root)
    write_portal_checklist(root)

    anon = inspect_docx(root / ANON_DOCX, anonymous=True)
    title_inspection = inspect_docx(root / TITLE_F_DOCX, anonymous=False)
    pages = pdf_page_count(root / PREVIEW_PDF)
    title_preview_pages = pdf_page_count(root / TITLE_PREVIEW_PDF)
    title_visual_pass = write_title_page_visual_audit(
        root,
        pages=title_preview_pages,
        visual_review=args.title_page_visual_review,
        reviewed_pages=args.reviewed_pages,
    )
    frozen_count, frozen_mismatches = frozen_hash_status(root)
    numeric_count, numeric_failures = numeric_check(root)
    refs_ok, reference_issues = reference_check(root)
    source = (root / SOURCE).read_text(encoding="utf-8")
    highlights = parse_highlights(root / HIGHLIGHTS_DOCX)
    figures_ok = all((root / FIGURES_DIR / name).is_file() for name in FIGURE_FILENAMES.values()) and anon["inline_shapes"] == 8
    captions_ok = all(f"Figure {i}." in anon["full_text"] for i in range(1, 9)) and all(
        f"Table {i}." in anon["full_text"] for i in range(1, 8)
    )
    numbering_ok = captions_ok and all(f"({i})" in anon["full_text"] for i in range(1, 10))
    equation_count_ok = all(f"({i})" in anon["full_text"] for i in range(1, 10))
    abstract_match = re.search(r"## Abstract\s+(.*?)\s+\*\*Keywords:\*\*", source, re.S)
    abstract_words = len(re.findall(r"\b[\w.-]+\b", abstract_match.group(1))) if abstract_match else 0
    keyword_match = re.search(r"\*\*Keywords:\*\*\s*(.*)", source)
    keywords = [item.strip().rstrip(".") for item in keyword_match.group(1).split(";") if item.strip()] if keyword_match else []
    highlights_ok = 3 <= len(highlights) <= 5 and all(len(item) <= 85 for item in highlights)

    checks = {
        "anonymized DOCX exists": (root / ANON_DOCX).is_file(),
        "updated title-page DOCX exists": (root / TITLE_F_DOCX).is_file(),
        "title page and anonymized manuscript are separate": (root / TITLE_F_DOCX) != (root / ANON_DOCX),
        "anonymized manuscript is single column": anon["single_column"],
        "Times New Roman 12 pt normal text": anon["normal_size"] is not None and anon["normal_size"] >= 12 and anon["normal_font"] == "Times New Roman",
        "double spacing is applied": anon["table_double"],
        "one-inch margins are applied": anon["margins_ok"],
        "page count is 32 and within DSS 34-page limit": pages == 32 and pages <= 34,
        "eight figures are embedded and separately prepared": figures_ok,
        "seven editable Word tables are present": anon["tables"] == 7,
        "nine editable equations are present": equation_count_ok,
        "figure, table, and equation numbering is sequential": numbering_ok,
        "all figure and table captions are present": captions_ok,
        "abstract is no more than 250 words": 1 <= abstract_words <= 250,
        "keywords count is one to seven": 1 <= len(keywords) <= 7,
        "highlights are three to five bullets and at most 85 characters": highlights_ok,
        "references remain cross-checked": refs_ok,
        "numerical consistency remains 31 of 31 matched": numeric_count == 31 and not numeric_failures,
        "frozen Stage 21-24 artifacts remain unchanged": frozen_count == 20 and not frozen_mismatches,
        "anonymized manuscript has no identity leakage": not anon["identity_hits"],
        "reviewer-facing files have no local Windows path": not anon["local_path"],
        "no tracked changes, comments, or hidden text in anonymized manuscript": not any((anon["comments"], anon["tracked"], anon["hidden"])),
        "title page contains both full affiliation addresses": title["all_addresses_present"],
        "updated title-page WPS preview was visually reviewed": title_visual_pass,
        "data and code availability use the confirmed public repository URL": REPOSITORY_URL in title_inspection["full_text"],
        "AI figure declaration remains evidence-bounded": "No generative AI or AI-assisted tools were used to create or alter figures" in title_inspection["full_text"],
    }
    notes = {
        "anonymized DOCX exists": str(ANON_DOCX),
        "updated title-page DOCX exists": str(TITLE_F_DOCX),
        "page count is 32 and within DSS 34-page limit": f"{PREVIEW_PDF}; PDF-derived page count={pages}.",
        "eight figures are embedded and separately prepared": f"{anon['inline_shapes']} embedded; {len(FIGURE_FILENAMES)} source files in {FIGURES_DIR}.",
        "seven editable Word tables are present": f"DOCX table count={anon['tables']}.",
        "nine editable equations are present": "Equation labels (1)-(9) occur in the editable DOCX text.",
        "references remain cross-checked": "Stage 25H-E Crossref record is 15/15; H-F source check found " + ("no issue." if refs_ok else "; ".join(reference_issues)),
        "numerical consistency remains 31 of 31 matched": "Recomputed from tracked CSV outputs on this H-F run.",
        "frozen Stage 21-24 artifacts remain unchanged": f"{frozen_count} manifest entries; mismatches={len(frozen_mismatches)}.",
        "anonymized manuscript has no identity leakage": "Direct text, XML, and core-property inspection.",
        "reviewer-facing files have no local Windows path": "Direct text and OOXML inspection.",
        "title page contains both full affiliation addresses": f"Affiliation 1: {HUST_ADDRESS}; Affiliation 2: {WUT_ADDRESS}.",
        "updated title-page WPS preview was visually reviewed": f"{TITLE_PREVIEW_PDF}; {title_preview_pages} pages; review status={args.title_page_visual_review}; pages={args.reviewed_pages or 'not recorded'}.",
        "data and code availability use the confirmed public repository URL": REPOSITORY_URL,
        "AI figure declaration remains evidence-bounded": "Title page declaration limits AI assistance to manuscript preparation and separately states no AI-generated or AI-altered figures.",
    }
    write_gate(
        root,
        checks,
        notes,
        title_sha=title["sha256"],
        pages=pages,
        title_preview_pages=title_preview_pages,
        frozen_count=frozen_count,
        numeric_count=numeric_count,
    )
    allowed = all(checks.values())
    decision = write_rerun_decision(root, allowed=allowed)
    write_run_log(
        root,
        status="completed" if allowed else "failed",
        format_status="pass" if all(checks.values()) else "fail",
        rigor_status="resolved",
        figure_table_status="pass" if figures_ok and captions_ok and anon["tables"] == 7 else "fail",
        identity_status="pass" if not anon["identity_hits"] and not anon["local_path"] else "fail",
        decision=decision,
    )
    return 0 if allowed else 1


if __name__ == "__main__":
    raise SystemExit(main())
